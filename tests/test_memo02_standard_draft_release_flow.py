from __future__ import annotations

import json
import sys
from datetime import datetime, timezone
from pathlib import Path

import pytest
from docx import Document

from scripts import run_memo02_standard_draft_release_flow as flow
from tests.test_monthly_plan_fact_standard_draft_release_pilot import (
    FORBIDDEN_VISIBLE_ENGLISH_PHRASES,
    PRODUCTION_GENERATOR_MODULES,
    assert_no_forbidden_visible_english,
    fake_blocked_visual_qa,
    fake_pass_visual_qa,
    write_source_files,
)


FORBIDDEN_OUTPUT_PARTS = {
    "01_raw",
    "02_stage",
    "03_marts",
    "04_charts",
    "04_signals",
    "05_evidence",
    "05_llm_package",
    "06_reports",
    "07_qa",
    "99_archive",
    "final",
}


def test_flow_writes_expected_outputs_under_provided_out_dir(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setattr(flow, "diagnose_docx_visual_quality", fake_pass_visual_qa)
    source_paths = write_source_files(tmp_path / "sources")
    out_dir = tmp_path / "draft_out"

    result = flow.run_memo02_standard_draft_release_flow(out_dir, source_paths)

    assert result["status"] == "pass"
    assert Path(result["output_dir"]) == out_dir
    assert (out_dir / "memo_display_contract.json").exists()
    assert (out_dir / "monthly_plan_fact_memo__standard__draft.docx").exists()
    assert (out_dir / "visual_qa" / "defects.json").exists()
    assert (out_dir / "visual_qa" / "diagnostic_report.md").exists()
    assert (out_dir / "release_manifest.json").exists()


def test_flow_default_output_dir_is_ignored_draft_flow_folder() -> None:
    now = datetime(2026, 5, 25, 19, 15, tzinfo=timezone.utc)

    assert flow.default_output_dir(now) == flow.DRAFT_FLOW_ROOT / "20260525_191500Z"
    assert "artifacts/draft_flows/memo02_standard" in str(flow.default_output_dir(now))


def test_flow_does_not_write_to_forbidden_layers(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(flow, "diagnose_docx_visual_quality", fake_pass_visual_qa)
    source_paths = write_source_files(tmp_path / "sources")
    before = sorted(path.relative_to(tmp_path) for path in tmp_path.rglob("*"))

    flow.run_memo02_standard_draft_release_flow(tmp_path / "draft_out", source_paths)

    after = sorted(path.relative_to(tmp_path) for path in tmp_path.rglob("*"))
    created = [path for path in after if path not in before]
    assert created
    assert all(path == Path("draft_out") or path.is_relative_to(Path("draft_out")) for path in created)
    assert all(not FORBIDDEN_OUTPUT_PARTS.intersection(path.parts) for path in created)


def test_flow_docx_visible_body_is_russian_without_forbidden_terms(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setattr(flow, "diagnose_docx_visual_quality", fake_pass_visual_qa)
    source_paths = write_source_files(tmp_path / "sources")

    result = flow.run_memo02_standard_draft_release_flow(tmp_path / "draft_out", source_paths)
    text = docx_text(Path(result["docx_path"]))

    assert "Назначение чернового пилота" in text
    assert "Использование исходного контекста" in text
    assert "Сводка готовности черновика" in text
    assert_no_forbidden_visible_english(text)
    for term in ["DOCX", "Git", "--out", "memo02 standard"]:
        assert term not in text
    assert not any(term in text for term in FORBIDDEN_VISIBLE_ENGLISH_PHRASES)


def test_flow_release_manifest_pass_when_visual_qa_passes(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setattr(flow, "diagnose_docx_visual_quality", fake_pass_visual_qa)
    source_paths = write_source_files(tmp_path / "sources")

    result = flow.run_memo02_standard_draft_release_flow(tmp_path / "draft_out", source_paths)
    manifest = json.loads(Path(result["release_manifest_path"]).read_text(encoding="utf-8"))

    assert manifest["decision"]["release_status"] == "pass"
    assert manifest["qa_status"]["visual_qa_status"] == "pass"
    assert result["manual_approval_required"] is True


def test_flow_blocks_when_visual_qa_blocks(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(flow, "diagnose_docx_visual_quality", fake_blocked_visual_qa)
    source_paths = write_source_files(tmp_path / "sources")

    result = flow.run_memo02_standard_draft_release_flow(tmp_path / "draft_out", source_paths)
    manifest = json.loads(Path(result["release_manifest_path"]).read_text(encoding="utf-8"))

    assert result["status"] == "blocked"
    assert manifest["decision"]["release_status"] == "blocked"
    assert "VQA-001" in manifest["decision"]["release_blockers"]


def test_flow_missing_source_writes_blocked_manifest_without_fake_source_data(tmp_path: Path) -> None:
    source_paths = write_source_files(tmp_path / "sources")
    source_paths.chart_metadata.unlink()

    result = flow.run_memo02_standard_draft_release_flow(tmp_path / "draft_out", source_paths)

    assert result["status"] == "blocked"
    assert result["missing_sources"] == ["chart_metadata"]
    assert not source_paths.chart_metadata.exists()
    assert not Path(result["contract_path"]).exists()
    assert not Path(result["docx_path"]).exists()
    assert Path(result["release_manifest_path"]).exists()


def test_flow_production_generators_are_not_imported_or_called(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setattr(flow, "diagnose_docx_visual_quality", fake_pass_visual_qa)
    source_paths = write_source_files(tmp_path / "sources")
    before = set(sys.modules)

    flow.run_memo02_standard_draft_release_flow(tmp_path / "draft_out", source_paths)

    imported = set(sys.modules) - before
    assert PRODUCTION_GENERATOR_MODULES.isdisjoint(imported)


def docx_text(path: Path) -> str:
    document = Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    return "\n".join(paragraphs + table_cells)
