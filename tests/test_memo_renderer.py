from __future__ import annotations

import base64
from pathlib import Path
import zipfile

import pytest
from docx import Document

from src.memo_display_contract import (
    MemoActionItem,
    MemoBlock,
    MemoChart,
    MemoDisplayContract,
    MemoKpiCard,
    MemoLimitation,
    MemoSection,
    MemoTable,
)
from src.memo_renderer import render_memo_contract_to_docx


def synthetic_contract() -> MemoDisplayContract:
    return MemoDisplayContract(
        memo_id="memo_renderer_test",
        memo_profile="synthetic_profile",
        depth_mode="standard",
        period="2026-04",
        audience="CFO",
        title="Synthetic Budget Memo",
        subtitle="Synthetic renderer fixture",
        status_line="Status: test only",
        sections=[
            MemoSection(
                section_id="summary",
                title="Executive Summary",
                blocks=[
                    MemoBlock(block_type="paragraph", text="Synthetic paragraph text."),
                    MemoBlock(block_type="bullet_list", bullets=["First finding", "Second finding"]),
                    MemoBlock(
                        block_type="kpi_cards",
                        kpi_cards=[MemoKpiCard(label="Revenue", value="100", status="test")],
                    ),
                    MemoBlock(
                        block_type="table",
                        table=MemoTable(headers=["Metric", "Value"], rows=[["Revenue", "100"], ["Cost", "80"]]),
                    ),
                    MemoBlock(block_type="chart", chart=MemoChart(chart_id="chart_01", title="Budget trend")),
                    MemoBlock(
                        block_type="limitation_box",
                        limitation=MemoLimitation(title="Limitations", text="Synthetic limitation text."),
                    ),
                    MemoBlock(
                        block_type="action_table",
                        action_items=[
                            MemoActionItem(
                                action="Review variance",
                                owner="Finance",
                                status="open",
                                marker="candidate",
                                due_date="2026-05-31",
                                evidence_ref="evidence-1",
                            )
                        ],
                    ),
                ],
            )
        ],
        appendix=[
            MemoBlock(
                block_type="evidence_appendix",
                text="Synthetic evidence appendix text.",
                evidence_refs=["evidence-1"],
                appendix_only=True,
            )
        ],
    )


def docx_text(path: Path) -> str:
    document = Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    return "\n".join(paragraphs + table_cells)


def test_renders_synthetic_minimal_contract_to_docx(tmp_path: Path) -> None:
    output_path = tmp_path / "memo.docx"

    result = render_memo_contract_to_docx(synthetic_contract(), output_path)

    assert result == output_path
    assert output_path.exists()
    assert output_path.stat().st_size > 0


def test_docx_contains_title_section_and_paragraph_text(tmp_path: Path) -> None:
    output_path = render_memo_contract_to_docx(synthetic_contract(), tmp_path / "memo.docx")
    text = docx_text(output_path)

    assert "Synthetic Budget Memo" in text
    assert "Executive Summary" in text
    assert "Synthetic paragraph text." in text


def test_docx_contains_table_headers_and_row_values(tmp_path: Path) -> None:
    output_path = render_memo_contract_to_docx(synthetic_contract(), tmp_path / "memo.docx")
    text = docx_text(output_path)

    assert "Metric" in text
    assert "Value" in text
    assert "Revenue" in text
    assert "100" in text


def test_docx_contains_limitation_text(tmp_path: Path) -> None:
    output_path = render_memo_contract_to_docx(synthetic_contract(), tmp_path / "memo.docx")

    assert "Synthetic limitation text." in docx_text(output_path)


def test_docx_contains_action_status_and_marker(tmp_path: Path) -> None:
    output_path = render_memo_contract_to_docx(synthetic_contract(), tmp_path / "memo.docx")
    text = docx_text(output_path)

    assert "open" in text
    assert "кандидат" in text


def test_docx_contains_appendix_and_evidence_text(tmp_path: Path) -> None:
    output_path = render_memo_contract_to_docx(synthetic_contract(), tmp_path / "memo.docx")
    text = docx_text(output_path)

    assert "Приложение / подтверждения" in text
    assert "Synthetic evidence appendix text." in text
    assert "evidence-1" in text


def test_invalid_contract_raises_value_error(tmp_path: Path) -> None:
    contract = MemoDisplayContract(
        memo_id="bad",
        memo_profile="synthetic_profile",
        depth_mode="standard",
        period="2026-04",
        audience="CFO",
        title="",
        status_line="invalid",
        sections=[],
    )

    with pytest.raises(ValueError, match="Invalid MemoDisplayContract"):
        render_memo_contract_to_docx(contract, tmp_path / "bad.docx")

    assert not (tmp_path / "bad.docx").exists()


def test_renderer_writes_only_output_path_in_tmp_path(tmp_path: Path) -> None:
    before = sorted(path.relative_to(tmp_path) for path in tmp_path.rglob("*"))
    output_path = tmp_path / "memo.docx"

    render_memo_contract_to_docx(synthetic_contract(), output_path)

    after = sorted(path.relative_to(tmp_path) for path in tmp_path.rglob("*"))
    assert before == []
    assert after == [Path("memo.docx")]


def test_renderer_embeds_chart_image_when_source_ref_is_image_path(tmp_path: Path) -> None:
    image_path = tmp_path / "chart.png"
    image_path.write_bytes(base64.b64decode("iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="))
    contract = synthetic_contract()
    chart_block = contract.sections[0].blocks[4]
    contract.sections[0].blocks[4] = MemoBlock(
        block_type="chart",
        chart=MemoChart(chart_id=chart_block.chart.chart_id, title=chart_block.chart.title, source_ref=str(image_path)),
    )

    output_path = render_memo_contract_to_docx(contract, tmp_path / "memo.docx")

    with zipfile.ZipFile(output_path) as archive:
        media = [name for name in archive.namelist() if name.startswith("word/media/")]
    assert media
