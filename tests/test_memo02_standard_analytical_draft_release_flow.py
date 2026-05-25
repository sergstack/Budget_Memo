from __future__ import annotations

import base64
import json
import sys
from pathlib import Path

import pytest
from docx import Document

from scripts import run_memo02_standard_analytical_draft_release_flow as flow
from tests.test_monthly_plan_fact_standard_draft_release_pilot import (
    PRODUCTION_GENERATOR_MODULES,
    assert_no_forbidden_visible_english,
    fake_pass_visual_qa,
)


FORBIDDEN_OUTPUT_PARTS = {
    "01_raw",
    "02_stage",
    "03_marts",
    "04_charts",
    "04_signals",
    "05_evidence",
    "05_llm_package",
    "06_reports",
    "07_qa",
    "99_archive",
    "final",
}

FORBIDDEN_VISIBLE_TECHNICAL_TERMS = [
    "memo02",
    "narrative",
    "memo02_standard_final_text",
    "standard_final_chart_metadata",
    "evidence_map",
    "package_qa",
    "chart_id",
    "planning_quality_frequency_impact",
    "Git",
    "--out",
    "memo02 standard",
]

ALLOWED_SOURCE_BUSINESS_TERMS = [
    "ДР offline marketing",
    "Delta EUR",
    "p-fact",
    "refunds",
    "IN/OUT",
]


def test_analytical_flow_creates_expected_outputs_under_out_only(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setattr(flow, "diagnose_docx_visual_quality", fake_pass_visual_qa)
    source_paths = write_analytical_sources(tmp_path / "sources")
    out_dir = tmp_path / "draft_out"

    result = flow.run_memo02_standard_analytical_draft_release_flow(out_dir, source_paths)

    assert result["status"] == "pass"
    assert Path(result["output_dir"]) == out_dir
    assert (out_dir / "memo_display_contract.json").exists()
    assert (out_dir / "monthly_plan_fact_memo__standard__analytical_draft.docx").exists()
    assert (out_dir / "visual_qa" / "defects.json").exists()
    assert (out_dir / "visual_qa" / "diagnostic_report.md").exists()
    assert (out_dir / "release_manifest.json").exists()
    assert (out_dir / "chart_interpretations.json").exists()
    assert (out_dir / "narrative_source.md").exists()


def test_analytical_flow_does_not_write_to_forbidden_layers(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setattr(flow, "diagnose_docx_visual_quality", fake_pass_visual_qa)
    source_paths = write_analytical_sources(tmp_path / "sources")
    before = sorted(path.relative_to(tmp_path) for path in tmp_path.rglob("*"))

    flow.run_memo02_standard_analytical_draft_release_flow(tmp_path / "draft_out", source_paths)

    after = sorted(path.relative_to(tmp_path) for path in tmp_path.rglob("*"))
    created = [path for path in after if path not in before]
    assert created
    assert all(path == Path("draft_out") or path.is_relative_to(Path("draft_out")) for path in created)
    assert all(not FORBIDDEN_OUTPUT_PARTS.intersection(path.parts) for path in created)


def test_analytical_flow_reads_sources_without_modifying_them(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setattr(flow, "diagnose_docx_visual_quality", fake_pass_visual_qa)
    source_paths = write_analytical_sources(tmp_path / "sources")
    source_files = [path for _, path in flow._source_items(source_paths)]
    before = {path: (path.stat().st_mtime_ns, path.stat().st_size) for path in source_files}

    flow.run_memo02_standard_analytical_draft_release_flow(tmp_path / "draft_out", source_paths)

    after = {path: (path.stat().st_mtime_ns, path.stat().st_size) for path in source_files}
    assert after == before


def test_analytical_docx_visible_body_is_russian_without_forbidden_terms(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setattr(flow, "diagnose_docx_visual_quality", fake_pass_visual_qa)
    source_paths = write_analytical_sources(tmp_path / "sources")

    result = flow.run_memo02_standard_analytical_draft_release_flow(tmp_path / "draft_out", source_paths)
    text = docx_text(Path(result["docx_path"]))

    for heading in [
        "Цель записки",
        "Период и статус источников",
        "Ключевые выводы",
        "План-Факт: смысл отклонения",
        "Графики и интерпретации",
        "Ограничения",
        "Следующие действия",
    ]:
        assert heading in text
    assert_no_forbidden_visible_english(text)
    for term in FORBIDDEN_VISIBLE_TECHNICAL_TERMS:
        assert term not in text
    for term in [
        "управленческий текст",
        "финальный текст стандартной записки",
        "карта подтверждений",
        "реестр графиков стандартного пакета",
        "проверка пакета",
    ]:
        assert term in text


def test_source_business_labels_are_preserved_and_do_not_fail_language_gate(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setattr(flow, "diagnose_docx_visual_quality", fake_pass_visual_qa)
    source_paths = write_analytical_sources(tmp_path / "sources")

    result = flow.run_memo02_standard_analytical_draft_release_flow(tmp_path / "draft_out", source_paths)
    text = docx_text(Path(result["docx_path"]))

    for term in ALLOWED_SOURCE_BUSINESS_TERMS:
        assert term in text
    for term in FORBIDDEN_VISIBLE_TECHNICAL_TERMS:
        assert term not in text


def test_grammar_blockers_are_absent_from_visible_body(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setattr(flow, "diagnose_docx_visual_quality", fake_pass_visual_qa)
    source_paths = write_analytical_sources(tmp_path / "sources")

    result = flow.run_memo02_standard_analytical_draft_release_flow(tmp_path / "draft_out", source_paths)
    text = docx_text(Path(result["docx_path"]))

    assert "отрицательную отклонение" not in text
    assert "направление отклонение" not in text
    assert "отрицательное отклонение" in text
    assert "направление отклонения" in text


def test_chart_interpretations_and_metadata_are_reflected_in_contract(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setattr(flow, "diagnose_docx_visual_quality", fake_pass_visual_qa)
    source_paths = write_analytical_sources(tmp_path / "sources")

    result = flow.run_memo02_standard_analytical_draft_release_flow(tmp_path / "draft_out", source_paths)
    contract = json.loads(Path(result["contract_path"]).read_text(encoding="utf-8"))
    interpretations = json.loads((tmp_path / "draft_out" / "chart_interpretations.json").read_text(encoding="utf-8"))

    assert interpretations[0]["chart_id"] == "planning_quality_frequency_impact"
    for key in ["what_shows", "management_meaning", "what_to_check", "limitation"]:
        assert interpretations[0][key]
    assert "частоту сигналов качества планирования" in interpretations[0]["what_shows"]
    assert "массовые плановые расхождения" in interpretations[0]["management_meaning"]
    chart_blocks = [
        block
        for section in contract["sections"]
        for block in section["blocks"]
        if block["block_type"] == "chart"
    ]
    assert chart_blocks
    assert chart_blocks[0]["chart"]["chart_id"] == "planning_quality_frequency_impact"
    assert chart_blocks[0]["chart"]["source_ref"].endswith("planning_quality_frequency_impact.png")


def test_chart_interpretation_blocks_are_business_readable(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setattr(flow, "diagnose_docx_visual_quality", fake_pass_visual_qa)
    source_paths = write_analytical_sources(tmp_path / "sources")

    result = flow.run_memo02_standard_analytical_draft_release_flow(tmp_path / "draft_out", source_paths)
    text = docx_text(Path(result["docx_path"]))

    assert "Что показывает график:" in text
    assert "Управленческий смысл:" in text
    assert "Что проверить:" in text
    assert "Ограничение:" in text
    assert "Идентификатор графика" not in text


def test_release_manifest_pass_when_visual_qa_passes(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setattr(flow, "diagnose_docx_visual_quality", fake_pass_visual_qa)
    source_paths = write_analytical_sources(tmp_path / "sources")

    result = flow.run_memo02_standard_analytical_draft_release_flow(tmp_path / "draft_out", source_paths)
    manifest = json.loads(Path(result["release_manifest_path"]).read_text(encoding="utf-8"))

    assert manifest["decision"]["release_status"] == "pass"
    assert manifest["qa_status"]["visual_qa_status"] == "pass"
    assert result["manual_approval_required"] is True


def test_missing_required_source_writes_blocked_manifest_without_fake_source_data(tmp_path: Path) -> None:
    source_paths = write_analytical_sources(tmp_path / "sources")
    source_paths.standard_chart_metadata.unlink()

    result = flow.run_memo02_standard_analytical_draft_release_flow(tmp_path / "draft_out", source_paths)

    assert result["status"] == "blocked"
    assert result["missing_sources"] == ["standard_chart_metadata"]
    assert not source_paths.standard_chart_metadata.exists()
    assert not Path(result["contract_path"]).exists()
    assert not Path(result["docx_path"]).exists()
    assert Path(result["release_manifest_path"]).exists()


def test_production_generators_are_not_imported_or_called(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setattr(flow, "diagnose_docx_visual_quality", fake_pass_visual_qa)
    source_paths = write_analytical_sources(tmp_path / "sources")
    before = set(sys.modules)

    flow.run_memo02_standard_analytical_draft_release_flow(tmp_path / "draft_out", source_paths)

    imported = set(sys.modules) - before
    assert PRODUCTION_GENERATOR_MODULES.isdisjoint(imported)


def test_default_run_does_not_call_ollama(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(flow, "diagnose_docx_visual_quality", fake_pass_visual_qa)
    source_paths = write_analytical_sources(tmp_path / "sources")
    before = set(sys.modules)

    result = flow.run_memo02_standard_analytical_draft_release_flow(tmp_path / "draft_out", source_paths)

    imported = set(sys.modules) - before
    assert result["status"] == "pass"
    assert "src.run_ollama_memo_pipeline" not in imported


def test_run_ollama_is_explicit_and_blocked_without_safe_implementation(tmp_path: Path) -> None:
    source_paths = write_analytical_sources(tmp_path / "sources")

    result = flow.run_memo02_standard_analytical_draft_release_flow(
        tmp_path / "draft_out",
        source_paths,
        run_ollama=True,
    )
    manifest = json.loads(Path(result["release_manifest_path"]).read_text(encoding="utf-8"))

    assert result["status"] == "blocked"
    assert "ollama_generation_not_implemented_in_draft_flow_v1" in manifest["decision"]["release_blockers"]
    assert not Path(result["docx_path"]).exists()


def docx_text(path: Path) -> str:
    document = Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    return "\n".join(paragraphs + table_cells)


def write_analytical_sources(root: Path) -> flow.Memo02AnalyticalSourcePaths:
    base = root / "memo02"
    final_md = base / "final" / "02_monthly_plan_fact_memo__standard__2026-04__final.md"
    ollama_md = base / "ollama_variants" / "02_monthly_plan_fact_memo__standard__ollama_revisor.md"
    package_qa = base / "qa" / "package_qa.md"
    evidence_map = base / "evidence" / "evidence_map.csv"
    chart_dir = base / "charts" / "standard_final"
    image_dir = chart_dir / "images"
    standard_chart_metadata = chart_dir / "standard_final_chart_metadata.json"
    chart_manifest = base / "charts" / "chart_manifest.json"
    chart_metadata_csv = base / "charts" / "chart_metadata.csv"
    for path in [final_md, ollama_md, package_qa, evidence_map, standard_chart_metadata, chart_manifest, chart_metadata_csv]:
        path.parent.mkdir(parents=True, exist_ok=True)
    image_dir.mkdir(parents=True, exist_ok=True)
    image_path = image_dir / "planning_quality_frequency_impact.png"
    image_path.write_bytes(base64.b64decode("iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="))
    final_md.write_text(
        "# Ежемесячная аналитическая записка План-Факт за апрель 2026\n\n"
        "## Главное за период\n\n"
        "Период близок к плану на уровне общего итога, но структура отклонений требует управленческого просмотра.\n\n"
        "## Ключевые показатели\n\n"
        "- План и факт сопоставляются за апрель 2026.\n"
        "- По статье ДР offline marketing виден сигнал Delta EUR; p-fact, refunds и IN/OUT остаются исходными бизнес-метками, а отрицательную отклонение и направление отклонение требуют проверки без вывода о причине.\n"
        "- Основной маршрут проверки проходит через статьи, ЦФО, факт без плана и план без факта.\n"
        "- Причины отклонений не подтверждаются без комментариев владельцев бюджета.\n",
        encoding="utf-8",
    )
    ollama_md.write_text(final_md.read_text(encoding="utf-8"), encoding="utf-8")
    package_qa.write_text("# Проверка пакета\n\nqa_status: pass\n", encoding="utf-8")
    evidence_map.write_text("claim_id,source\nM02-001,accepted package\n", encoding="utf-8")
    chart_rows = [
        {
            "chart_id": "planning_quality_frequency_impact",
            "title_ru": "Качество планирования: частота и ABS влияние, 2026-04",
            "image_path": str(image_path),
            "period": "2026-04",
            "source": "принятый пакет данных: срезы исполнения и разрывов",
            "limitation": "ABS показывает масштаб, а Delta не подтверждает причину.",
        }
    ]
    standard_chart_metadata.write_text(json.dumps(chart_rows, ensure_ascii=False, indent=2), encoding="utf-8")
    chart_manifest.write_text(json.dumps(chart_rows, ensure_ascii=False, indent=2), encoding="utf-8")
    chart_metadata_csv.write_text("chart_id,title\nplanning_quality_frequency_impact,Качество планирования\n", encoding="utf-8")
    return flow.Memo02AnalyticalSourcePaths(
        final_md=final_md,
        ollama_revisor_md=ollama_md,
        package_qa=package_qa,
        evidence_map=evidence_map,
        standard_chart_metadata=standard_chart_metadata,
        chart_manifest=chart_manifest,
        chart_metadata_csv=chart_metadata_csv,
    )
