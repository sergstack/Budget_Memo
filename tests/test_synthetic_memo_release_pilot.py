from __future__ import annotations

import json
import sys
from pathlib import Path

import pytest
from docx import Document

from scripts import run_synthetic_memo_release_pilot as pilot


FORBIDDEN_VISIBLE_ENGLISH_PHRASES = [
    "Draft Pilot Scope",
    "Source Context Summary",
    "Chart placeholder",
    "Source files are read-only",
    "No production DOCX generator is called",
    "No raw, stage, mart, chart",
    "Release Manifest",
    "Synthetic data only",
    "Status: draft",
    "Period:",
    "Audience:",
]

PRODUCTION_GENERATOR_MODULES = {
    "src.regenerate_clean_memo_narratives",
    "src.build_docx_report",
    "src.polish_docx_report",
    "src.build_depth_mode_outputs",
}


def test_pilot_creates_docx_visual_qa_and_release_manifest(tmp_path: Path) -> None:
    result = pilot.run_synthetic_memo_release_pilot(tmp_path)

    assert Path(result["docx_path"]).exists()
    assert Path(result["visual_qa_path"]).exists()
    assert Path(result["release_manifest_path"]).exists()


def test_release_manifest_references_generated_docx_and_visual_qa(tmp_path: Path) -> None:
    result = pilot.run_synthetic_memo_release_pilot(tmp_path)
    manifest = json.loads(Path(result["release_manifest_path"]).read_text(encoding="utf-8"))

    assert manifest["artifact_paths"]["docx_path"] == result["docx_path"]
    assert manifest["artifact_paths"]["visual_qa_path"] == result["visual_qa_path"]
    assert manifest["artifact_paths"]["release_manifest_path"] == result["release_manifest_path"]


def test_synthetic_pilot_docx_visible_body_has_no_forbidden_english_phrases(tmp_path: Path) -> None:
    result = pilot.run_synthetic_memo_release_pilot(tmp_path)
    text = docx_text(Path(result["docx_path"]))

    assert "Синтетический пилот выпуска записки" in text
    for phrase in FORBIDDEN_VISIBLE_ENGLISH_PHRASES:
        assert phrase not in text


def test_soffice_unavailable_blocks_release_manifest(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(pilot, "diagnose_docx_visual_quality", fake_blocked_visual_qa)

    result = pilot.run_synthetic_memo_release_pilot(tmp_path)
    manifest = json.loads(Path(result["release_manifest_path"]).read_text(encoding="utf-8"))

    assert manifest["decision"]["release_status"] == "blocked"
    assert manifest["qa_status"]["visual_qa_status"] == "blocked"
    assert "VQA-001" in manifest["decision"]["release_blockers"]


def test_pilot_writes_only_under_out_dir(tmp_path: Path) -> None:
    out_dir = tmp_path / "pilot"
    before = sorted(path.relative_to(tmp_path) for path in tmp_path.rglob("*"))

    pilot.run_synthetic_memo_release_pilot(out_dir)

    after = sorted(path.relative_to(tmp_path) for path in tmp_path.rglob("*"))
    created = [path for path in after if path not in before]
    assert created
    assert all(path == Path("pilot") or path.is_relative_to(Path("pilot")) for path in created)


def test_production_generators_are_not_imported_or_called(tmp_path: Path) -> None:
    before = set(sys.modules)

    pilot.run_synthetic_memo_release_pilot(tmp_path)

    imported = set(sys.modules) - before
    assert PRODUCTION_GENERATOR_MODULES.isdisjoint(imported)


def docx_text(path: Path) -> str:
    document = Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    return "\n".join(paragraphs + table_cells)


def fake_blocked_visual_qa(docx_path: Path, out_dir: Path, soffice_bin: str | None = None) -> dict:
    out_dir.mkdir(parents=True, exist_ok=True)
    payload = {
        "verdicts": {
            "docx_render_status": "blocked",
            "visual_layout_status": "pass",
            "executive_readability_status": "pass",
            "publishing_hygiene_status": "pass",
            "overall_visual_release_status": "blocked",
        },
        "release_blockers": [{"id": "VQA-001", "description": "LibreOffice/soffice is unavailable."}],
        "defects": [],
    }
    (out_dir / "defects.json").write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    (out_dir / "diagnostic_report.md").write_text("# fake visual qa\n", encoding="utf-8")
    return payload
