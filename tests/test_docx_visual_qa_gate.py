from __future__ import annotations

import json
import subprocess
import sys
from pathlib import Path

from docx import Document

from scripts.diagnose_docx_visual_quality import diagnose_docx_visual_quality


REQUIRED_STATUS_FIELDS = {
    "docx_render_status",
    "visual_layout_status",
    "executive_readability_status",
    "publishing_hygiene_status",
    "overall_visual_release_status",
    "release_blockers",
    "defects",
}


def write_docx(path: Path, text: str = "Synthetic memo text") -> Path:
    document = Document()
    document.add_heading("Synthetic memo", level=0)
    document.add_paragraph(text)
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = "Value"
    document.save(path)
    return path


def load_defects(out_dir: Path) -> dict:
    return json.loads((out_dir / "defects.json").read_text(encoding="utf-8"))


def test_valid_synthetic_docx_writes_diagnostics(tmp_path: Path) -> None:
    docx_path = write_docx(tmp_path / "memo.docx")
    out_dir = tmp_path / "diagnostics"

    result = diagnose_docx_visual_quality(docx_path, out_dir)

    assert (out_dir / "defects.json").exists()
    assert (out_dir / "diagnostic_report.md").exists()
    assert result["metrics"]["paragraph_count"] >= 1
    assert result["metrics"]["tables_count"] == 1


def test_defects_json_has_required_status_fields(tmp_path: Path) -> None:
    docx_path = write_docx(tmp_path / "memo.docx")
    out_dir = tmp_path / "diagnostics"

    diagnose_docx_visual_quality(docx_path, out_dir)

    assert REQUIRED_STATUS_FIELDS.issubset(load_defects(out_dir))


def test_missing_docx_path_returns_blocked_status(tmp_path: Path) -> None:
    out_dir = tmp_path / "diagnostics"

    result = diagnose_docx_visual_quality(tmp_path / "missing.docx", out_dir)

    assert result["overall_visual_release_status"] == "blocked"
    assert "docx_missing" in result["release_blockers"]
    assert (out_dir / "defects.json").exists()


def test_local_debug_path_inside_docx_is_reported(tmp_path: Path) -> None:
    docx_path = write_docx(tmp_path / "memo.docx", "Debug path /Users/example/project/03_marts/file.parquet")
    out_dir = tmp_path / "diagnostics"

    result = diagnose_docx_visual_quality(docx_path, out_dir)

    defect_ids = {defect["id"] for defect in result["defects"]}
    assert "local_debug_path_present" in defect_ids
    assert result["publishing_hygiene_status"] == "blocked"


def test_source_docx_is_not_modified(tmp_path: Path) -> None:
    docx_path = write_docx(tmp_path / "memo.docx")
    before = (docx_path.stat().st_size, docx_path.stat().st_mtime_ns)

    diagnose_docx_visual_quality(docx_path, tmp_path / "diagnostics")

    after = (docx_path.stat().st_size, docx_path.stat().st_mtime_ns)
    assert after == before


def test_outputs_are_written_only_under_out_dir(tmp_path: Path) -> None:
    docx_path = write_docx(tmp_path / "memo.docx")
    out_dir = tmp_path / "diagnostics"
    before = sorted(path.relative_to(tmp_path) for path in tmp_path.rglob("*"))

    diagnose_docx_visual_quality(docx_path, out_dir)

    after = sorted(path.relative_to(tmp_path) for path in tmp_path.rglob("*"))
    created = [path for path in after if path not in before]
    assert created
    assert all(path == Path("diagnostics") or path.is_relative_to(Path("diagnostics")) for path in created)


def test_script_import_and_cli_missing_docx_have_no_side_effects_outside_out(tmp_path: Path) -> None:
    out_dir = tmp_path / "diagnostics"
    missing_docx = tmp_path / "missing.docx"
    command = [
        sys.executable,
        "scripts/diagnose_docx_visual_quality.py",
        "--docx",
        str(missing_docx),
        "--out",
        str(out_dir),
    ]

    completed = subprocess.run(command, cwd=Path.cwd(), capture_output=True, text=True, check=False)

    assert completed.returncode == 1
    assert (out_dir / "defects.json").exists()
    assert (out_dir / "diagnostic_report.md").exists()
