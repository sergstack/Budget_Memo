from __future__ import annotations

import argparse
import json
import shutil
import subprocess
import zipfile
from pathlib import Path
from typing import Any

from docx import Document


STATUS_FIELDS = [
    "docx_render_status",
    "visual_layout_status",
    "executive_readability_status",
    "publishing_hygiene_status",
    "overall_visual_release_status",
]

LOCAL_PATH_MARKERS = ["/Users/", "C:\\", "/tmp/", "src/", "03_marts/", "07_qa/"]
TRACKED_CHANGE_MARKERS = ["<w:ins", "<w:del", "<w:moveFrom", "<w:moveTo"]


def diagnose_docx_visual_quality(docx_path: Path, out_dir: Path) -> dict[str, Any]:
    docx_path = Path(docx_path)
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    result: dict[str, Any] = {
        "docx_path": str(docx_path),
        "out_dir": str(out_dir),
        "docx_render_status": "pass",
        "visual_layout_status": "pass",
        "executive_readability_status": "pass",
        "publishing_hygiene_status": "pass",
        "overall_visual_release_status": "pass",
        "release_blockers": [],
        "defects": [],
        "metrics": {
            "embedded_media_count": 0,
            "tables_count": 0,
            "paragraph_count": 0,
            "text_length": 0,
        },
    }

    if not docx_path.exists():
        _add_defect(result, "docx_missing", "blocker", f"DOCX does not exist: {docx_path}")
        _finalize_and_write(result, out_dir)
        return result
    if docx_path.stat().st_size == 0:
        _add_defect(result, "docx_empty", "blocker", f"DOCX is empty: {docx_path}")
        _finalize_and_write(result, out_dir)
        return result

    package_text = ""
    try:
        with zipfile.ZipFile(docx_path) as package:
            names = package.namelist()
            result["metrics"]["embedded_media_count"] = len([name for name in names if name.startswith("word/media/")])
            package_text = _read_relevant_package_text(package, names)
            _check_package_hygiene(names, package_text, result)
    except zipfile.BadZipFile:
        _add_defect(result, "docx_unreadable_package", "blocker", "DOCX package cannot be opened as a zip archive")
        _finalize_and_write(result, out_dir)
        return result

    try:
        document = Document(docx_path)
    except Exception as exc:  # pragma: no cover - exact python-docx exceptions vary by package defect.
        _add_defect(result, "docx_unreadable_document", "blocker", f"DOCX cannot be opened by python-docx: {exc}")
        _finalize_and_write(result, out_dir)
        return result

    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    visible_text = "\n".join(paragraphs + table_cells)
    result["metrics"]["tables_count"] = len(document.tables)
    result["metrics"]["paragraph_count"] = len(document.paragraphs)
    result["metrics"]["text_length"] = len(visible_text.strip())

    if not visible_text.strip():
        _add_defect(result, "docx_no_visible_text", "blocker", "DOCX has no visible paragraph or table text")
    _check_local_path_markers(visible_text + "\n" + package_text, result)
    _attempt_render(docx_path, out_dir, result)
    _finalize_and_write(result, out_dir)
    return result


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Run read-only DOCX visual/publishing diagnostics. Does not run live Kestra."
    )
    parser.add_argument("--docx", required=True, type=Path)
    parser.add_argument("--out", required=True, type=Path)
    args = parser.parse_args()

    result = diagnose_docx_visual_quality(args.docx, args.out)
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 1 if result["overall_visual_release_status"] == "blocked" else 0


def _read_relevant_package_text(package: zipfile.ZipFile, names: list[str]) -> str:
    chunks: list[str] = []
    for name in names:
        if not (name.endswith(".xml") or name.endswith(".rels")):
            continue
        try:
            chunks.append(package.read(name).decode("utf-8", errors="ignore"))
        except KeyError:
            continue
    return "\n".join(chunks)


def _check_package_hygiene(names: list[str], package_text: str, result: dict[str, Any]) -> None:
    comment_parts = [name for name in names if "comments" in name.lower()]
    if comment_parts:
        _add_defect(result, "docx_comments_present", "blocker", "DOCX package contains comments XML parts")
    if any(marker in package_text for marker in TRACKED_CHANGE_MARKERS):
        _add_defect(result, "docx_tracked_changes_present", "blocker", "DOCX package contains tracked change markers")


def _check_local_path_markers(text: str, result: dict[str, Any]) -> None:
    for marker in LOCAL_PATH_MARKERS:
        if marker in text:
            _add_defect(result, "local_debug_path_present", "blocker", f"DOCX contains local/debug path marker: {marker}")


def _attempt_render(docx_path: Path, out_dir: Path, result: dict[str, Any]) -> None:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        _add_defect(result, "soffice_unavailable", "blocker", "LibreOffice/soffice is unavailable for DOCX to PDF render")
        return

    logs_dir = out_dir / "logs"
    logs_dir.mkdir(parents=True, exist_ok=True)
    command = [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(out_dir), str(docx_path)]
    completed = subprocess.run(command, capture_output=True, text=True, timeout=60, check=False)
    (logs_dir / "libreoffice_stdout.log").write_text(completed.stdout, encoding="utf-8")
    (logs_dir / "libreoffice_stderr.log").write_text(completed.stderr, encoding="utf-8")

    expected_pdf = out_dir / f"{docx_path.stem}.pdf"
    if completed.returncode != 0 or not expected_pdf.exists() or expected_pdf.stat().st_size == 0:
        _add_defect(result, "docx_to_pdf_render_failed", "blocker", "LibreOffice DOCX to PDF conversion failed")


def _add_defect(result: dict[str, Any], defect_id: str, severity: str, message: str) -> None:
    result["defects"].append({"id": defect_id, "severity": severity, "message": message})
    if severity == "blocker" and defect_id not in result["release_blockers"]:
        result["release_blockers"].append(defect_id)


def _finalize_and_write(result: dict[str, Any], out_dir: Path) -> None:
    blocker_ids = set(result["release_blockers"])
    if {"docx_missing", "docx_empty", "docx_unreadable_package", "docx_unreadable_document", "soffice_unavailable", "docx_to_pdf_render_failed"} & blocker_ids:
        result["docx_render_status"] = "blocked"
    if {"local_debug_path_present", "docx_comments_present", "docx_tracked_changes_present"} & blocker_ids:
        result["publishing_hygiene_status"] = "blocked"
    if "docx_no_visible_text" in blocker_ids:
        result["executive_readability_status"] = "blocked"
        result["visual_layout_status"] = "blocked"

    statuses = [result[field] for field in STATUS_FIELDS if field != "overall_visual_release_status"]
    if any(status == "blocked" for status in statuses):
        result["overall_visual_release_status"] = "blocked"
    elif any(status == "revise" for status in statuses):
        result["overall_visual_release_status"] = "revise"
    else:
        result["overall_visual_release_status"] = "pass"

    (out_dir / "defects.json").write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    (out_dir / "diagnostic_report.md").write_text(_markdown_report(result), encoding="utf-8")


def _markdown_report(result: dict[str, Any]) -> str:
    lines = [
        "# DOCX Visual QA Diagnostic Report",
        "",
        f"- DOCX: `{result['docx_path']}`",
        f"- Overall visual release status: `{result['overall_visual_release_status']}`",
        f"- DOCX render status: `{result['docx_render_status']}`",
        f"- Visual layout status: `{result['visual_layout_status']}`",
        f"- Executive readability status: `{result['executive_readability_status']}`",
        f"- Publishing hygiene status: `{result['publishing_hygiene_status']}`",
        "",
        "## Metrics",
        "",
    ]
    for key, value in result["metrics"].items():
        lines.append(f"- {key}: {value}")
    lines.extend(["", "## Release Blockers", ""])
    if result["release_blockers"]:
        for blocker in result["release_blockers"]:
            lines.append(f"- {blocker}")
    else:
        lines.append("- none")
    lines.extend(["", "## Defects", ""])
    if result["defects"]:
        for defect in result["defects"]:
            lines.append(f"- `{defect['severity']}` `{defect['id']}`: {defect['message']}")
    else:
        lines.append("- none")
    lines.append("")
    return "\n".join(lines)


if __name__ == "__main__":
    raise SystemExit(main())
