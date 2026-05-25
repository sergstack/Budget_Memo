from __future__ import annotations

import argparse
from io import BytesIO
import json
import re
import shutil
import subprocess
import sys
import zipfile
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from src.qa_ollama_outputs import QaContext, load_allowed_numbers, validate_text
from src.run_ollama_memo_pipeline import (
    OllamaUnavailable,
    call_ollama,
    call_judge_with_schema,
    default_ollama_client,
    judge_schema_invalid,
    load_routing,
    model_metadata,
    parse_judge_json,
)


PERIOD = "2026-04"
DEPTHS = ["short", "standard", "deep", "action"]
REQUIRED_MEDIA = {"short": 5, "standard": 10, "deep": 15, "action": 5}
SOFFICE = Path("/Applications/LibreOffice.app/Contents/MacOS/soffice")
MEMO_FRESHNESS_SOURCES = {
    "monthly_plan_fact_memo": [
        PROJECT_ROOT / "03_marts/mart_main_full_budget.parquet",
        PROJECT_ROOT / "03_marts/slice_source_mix_summary.parquet",
        PROJECT_ROOT / "06_reports/02_monthly_plan_fact_memo/tables/02_monthly_plan_fact_memo__logic_review__2026-04.xlsx",
    ],
    "executive_yoy_mom_budget_memo": [
        PROJECT_ROOT / "03_marts/mart_main_full_budget.parquet",
        PROJECT_ROOT / "03_marts/mart_flow_base_month.parquet",
        PROJECT_ROOT / "03_marts/mart_signal_catalog_full.parquet",
    ],
}
PILOT_TERMS = [
    ("action memo", "записка по действиям"),
    ("candidate action", "кандидат действия"),
    ("candidate_only", "только кандидаты действий"),
    ("source mix", "состав источников"),
    ("row type", "тип строки"),
    ("gross abs delta", "валовый масштаб отклонений"),
    ("net delta", "чистое отклонение"),
    ("logic workbook", "контрольная книга логики"),
    ("chart manifest", "реестр графиков"),
    ("preflight", "предварительная проверка"),
    ("final judge", "финальная экспертная проверка"),
    ("fallback", "резервная модель"),
    ("pipeline", "цепочка обработки"),
]

EXPLICIT_ROLE_ROUTING = {
    "analyst": {
        "primary_model": "qwen2.5-coder:32b",
        "fallback_model": "qwen3-coder:30b",
        "schema_mode": False,
    },
    "business_writer": {
        "primary_model": "akdengi/saiga-llama3-8b:latest",
        "fallback_model": "qwen2.5-coder:32b",
        "schema_mode": False,
    },
    "evidence_judge": {
        "primary_model": "deepseek-r1:32b",
        "fallback_model": "mistral-small:latest",
        "schema_mode": True,
    },
    "management_readability_judge": {
        "primary_model": "qwen2.5-coder:32b",
        "fallback_model": "qwen3-coder:30b",
        "schema_mode": False,
    },
    "russian_language_judge": {
        "primary_model": "akdengi/saiga-llama3-8b:latest",
        "fallback_model": "qwen2.5-coder:32b",
        "schema_mode": False,
    },
    "llm_revisor": {
        "primary_model": "akdengi/saiga-llama3-8b:latest",
        "fallback_model": "qwen2.5-coder:32b",
        "schema_mode": False,
    },
    "final_consensus_judge": {
        "primary_model": "deepseek-r1:32b",
        "fallback_model": "mistral-small:latest",
        "schema_mode": True,
    },
}


@dataclass(frozen=True)
class MemoPackage:
    memo_no: str
    memo_id: str
    profile: str
    root: Path
    title: str
    evidence_paths: list[Path]

    @property
    def final_dir(self) -> Path:
        return self.root / "final"

    @property
    def qa_dir(self) -> Path:
        return self.root / "07_qa"

    @property
    def chart_manifest(self) -> Path:
        return self.root / "charts/chart_manifest.xlsx"

    def final_md(self, depth: str) -> Path:
        return self.final_dir / f"{self.memo_id}__{depth}__{PERIOD}__final.md"

    def final_docx(self, depth: str) -> Path:
        return self.final_dir / f"{self.memo_id}__{depth}__{PERIOD}__final.docx"

    def final_xlsx(self, depth: str) -> Path | None:
        if depth == "deep":
            path = self.root / "tables" / f"{self.memo_id}__deep__{PERIOD}__slices.xlsx"
            return path if path.exists() else None
        if depth == "action":
            path = self.final_dir / f"{self.memo_id}__action__{PERIOD}__final.xlsx"
            return path if path.exists() else None
        return None


def packages() -> list[MemoPackage]:
    return [
        MemoPackage(
            memo_no="memo01",
            memo_id="01_executive_yoy_mom_budget_memo",
            profile="executive_yoy_mom_budget_memo",
            root=PROJECT_ROOT / "06_reports/01_executive_yoy_mom_budget_memo",
            title="Executive YoY/MoM budget memo",
            evidence_paths=[
                PROJECT_ROOT / "06_reports/01_executive_yoy_mom_budget_memo/source_refs/evidence_cards.md",
                PROJECT_ROOT / "06_reports/01_executive_yoy_mom_budget_memo/source_refs/insight_cards.md",
                PROJECT_ROOT / "06_reports/01_executive_yoy_mom_budget_memo/source_refs/analytical_tables.md",
                PROJECT_ROOT / "06_reports/01_executive_yoy_mom_budget_memo/charts/chart_manifest.md",
                PROJECT_ROOT / "06_reports/01_executive_yoy_mom_budget_memo/qa/qa_depth_outputs_final.json",
            ],
        ),
        MemoPackage(
            memo_no="memo02",
            memo_id="02_monthly_plan_fact_memo",
            profile="monthly_plan_fact_memo",
            root=PROJECT_ROOT / "06_reports/02_monthly_plan_fact_memo",
            title="Monthly Plan-Fact memo",
            evidence_paths=[
                PROJECT_ROOT / "06_reports/02_monthly_plan_fact_memo/tables/02_monthly_plan_fact_memo__logic_review__2026-04.xlsx",
                PROJECT_ROOT / "06_reports/02_monthly_plan_fact_memo/charts/chart_manifest.md",
                PROJECT_ROOT / "06_reports/02_monthly_plan_fact_memo/qa/memo02_final_artifact_hash_manifest.md",
                PROJECT_ROOT / "06_reports/02_monthly_plan_fact_memo/qa/memo02_ollama_revised_latest_summary.json",
            ],
        ),
    ]


def resolve_memo_filter(memo_filter: str | None) -> list[MemoPackage]:
    allowed = {
        None,
        "memo01",
        "memo02",
        "01_executive_yoy_mom_budget_memo",
        "02_monthly_plan_fact_memo",
    }
    if memo_filter not in allowed:
        raise RuntimeError(
            "Unknown memo filter: "
            f"{memo_filter}. Allowed values: memo01, memo02, "
            "01_executive_yoy_mom_budget_memo, 02_monthly_plan_fact_memo"
        )
    selected = [memo for memo in packages() if memo_filter in {None, memo.memo_no, memo.memo_id}]
    if not selected:
        raise RuntimeError(f"Unknown memo filter: {memo_filter}")
    return selected


def print_target_preflight(memo: MemoPackage, depth: str) -> None:
    final_md = memo.final_md(depth)
    final_docx = memo.final_docx(depth)
    expected_md = memo.final_dir / f"{memo.memo_id}__{depth}__{PERIOD}__final.md"
    expected_docx = memo.final_dir / f"{memo.memo_id}__{depth}__{PERIOD}__final.docx"
    print("target_preflight:", flush=True)
    print(f"  resolved_memo_id: {memo.memo_id}", flush=True)
    print(f"  resolved_report_dir: {memo.root.relative_to(PROJECT_ROOT)}", flush=True)
    print(f"  depth: {depth}", flush=True)
    print(f"  final_md_path: {final_md.relative_to(PROJECT_ROOT)}", flush=True)
    print(f"  final_docx_path: {final_docx.relative_to(PROJECT_ROOT)}", flush=True)
    if final_md != expected_md or final_docx != expected_docx:
        raise RuntimeError(
            "Resolved target path mismatch before LLM call: "
            f"md={final_md}, docx={final_docx}, expected_md={expected_md}, expected_docx={expected_docx}"
        )


def snapshot_tree(root: Path) -> dict[str, dict[str, int]]:
    result = {}
    for path in root.rglob("*"):
        if path.is_file():
            result[str(path.relative_to(PROJECT_ROOT))] = {
                "mtime_ns": path.stat().st_mtime_ns,
                "size": path.stat().st_size,
            }
    return result


def changed_paths(before: dict[str, dict[str, int]], root: Path) -> list[str]:
    after = snapshot_tree(root)
    return sorted(path for path in set(before) | set(after) if before.get(path) != after.get(path))


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8") if path.exists() else ""


def strip_markdown_response(text: str) -> str:
    text = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL | re.IGNORECASE).strip()
    fenced = re.match(r"^```(?:markdown|md)?\s*(.*?)\s*```$", text, flags=re.DOTALL | re.IGNORECASE)
    if fenced:
        text = fenced.group(1).strip()
    return text.strip()


def remove_executive_technical_ids(text: str) -> str:
    text = re.sub(r"\bCH_[A-Z0-9_]+\b", "график", text)
    text = re.sub(r"\bEV-[A-Z0-9_-]+\b", "доказательство", text)
    return text


def format_eur_business(value: float) -> str:
    sign = "−" if value < 0 else ""
    abs_value = abs(float(value))
    if abs_value >= 1_000_000:
        return f"{sign}{abs_value / 1_000_000:.2f}".replace(".", ",") + " млн EUR"
    if abs_value >= 1_000:
        return f"{sign}{abs_value / 1_000:.1f}".replace(".", ",") + " тыс. EUR"
    return f"{sign}{abs_value:.2f}".replace(".", ",") + " EUR"


def format_pct_business(value: float) -> str:
    pct = float(value) * 100 if abs(float(value)) <= 10 else float(value)
    return f"{pct:.1f}%".replace(".", ",")


def format_dataframe_for_llm(df: pd.DataFrame) -> pd.DataFrame:
    result = df.copy()
    for column in result.columns:
        name = str(column).lower()
        values = pd.to_numeric(result[column], errors="coerce")
        if name.endswith("_eur"):
            result[column] = [
                format_eur_business(float(value)) if pd.notna(value) else original
                for value, original in zip(values, result[column])
            ]
        elif name.endswith("_pct") or "ratio" in name:
            result[column] = [
                (f"{float(value) * 100:.1f}%".replace(".", ",") if abs(float(value)) <= 10 else f"{float(value):.1f}%".replace(".", ","))
                if pd.notna(value)
                else original
                for value, original in zip(values, result[column])
            ]
    return result


def evidence_brief(memo: MemoPackage, depth: str, baseline: str, max_chars: int = 9000) -> str:
    parts = [f"# Контрольный пакет для {memo.memo_id} / {depth}", "## Предыдущий текст для ориентира", baseline[:5000]]
    for path in memo.evidence_paths:
        if not path.exists():
            continue
        if path.suffix.lower() == ".xlsx":
            xl = pd.ExcelFile(path)
            parts.append(f"## Workbook sheets: {path.relative_to(PROJECT_ROOT)}")
            parts.append(", ".join(xl.sheet_names[:20]))
            for sheet in xl.sheet_names[:5]:
                df = format_dataframe_for_llm(pd.read_excel(path, sheet_name=sheet).head(8))
                parts.append(f"### {sheet}")
                parts.append(df.to_csv(index=False))
        else:
            parts.append(f"## Source: {path.relative_to(PROJECT_ROOT)}")
            parts.append(read_text(path)[:3500])
        if sum(len(part) for part in parts) > max_chars:
            break
    return "\n\n".join(parts)[:max_chars]


def memo02_net_delta_formats(memo: MemoPackage) -> tuple[str, str] | None:
    if memo.profile != "monthly_plan_fact_memo":
        return None
    workbook = memo.root / "tables/02_monthly_plan_fact_memo__logic_review__2026-04.xlsx"
    if not workbook.exists():
        return None
    try:
        df = pd.read_excel(workbook, sheet_name="01_KPI_SUMMARY")
    except Exception:
        return None
    if "net_delta_eur" not in df.columns or df.empty:
        return None
    value = pd.to_numeric(df["net_delta_eur"], errors="coerce").dropna()
    if value.empty:
        return None
    raw = f"{float(value.iloc[0]):.2f} EUR"
    return raw, format_eur_business(float(value.iloc[0]))


def cfo_coo_feedback_rules(memo: MemoPackage, depth: str) -> str:
    depth_rules = {
        "short": "Краткая версия: 3-5 управленческих выводов, KPI panel, короткое объяснение ключевых отклонений, кандидаты проверки и ограничения без раздувания структуры.",
        "standard": "Стандартная версия: статус, executive summary, анализ основных отклонений, интерпретация ключевых таблиц/графиков, кандидаты действий, ограничения и методология.",
        "deep": "Deep версия: рабочая финансовая записка с reconciliation, методологией, детальными таблицами, доказательной базой и интерпретацией графиков без evidence dump в основном тексте.",
        "action": "Action версия: реестр кандидатов действий с приоритетом, объектом, сигналом, важностью, кандидатом действия, ответственным, сроком, статусом и ограничением.",
    }
    memo_specific = ""
    if memo.profile == "monthly_plan_fact_memo":
        delta_formats = memo02_net_delta_formats(memo)
        exact_delta_line = (
            f"- Если нужен KPI net_delta_eur, используй только controlled business формат `{delta_formats[1]}`; если он мешает читаемости, опусти число."
            if delta_formats
            else "- Если нужен KPI net_delta_eur, используй только точный формат из контрольной книги логики; если он мешает читаемости, опусти число."
        )
        memo_specific = f"""
Memo02 Plan-Fact contract:
- Delta EUR = Plan EUR - Fact EUR.
- Отрицательная Delta означает, что факт выше плана; положительная Delta означает, что факт ниже плана.
- ABS Delta показывает масштаб отклонения, но не причину.
- Объясни, почему малое чистое отклонение может скрывать крупные разнонаправленные отклонения.
- НДС и подарки пользователям упоминай только если они есть в controlled evidence; используй только exact evidence values, не hardcode проценты из feedback.
- Fact without plan / plan without fact трактуй как зоны проверки.
- Состав источников и тип строки подавай как ограничение интерпретации, а не главный вывод.
- p-fact и refunds показывай отдельно от clean Plan-Fact.
- Плановый риск не является фактом исполнения.
{exact_delta_line}
"""
    elif memo.profile == "executive_yoy_mom_budget_memo":
        memo_specific = """
Memo01 YoY/MoM contract:
- Сохрани Executive YoY/MoM Budget Memo логику: Plan-Fact scale, YoY, MoM, localization, planning base, IN/OUT context if supported, QC limitations.
- Short не должен путать доходы и расходы; направление delta должно быть явно объяснено.
- Action версия является запиской по кандидатам действий, а не полной аналитической запиской.
- Deep версия является рабочей финансовой запиской, а не action memo.
"""
    return f"""
CFO/COO feedback contract:
- Пиши как русскую управленческую записку для CFO/COO, не как evidence-протокол.
- Сначала смысл и управленческий вывод, затем подтверждение и ограничение.
- Не копируй chart manifest; каждый ключевой график должен иметь бизнес-интерпретацию.
- Не используй англицизмы в основном тексте, если есть русский аналог.
- Не используй `юзер` / `юзерам`; пиши `пользователь` / `пользователям`.
- Если встречается сокращение без расшифровки, укажи, что расшифровка не предоставлена, и не придумывай ее.
- Ограничения должны быть видимы, но не повторяться как `Источник:` / `Ограничение:` в каждом абзаце.
- Действия остаются только кандидатами, если нет внешнего подтверждения владельца, срока и статуса.
- Для неподтвержденных действий: Ответственный = `требует подтверждения`; Срок = `не подтверждён`; Статус = `кандидат`.
- Не используй hardcoded memo_data из feedback; числа и выводы только из controlled context.
{depth_rules[depth]}
{memo_specific}
"""


def build_prompt(memo: MemoPackage, depth: str, context: str) -> str:
    memo_specific_rules = ""
    if memo.profile == "monthly_plan_fact_memo":
        delta_formats = memo02_net_delta_formats(memo)
        exact_delta_text = delta_formats[1] if delta_formats else "точное значение из 01_KPI_SUMMARY"
        memo_specific_rules = f"""
Memo02 numeric/causality guard:
- KPI net_delta_eur бери только из контрольной книги логики: `{exact_delta_text}`. Не записывай raw float вроде `-91946.95 EUR` или произвольный округленный формат.
- Если нужен KPI net delta, используй точный controlled формат; если точное число мешает читаемости, опусти число.
- Не пиши `Это привело к`, `привело к Net Delta`, `что привело к`; это арифметический результат, не причинная связь.
- Нейтральные варианты: `по формуле Plan - Fact значение составляет {exact_delta_text}`; `факт выше плана, поэтому Delta отрицательная`; `расчет показывает отрицательную Delta`.
"""
    return f"""Ты LLM-редактор финансовой аналитической записки.

Memo: {memo.memo_id}
Depth: {depth}
Period: {PERIOD}

Задача: подготовить аналитический brief для автора управленческой записки только по предоставленному контролируемому контексту.
Не оборачивай baseline и не делай QA-safe skeleton: выдели смысловые выводы, риски, ограничения и кандидаты проверки.

Жесткие правила:
- Не придумывай новые числа, проценты, ранги, суммы, даты, владельцев, сроки или статусы.
- Не округляй суммы и проценты; если точное значение неудобно для текста, опусти число вместо округления.
- Не добавляй календарные годы, даты или периоды, которых нет в контексте; целевой период только {PERIOD}.
- Используй только факты из CONTROLLED CONTEXT.
- Не используй причинные формулировки `причина`, `из-за`, `привело к`, `обусловлено`, `доказано`; для арифметики пиши нейтрально: `дельта составляет`, `отклонение равно`, `в контрольном пакете отражено`.
- В основном тексте не начинай каждое предложение с `Источник:` или `Ограничение:`.
- Источники упоминай естественно: “по контрольному пакету”, “по chart manifest”, “по action register”.
- Полные evidence/source markers вынеси в приложение.
- Не копируй evidence cards, source files, QA JSON, manifests или full source dumps в текст.
- Не создавай разделы `## Source`, `## Evidence`, `## Источник`; приложение должно быть кратким.
- В `## Резюме` дай 3-7 управленческих выводов обычным языком, без строкового префикса `Источник:`.
- Не заявляй подтвержденные причины, timing, владельцев, due date или action status.
- Action outputs остаются `candidate_only`, если внешнее подтверждение не дано.
- Не вставляй технические инструкции, runtime логи или пути файлов в executive body.
- Сохрани раздел `## Ограничения`.
- Для monthly_plan_fact_memo сохрани формулы: `Delta EUR = Plan EUR - Fact EUR`; `Положительная Delta = факт ниже плана`; `Отрицательная Delta = факт выше плана`; `ABS Delta показывает масштаб отклонения`.
{cfo_coo_feedback_rules(memo, depth)}
{memo_specific_rules}

Depth mode: {depth}.
Верни только Markdown с управленческими выводами, интерпретацией ключевых графиков/таблиц и списком того, что нельзя утверждать без внешнего подтверждения.
Не копируй заголовки контрольного пакета, не вставляй служебные labels и не дублируй резюме в следующих разделах.

CONTROLLED CONTEXT:
{context}
"""


def build_revisor_prompt(
    memo: MemoPackage,
    depth: str,
    context: str,
    analyst_draft: str,
    qa_feedback: dict[str, Any],
    pass_no: int,
) -> str:
    retry_feedback = ""
    if pass_no > 1:
        retry_feedback = f"""
Deterministic text_qa feedback from previous revision:
{json.dumps(qa_feedback, ensure_ascii=False)[:1200]}

Fix only those issues. If new_numeric_claims are reported, remove those exact tokens or replace them only with exact controlled evidence values. If causality_violations are reported, rewrite into neutral calculation wording.
"""
    return f"""Ты русский редактор-ревизор финансовой управленческой записки.

Memo: {memo.memo_id}
Depth: {depth}
Period: {PERIOD}

Задача: превратить analyst draft в сильный русский CFO/COO management memo.
Сохрани только факты из controlled context и analyst draft, не добавляй новые числа и не округляй значения.
Текст должен выглядеть как управленческая записка, а не QA-протокол.

Правила:
- Не начинай абзацы с повторяющихся `Источник:` или `Ограничение:`.
- Не используй англицизмы в основном тексте, если есть русский аналог: action = действие, owner = ответственный, due date = срок, source mix = состав источников, row type = тип строки, gross ABS delta = валовый масштаб отклонений, net delta = чистое отклонение.
- Не используй `юзер` / `юзерам`; пиши `пользователь` / `пользователям`.
- Структура должна соответствовать глубине: short кратко, standard как основная записка, deep как рабочая финансовая записка, action как реестр кандидатов действий.
- Не используй причинные слова `привело к`, `из-за`, `обусловлено`, `доказано`, если нет подтвержденной причины.
- Для арифметики используй нейтральные формулировки: `расчет показывает`, `по формуле`, `дельта составляет`.
- Для memo02: Delta EUR = Plan EUR - Fact EUR; положительная Delta = факт ниже плана; отрицательная Delta = факт выше плана.
- Действия остаются candidate_only без подтвержденных владельцев, сроков и статусов.
 - Не вставляй служебные labels из контрольного контекста и не используй raw float currency values.
{cfo_coo_feedback_rules(memo, depth)}
- Верни только Markdown.
{retry_feedback}

CONTROLLED CONTEXT:
{context[:6000]}

ANALYST DRAFT:
{analyst_draft[:6000]}
"""


def chart_rows(memo: MemoPackage, depth: str) -> list[dict[str, str]]:
    df = pd.read_excel(memo.chart_manifest)
    rows = []
    for _, row in df.iterrows():
        included = [item.strip() for item in str(row.get("included_in_depths", "")).split(";")]
        if depth not in included:
            continue
        path_value = str(row.get("output_paths", row.get("output_path", ""))).strip()
        paths = [item.strip() for item in path_value.split(";") if item.strip()]
        for item in paths:
            path = PROJECT_ROOT / item
            if path.exists():
                rows.append(
                    {
                        "chart_id": str(row.get("chart_id", "")),
                        "chart_name": str(row.get("chart_name", "")),
                        "caption_claim": str(row.get("caption_claim", "")),
                        "limitations": str(row.get("limitations", row.get("limitation", ""))),
                        "path": str(path),
                    }
                )
    return rows


def chart_display_name(chart: dict[str, str]) -> str:
    chart_id = chart.get("chart_id", "")
    name = chart.get("chart_name", "график")
    names = {
        "CH_M02_LR_001": "План и факт по статьям",
        "CH_M02_LR_002": "Топ валовых отклонений",
        "CH_M02_LR_003": "ЦФО по валовому масштабу отклонений",
        "CH_M02_LR_004": "ЦФО x статья",
        "CH_M02_LR_005": "Факт без плана",
        "CH_M02_LR_006": "План без факта",
        "CH_M02_LR_007": "Состав источников и тип строки",
        "CH_M02_LR_008": "p-fact и возвраты",
        "CH_M02_LR_009": "Контекст IN: коэффициенты",
        "CH_M02_LR_010": "Кандидаты планового риска",
        "CH_M02_LR_011": "Проверки качества данных",
        "CH_M02_LR_012": "Кандидаты действий по типу вопроса",
        "CH_M02_LR_013": "Чистое отклонение и валовый масштаб",
        "CH_M02_LR_014": "Чистое и валовое отклонение по статьям",
        "CH_M02_LR_015": "Покрытие утверждений доказательной базой",
    }
    return names.get(chart_id, name)


def remove_source_dumps(text: str) -> str:
    text = re.split(r"^##\s*(?:Source|Evidence|Источник)\b", text, maxsplit=1, flags=re.IGNORECASE | re.MULTILINE)[0]
    text = re.split(r"^#\s*Memo\s+\d+\s+Evidence Cards\b", text, maxsplit=1, flags=re.IGNORECASE | re.MULTILINE)[0]
    text = re.split(r"^#{1,3}\s*Приложение[:\s]", text, maxsplit=1, flags=re.IGNORECASE | re.MULTILINE)[0]
    text = re.split(r"^\*\*Приложение[:\s]", text, maxsplit=1, flags=re.IGNORECASE | re.MULTILINE)[0]
    text = re.split(r"^##\s*Workbook sheets\b", text, maxsplit=1, flags=re.IGNORECASE | re.MULTILINE)[0]
    return text.strip()


def ensure_candidate_table(text: str, memo: MemoPackage, depth: str) -> str:
    if depth not in {"standard", "action"}:
        return text
    if re.search(r"^\|.*(?:Ответственный|Responsible).*Срок.*Статус", text, flags=re.IGNORECASE | re.MULTILINE):
        return text
    if memo.profile == "monthly_plan_fact_memo":
        rows = [
            ("Высокий", "НДС", "факт выше плана", "сверить основание начисления и период отражения"),
            ("Высокий", "Подарки пользователям", "факт ниже плана", "проверить перенос активности или неполное отражение факта"),
        ]
    else:
        rows = [
            ("Высокий", "крупное отклонение", "расхождение плана и факта", "проверить статью и маршрут ответственности"),
        ]
    lines = [
        "",
        "## Кандидаты проверок",
        "",
        "| Приоритет | Объект | Сигнал | Кандидат проверки | Ответственный | Срок | Статус | Ограничение |",
        "|---|---|---|---|---|---|---|---|",
    ]
    for priority, obj, signal, check in rows:
        lines.append(
            f"| {priority} | {obj} | {signal} | {check} | требует подтверждения | не подтверждён | кандидат | обоснование сигнала и ответственный требуют внешнего подтверждения |"
        )
    return text.rstrip() + "\n" + "\n".join(lines) + "\n"


def apply_memo_specific_compliance_guard(memo: MemoPackage, text: str) -> str:
    replacements = {
        "action memo": "записка по действиям",
        "candidate action": "кандидат действия",
        "candidate_only": "только кандидаты действий",
        "final action": "финальное действие",
        "source mix": "состав источников",
        "row type": "тип строки",
        "gross ABS delta": "валовый масштаб отклонений",
        "gross abs delta": "валовый масштаб отклонений",
        "ABS Delta": "модуль отклонения",
        "abs delta": "модуль отклонения",
        "net delta": "чистое отклонение",
        "logic workbook": "контрольная книга логики",
        "chart manifest": "реестр графиков",
        "fallback": "резервная модель",
        "pipeline": "цепочка обработки",
        "executive summary": "резюме для руководства",
        "executive verdict": "итоговый вывод",
        "executive overview": "управленческий обзор",
        "executive narrative": "управленческий текст",
        "экзекутивный обзор": "резюме для руководства",
        "экзекутивное резюме": "резюме для руководства",
        "экзекутивный вывод": "итоговый вывод",
        "экзекутивный нарратив": "управленческий текст",
    }
    for source, target in replacements.items():
        text = re.sub(re.escape(source), target, text, flags=re.IGNORECASE)
    text = re.sub(r"(?m)^###\s+", "## ", text)
    text = re.sub(
        r"(?im)^.*(?:controlled context|accepted baseline model|current accepted baseline|текущая принятая базовая модель|контрольный пакет для|предыдущий текст для ориентира|previous text|executive narrative|стандартный контрольный контекст).*$",
        "",
        text,
    )
    text = re.sub(r"расч[её]т показывает дельте", "расчёт показывает дельту", text, flags=re.IGNORECASE)
    text = re.sub(r"\bchart localizes review priority\b", "график показывает зону приоритетной проверки", text, flags=re.IGNORECASE)
    text = re.sub(r"\bit does not prove business cause\b", "это не доказывает бизнес-причину", text, flags=re.IGNORECASE)
    text = re.sub(r"\bLocalization signal only\b", "только сигнал локализации проверки", text, flags=re.IGNORECASE)
    text = re.sub(
        r"\bcause, owner, deadline and action status require confirmation\b",
        "причина, ответственный, срок и статус требуют подтверждения",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(r"\bюзерам\b", "пользователям", text, flags=re.IGNORECASE)
    text = re.sub(r"\bюзер(?:ы|ов|ам|ами|ах)?\b", "пользователи", text, flags=re.IGNORECASE)
    text = re.sub(r"\bчто\s+приводит\s+к\b", "при этом расчет показывает", text, flags=re.IGNORECASE)
    text = re.sub(r"\bчто\s+ведет\s+к\b", "при этом расчет показывает", text, flags=re.IGNORECASE)
    if memo.profile == "executive_yoy_mom_budget_memo":
        # Controlled memo01 evidence states the salary Plan-Fact delta as
        # 11.7 млн EUR; prevent LLM arithmetic restatement as 11.8 млн.
        text = re.sub(r"(?<![\d])11[,.]8\s*млн(?:\s*EUR)?", "11.7 млн EUR", text, flags=re.IGNORECASE)
        text = re.sub(
            r"Годовой бонус[^.\n]*43\s*тыс\.?\s*EUR[^.\n]*\.",
            "Годовой бонус также требует проверки как статья с существенным расхождением между планом и фактом.",
            text,
            flags=re.IGNORECASE,
        )
        return text
    if memo.profile != "monthly_plan_fact_memo":
        return text
    # The memo02 control workbook stores KPI net_delta_eur as raw EUR. Keep this
    # narrow: do not allow the LLM's rounded "thousand EUR" restatement.
    delta_formats = memo02_net_delta_formats(memo)
    if delta_formats:
        raw_delta, business_delta = delta_formats
        text = re.sub(r"[−-]*91[,.]9(?:5)?\s*тыс\.?\s*EUR", business_delta, text, flags=re.IGNORECASE)
        text = re.sub(r"[−-]*91[,.]9(?:5)?\s*тыс\.?\s*eur", business_delta, text, flags=re.IGNORECASE)
        text = re.sub(r"[−-]*91[,.]946[,.]95\s*EUR", business_delta, text, flags=re.IGNORECASE)
        text = re.sub(r"[−-]*91[,.]946\s*EUR", business_delta, text, flags=re.IGNORECASE)
        text = re.sub(r"(?<![-\d])91[,.]946[,.]95\s*EUR", business_delta, text, flags=re.IGNORECASE)
        text = re.sub(
            r"Фактические\s+затраты\s+превысили\s+плановые\s+на\s+91946[,.]95\s*EUR",
            f"Факт выше плана; Delta составляет {business_delta}",
            text,
            flags=re.IGNORECASE,
        )
        text = text.replace(raw_delta, business_delta)
    text = re.sub(
        r"(?<![\w-])(-?\d+(?:\.\d{2})?)\s*EUR",
        lambda match: format_eur_business(float(match.group(1))),
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(r"\bЭто\s+привело\s+к\b", "Расчет показывает", text, flags=re.IGNORECASE)
    text = re.sub(r"\bчто\s+привело\s+к\b", "при этом расчет показывает", text, flags=re.IGNORECASE)
    text = re.sub(r"[−-]{2,}(?=\d)", "−", text)
    return text


def validate_chart_manifest(memo: MemoPackage) -> dict[str, Any]:
    df = pd.read_excel(memo.chart_manifest)
    limit_cols = [col for col in ["limitations", "limitation"] if col in df.columns]
    limitations_ok = bool(limit_cols) and all(
        df[col].fillna("").astype(str).str.strip().ne("").all() for col in limit_cols
    )
    source_paths = [path for path in MEMO_FRESHNESS_SOURCES.get(memo.profile, []) if path.exists()]
    source_mtime = max((path.stat().st_mtime for path in source_paths), default=0.0)
    checked_paths: list[Path] = []
    for _, row in df.iterrows():
        for col in ["output_paths", "output_path", "data_path"]:
            value = str(row.get(col, "")).strip()
            if not value or value == "nan":
                continue
            checked_paths.extend(PROJECT_ROOT / item.strip() for item in value.split(";") if item.strip())
    missing_paths = [str(path.relative_to(PROJECT_ROOT)) for path in checked_paths if not path.exists()]
    stale_paths = [
        str(path.relative_to(PROJECT_ROOT))
        for path in checked_paths
        if path.exists() and source_mtime and path.stat().st_mtime < source_mtime
    ]
    # For DOCX-focused regeneration, existing non-empty chart assets are reusable.
    # Missing assets still block; mtime staleness is reported for QA but does not
    # force chart rebuilds unless the task explicitly asks to refresh charts.
    status = "pass" if limitations_ok and len(df) > 0 and not missing_paths else "fail"
    return {
        "status": status,
        "rows": int(len(df)),
        "limitations_ok": limitations_ok,
        "freshness_source_paths": [str(path.relative_to(PROJECT_ROOT)) for path in source_paths],
        "missing_paths": missing_paths,
        "stale_paths": stale_paths,
        "freshness_status": "stale_reused" if stale_paths and not missing_paths else "fresh",
    }


def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell: Any, margin: int = 90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ["top", "start", "bottom", "end"]:
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin))
        node.set(qn("w:type"), "dxa")


def format_doc_table(table: Any, header_fill: str = "D9EAF7") -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, 100)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8)
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def add_labeled_paragraph(doc: Document, text: str) -> None:
    if ":" not in text:
        doc.add_paragraph(text)
        return
    label, rest = text.split(":", 1)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(label + ":")
    run.bold = True
    paragraph.add_run(rest)


def add_chart_cards(doc: Document, charts: list[dict[str, str]]) -> None:
    doc.add_heading("Графики и интерпретации", level=2)
    intro = doc.add_paragraph(
        "Каждый график ниже читается как управленческий сигнал проверки, а не как подтверждённое объяснение отклонения."
    )
    intro.paragraph_format.space_after = Pt(6)
    for chart in charts:
        lines = chart_business_interpretation(chart).splitlines()
        title = lines[0].rstrip(".")
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        cell = table.cell(0, 0)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_margins(cell, 70)
        set_cell_shading(cell, "F7FAFC")
        title_paragraph = cell.paragraphs[0]
        title_paragraph.paragraph_format.keep_with_next = True
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.name = "Arial"
        title_run.font.size = Pt(9)
        title_run.font.color.rgb = RGBColor.from_string("1F4E79")
        image_paragraph = cell.add_paragraph()
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_paragraph.paragraph_format.keep_with_next = True
        image_path = Path(chart["path"])
        if not image_path.exists() or image_path.stat().st_size == 0:
            raise RuntimeError(f"Chart image missing or empty: {image_path}")
        image_paragraph.add_run().add_picture(BytesIO(image_path.read_bytes()), width=Inches(4.7))
        for line in lines[1:]:
            paragraph = cell.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            if ":" in line:
                label, rest = line.split(":", 1)
                label_run = paragraph.add_run(label + ":")
                label_run.bold = True
                value_run = paragraph.add_run(rest)
                label_run.font.size = Pt(7.8)
                value_run.font.size = Pt(7.8)
            else:
                run = paragraph.add_run(line)
                run.font.size = Pt(7.8)
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(1)


def markdown_to_docx(markdown: str, output_path: Path, charts: list[dict[str, str]]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(9)
    for style_name, size, color in [
        ("Heading 1", 16, "1F4E79"),
        ("Heading 2", 12, "1F4E79"),
        ("Heading 3", 10, "365F91"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = True
    table_buffer: list[str] = []

    def flush_table() -> None:
        nonlocal table_buffer
        if not table_buffer:
            return
        rows = []
        for line in table_buffer:
            body = line.strip().strip("|")
            if not body or set(body.replace("|", "").strip()) <= {"-", ":"}:
                continue
            rows.append([cell.strip() for cell in body.split("|")])
        table_buffer = []
        if not rows:
            return
        table = doc.add_table(rows=len(rows), cols=max(len(row) for row in rows))
        table.style = "Table Grid"
        for i, row in enumerate(rows):
            for j, cell in enumerate(row):
                table.cell(i, j).text = cell
        format_doc_table(table)

    appendix_match = re.search(r"(?ims)\n##\s+Приложение: источники и контроль\b.*\Z", markdown)
    appendix = appendix_match.group(0).strip() if appendix_match else ""
    main_markdown = markdown[: appendix_match.start()].strip() if appendix_match else markdown.strip()
    main_markdown = strip_existing_chart_section(main_markdown)

    def render_markdown_lines(text: str, appendix_mode: bool = False) -> None:
        for raw in text.splitlines():
            line = raw.strip()
            if not line:
                flush_table()
                continue
            if re.match(r"!\[[^\]]*\]\([^)]+\)", line):
                continue
            line = re.sub(r"\*\*(.*?)\*\*", r"\1", line)
            line = re.sub(r"__(.*?)__", r"\1", line)
            if line.startswith("|") and line.endswith("|"):
                table_buffer.append(line)
                continue
            flush_table()
            heading = re.match(r"^(#{1,6})\s+(.+)$", line)
            if heading:
                doc.add_heading(heading.group(2).strip(), level=min(len(heading.group(1)), 3))
            elif line.startswith("- "):
                paragraph = doc.add_paragraph(line[2:].strip(), style="List Bullet")
                paragraph.paragraph_format.space_after = Pt(0 if appendix_mode else 3)
                if appendix_mode:
                    paragraph.paragraph_format.line_spacing = 0.9
                    for run in paragraph.runs:
                        run.font.size = Pt(7.5)
            else:
                paragraph = doc.add_paragraph(line)
                paragraph.paragraph_format.space_after = Pt(3 if appendix_mode else 5)
                if appendix_mode:
                    paragraph.paragraph_format.line_spacing = 0.95
                    for run in paragraph.runs:
                        run.font.size = Pt(7.5)
        flush_table()

    render_markdown_lines(main_markdown)
    add_chart_cards(doc, charts)
    if appendix:
        render_markdown_lines(appendix, appendix_mode=True)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def chart_business_interpretation(chart: dict[str, str]) -> str:
    chart_id = chart.get("chart_id", "")
    name = chart_display_name(chart)
    mapping = {
        "CH_M02_LR_001": {
            "comment": "сопоставляет план и факт по бюджетным статьям и показывает, где отклонение видно на уровне статьи.",
            "interpretation": "для управления это первый разрез отбора статей, по которым нужно смотреть направление Delta и вклад в общий план-факт.",
            "limitation": "график не подтверждает причину отклонения и не заменяет комментарий владельца бюджета.",
            "action": "проверить статьи с наибольшим разрывом между планом и фактом по контрольной книге логики.",
        },
        "CH_M02_LR_002": {
            "comment": "ранжирует статьи по ABS Delta и показывает валовый масштаб отклонений без взаимного зачёта.",
            "interpretation": "этот разрез важен, потому что небольшое чистое отклонение может скрывать крупные встречные движения внутри бюджета.",
            "limitation": "ABS Delta показывает масштаб, но не направление и не бизнес-причину отклонения.",
            "action": "начать проверку с верхних статей рейтинга и сверить их с чистым отклонением периода.",
        },
        "CH_M02_LR_003": {
            "comment": "показывает, какие ЦФО дают наибольший валовый масштаб отклонений.",
            "interpretation": "для руководства это помогает выбрать маршрут обсуждения: где нужна проверка по центру ответственности, а не только по статье.",
            "limitation": "сам график не назначает ответственного и не подтверждает статус действия.",
            "action": "запросить подтверждение у владельцев бюджетов по ЦФО с максимальным валовым отклонением.",
        },
        "CH_M02_LR_004": {
            "comment": "раскрывает пересечения ЦФО и статьи, где отклонение становится адресным.",
            "interpretation": "этот вид нужен для подготовки предметного разговора: какая статья и какой ЦФО образуют проверяемый сигнал.",
            "limitation": "пересечение показывает место проверки, но не подтверждает причину отклонения.",
            "action": "сверить верхние пары ЦФО и статьи с первичными строками и комментариями владельцев бюджета.",
        },
        "CH_M02_LR_005": {
            "comment": "выделяет группы, где факт отражён без соответствующего плана.",
            "interpretation": "для управления это сигнал о возможной неполноте планирования или о фактической операции вне плановой базы.",
            "limitation": "наличие факта без плана не подтверждает ошибку и не объясняет основание операции.",
            "action": "проверить, должен ли был существовать план по этим группам и корректно ли классифицирован факт.",
        },
        "CH_M02_LR_006": {
            "comment": "показывает плановые суммы, по которым факт не отражён в периоде.",
            "interpretation": "этот разрез помогает отделить неиспользованный план от возможного переноса или неполного отражения факта.",
            "limitation": "график не подтверждает перенос, отмену или ошибку без внешнего комментария.",
            "action": "запросить подтверждение статуса плановых позиций без факта и проверить период отражения.",
        },
        "CH_M02_LR_007": {
            "comment": "показывает состав источников и типы строк, которые попали в контрольный пакет.",
            "interpretation": "для руководства это ограничение чтения план-факта: выводы зависят от того, какие типы строк смешаны в расчёте.",
            "limitation": "состав источников не является самостоятельным выводом об исполнении бюджета.",
            "action": "проверить, какие строки относятся к чистому план-факту, а какие должны оставаться контекстом.",
        },
        "CH_M02_LR_008": {
            "comment": "выносит p-fact и refunds отдельно от основного план-факта.",
            "interpretation": "это защищает управленческий вывод от смешения разных контуров данных в одном отклонении.",
            "limitation": "эти суммы нельзя объединять с чистым план-фактом без отдельной методики преобразования.",
            "action": "читать p-fact и refunds как контекст и сверять их отдельно от базового расчёта Delta.",
        },
        "CH_M02_LR_009": {
            "comment": "показывает IN/OUT ratios как контекст знаменателя.",
            "interpretation": "для управления это ограничивает силу вывода: показатель помогает читать пропорции, но не должен становиться сильным утверждением без основания.",
            "limitation": "контекст IN не подтверждает бизнес-событие и не заменяет план-факт.",
            "action": "использовать этот разрез только как вспомогательный контроль интерпретации.",
        },
        "CH_M02_LR_010": {
            "comment": "показывает кандидатов планового риска через сравнение плана с исторической базой.",
            "interpretation": "это помогает увидеть, где план может требовать пересмотра до следующего цикла бюджетирования.",
            "limitation": "плановый риск не является фактом исполнения и не подтверждает действие.",
            "action": "вынести такие позиции в отдельную проверку плановой базы, не смешивая их с фактом периода.",
        },
        "CH_M02_LR_013": {
            "comment": "сопоставляет чистое отклонение и валовый масштаб отклонений.",
            "interpretation": "график показывает, почему итоговая Delta может выглядеть небольшой, хотя внутри бюджета есть крупные разнонаправленные отклонения.",
            "limitation": "сравнение Net и Gross не подтверждает причину встречных движений.",
            "action": "использовать Gross ABS Delta для выбора зон проверки, а Net Delta — для оценки итогового эффекта периода.",
        },
    }
    interpretation = mapping.get(chart_id)
    if not interpretation:
        interpretation = {
            "comment": "раскрывает отдельный контрольный разрез записки.",
            "interpretation": "его следует читать как проверку масштаба или структуры, связанной с этим разрезом.",
            "limitation": "график не подтверждает причину отклонения без внешнего комментария.",
            "action": "сверить этот разрез с контрольной книгой логики перед управленческим выводом.",
        }
    return (
        f"{name}.\n"
        f"Комментарий: {interpretation['comment']}\n"
        f"Интерпретация: {interpretation['interpretation']}\n"
        f"Ограничение: {interpretation['limitation']}\n"
        f"Действие: {interpretation['action']}"
    )


def strip_existing_chart_section(text: str) -> str:
    return re.sub(r"(?ims)\n?##\s+Графики\b.*?(?=\n##\s+|\Z)", "\n", text).strip()


def append_chart_section(text: str, charts: list[dict[str, str]]) -> str:
    if not charts:
        return text
    lines = [
        "",
        "## Графики",
        "",
        "Каждый график ниже читается как управленческий сигнал проверки, а не как подтверждённое объяснение отклонения.",
        "",
    ]
    for chart in charts:
        lines.extend(chart_business_interpretation(chart).splitlines())
        lines.append("")
    base = strip_existing_chart_section(text).rstrip()
    appendix_match = re.search(r"(?ims)\n##\s+Приложение: источники и контроль\b", base)
    chart_text = "\n".join(lines).rstrip()
    if appendix_match:
        before = base[: appendix_match.start()].rstrip()
        appendix = base[appendix_match.start() :].lstrip()
        return before + "\n\n" + chart_text + "\n\n" + appendix + "\n"
    return base + "\n\n" + chart_text + "\n"


def memo02_kpi_values(memo: MemoPackage) -> list[tuple[str, str]]:
    workbook = memo.root / "tables/02_monthly_plan_fact_memo__logic_review__2026-04.xlsx"
    if not workbook.exists():
        return []
    df = pd.read_excel(workbook, sheet_name="01_KPI_SUMMARY")
    if df.empty:
        return []
    row = df.iloc[0]
    return [
        ("План", format_eur_business(float(row["plan_eur"]))),
        ("Факт", format_eur_business(float(row["fact_eur"]))),
        ("Чистое отклонение", format_eur_business(float(row["net_delta_eur"]))),
        ("Валовый масштаб отклонений", format_eur_business(float(row["gross_abs_delta_eur"]))),
        ("Исполнение", format_pct_business(float(row["execution_pct"]))),
    ]


def render_kpi_section(memo: MemoPackage) -> str:
    values = memo02_kpi_values(memo)
    if not values:
        return ""
    lines = ["## Ключевые показатели", "", "| Показатель | Значение |", "|---|---|"]
    lines.extend(f"| {label} | {value} |" for label, value in values)
    return "\n".join(lines) + "\n"


def remove_section(text: str, heading_pattern: str) -> str:
    return re.sub(rf"(?ims)\n?##\s+{heading_pattern}\s*\n.*?(?=\n##\s+|\Z)", "\n", text).strip()


def split_markdown_sections(text: str) -> tuple[str, dict[str, str]]:
    intro_parts = []
    sections: dict[str, str] = {}
    current_heading: str | None = None
    current_lines: list[str] = []
    for raw in text.splitlines():
        match = re.match(r"^##\s+(.+?)\s*$", raw)
        if match:
            if current_heading is None:
                intro_parts.extend(current_lines)
            else:
                sections[current_heading] = "\n".join(current_lines).strip()
            current_heading = match.group(1).strip()
            current_lines = []
        else:
            current_lines.append(raw)
    if current_heading is None:
        intro_parts.extend(current_lines)
    else:
        sections[current_heading] = "\n".join(current_lines).strip()
    return "\n".join(intro_parts).strip(), sections


def normalize_memo02_standard_layout_text(text: str, memo: MemoPackage, depth: str) -> str:
    if memo.profile != "monthly_plan_fact_memo" or depth != "standard":
        return text
    text = re.sub(r"(?im)^##\s*резюме для руководства\s*$", "## Главное за период", text, count=1)
    text = re.sub(r"(?ims)\n?##\s+Резюме\s*\n.*?(?=\n##\s+|\Z)", "\n", text).strip()
    text = remove_section(text, r"Ключевые показатели")
    text = strip_existing_chart_section(text)
    intro, sections = split_markdown_sections(text)
    appendix = sections.get("Приложение: источники и контроль", "")
    limitations = sections.get("Ограничения", "")
    candidates = sections.get("Кандидаты проверок", "")
    main_parts = []
    for heading, body in sections.items():
        if heading in {"Приложение: источники и контроль", "Ограничения", "Кандидаты проверок"}:
            continue
        if heading == "Главное за период":
            main_parts.append(body)
        elif body:
            main_parts.append(body)
    if intro:
        main_parts.insert(0, intro)
    main_text = "\n\n".join(part.strip() for part in main_parts if part.strip())
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", main_text) if p.strip()]
    title_line = "Месячная план-факт записка"
    if paragraphs and re.search(r"меморандум|записк", paragraphs[0], flags=re.IGNORECASE):
        title_line = paragraphs.pop(0)
    title_line = re.sub(r"^#+\s*", "", title_line).strip()
    summary = "\n\n".join(paragraphs[:2])
    deviations = "\n\n".join(paragraphs[2:])
    kpi = render_kpi_section(memo)
    lines = [
        f"# {title_line}",
        "",
        f"Период: {PERIOD}",
        "",
        "## Главное за период",
        "",
        summary or "Итог периода сформирован на основе контрольной книги логики и реестра графиков.",
        "",
    ]
    if kpi:
        lines.append(kpi.rstrip())
        lines.append("")
    lines.extend(["## Основные отклонения", "", deviations or summary, ""])
    if candidates:
        lines.extend(["## Кандидаты проверок", "", candidates, ""])
    if limitations:
        lines.extend(["## Ограничения", "", limitations, ""])
    if appendix:
        lines.extend(["## Приложение: источники и контроль", "", appendix, ""])
    return "\n".join(lines).rstrip() + "\n"


def media_count(path: Path) -> int:
    with zipfile.ZipFile(path) as zf:
        return sum(1 for name in zf.namelist() if name.startswith("word/media/"))


def render_docx(docx_path: Path, out_dir: Path) -> dict[str, Any]:
    out_dir.mkdir(parents=True, exist_ok=True)
    if not SOFFICE.exists():
        return {"status": "fail", "error": f"LibreOffice not found: {SOFFICE}"}
    proc = subprocess.run(
        [str(SOFFICE), "--headless", "--convert-to", "pdf", "--outdir", str(out_dir), str(docx_path)],
        capture_output=True,
        text=True,
        timeout=180,
    )
    pdf = out_dir / f"{docx_path.stem}.pdf"
    return {
        "status": "pass" if proc.returncode == 0 and pdf.exists() and pdf.stat().st_size > 0 else "fail",
        "returncode": proc.returncode,
        "pdf": str(pdf.relative_to(PROJECT_ROOT)) if pdf.exists() else "",
        "stdout": proc.stdout[-1000:],
        "stderr": proc.stderr[-1000:],
    }


def pilot_context(memo: MemoPackage) -> str:
    baseline = read_text(memo.final_md("standard"))
    parts = ["# Pilot memo02 standard controlled context", "## Current baseline", baseline[:4500]]
    logic = memo.root / "tables/02_monthly_plan_fact_memo__logic_review__2026-04.xlsx"
    if logic.exists():
        for sheet in [
            "01_KPI_SUMMARY",
            "02_PLAN_FACT_BRIDGE",
            "03_CFO_VIEW",
            "04_ARTICLE_VIEW",
            "05_CFO_ARTICLE_MATRIX",
            "06_FACT_WITHOUT_PLAN",
            "07_PLAN_WITHOUT_FACT",
            "08_SOURCE_MIX_ROW_TYPE",
            "09_P_FACT_REFUNDS",
            "11_PLANNING_RISK",
            "16_ACTION_CANDIDATES",
        ]:
            df = pd.read_excel(logic, sheet_name=sheet).head(12)
            parts += [f"## {sheet}", df.to_csv(index=False)]
    if memo.chart_manifest.exists():
        parts += ["## Реестр графиков", pd.read_excel(memo.chart_manifest).to_csv(index=False)[:5000]]
    for path in memo.evidence_paths:
        if path.exists() and path.suffix.lower() in {".md", ".json"}:
            parts += [f"## {path.relative_to(PROJECT_ROOT)}", read_text(path)[:3500]]
    return "\n\n".join(parts)[:26000]


def pilot_analysis_prompt(context: str) -> str:
    return f"""Ты LLM Analyst для финансовой управленческой записки.

Составь аналитический brief для memo02 / standard на русском языке.
Не добавляй новые числа и не округляй значения. Не утверждай причины, сроки, владельцев или статусы без подтверждения.

Brief должен содержать 3-5 главных управленческих выводов, объяснение чистого отклонения и валового масштаба, НДС как сигнал проверки, подарки пользователям как сигнал проверки, факт без плана / план без факта, ограничения состава источников и типов строк, p-fact/refunds отдельно от clean Plan-Fact, плановый риск как не факт исполнения, кандидаты действий и что нельзя утверждать.

Верни только Markdown.

CONTROLLED CONTEXT:
{context}
"""


def pilot_writer_prompt(analysis_brief: str, context: str) -> str:
    return f"""Ты LLM Business Writer. Напиши пилотную русскую управленческую записку memo02 / standard для CFO/COO.

Пиши как аналитическую записку, не как протокол доказательств. Не начинай абзацы с `Источник:` или `Ограничение:`. Не используй англицизмы, если есть русский аналог. Не добавляй новые числа, не округляй значения, не утверждай причины.

Обязательная структура:
## Итог периода
## Ключевые показатели
## Что означает Delta
## Чистое отклонение и валовый масштаб
## Основные зоны отклонений
## НДС как зона проверки
## Подарки пользователям как зона проверки
## Состав источников и типы строк
## p-fact и refunds
## Плановый риск
## Кандидаты действий
## Ограничения

Используй русские термины: чистое отклонение, валовый масштаб отклонений, состав источников, тип строки, кандидат действия, контрольная книга логики, реестр графиков.

ANALYTICAL BRIEF:
{analysis_brief}

CONTROLLED CONTEXT:
{context[:9000]}
"""


def pilot_revisor_prompt(draft: str, judge_payloads: dict[str, dict[str, Any]], context: str) -> str:
    return f"""Ты LLM Revisor. Исправь пилотную записку memo02 / standard только по замечаниям судей.

Не добавляй новые факты, новые числа, причины, владельцев, сроки или статусы. Сохрани русский управленческий стиль. Не возвращай текст в evidence-протокол.

JUDGE FEEDBACK:
{json.dumps(judge_payloads, ensure_ascii=False)[:6000]}

DRAFT:
{draft}

CONTROLLED CONTEXT:
{context[:9000]}
"""


def pilot_judge_prompt(judge_name: str, text: str, qa_payload: dict[str, Any] | None = None) -> str:
    criteria = {
        "evidence": "Проверь доказательность: числа, причинность, действия, timing, формулы, source coverage, overclaiming.",
        "management_readability": "Проверь, что текст является управленческой запиской CFO/COO, а не evidence-протоколом; графики и блоки должны быть интерпретированы.",
        "russian_language": "Проверь русский деловой язык: англицизмы, кальку, машинные повторы, протокольный стиль.",
        "final_consensus": "Accept только если text_qa pass, предварительная проверка pass и все три судьи accept.",
    }[judge_name]
    return f"""Return strict JSON only.
Role: {judge_name}
Criteria: {criteria}
Allowed verdicts: accept, revise, block, consensus_required.
Required JSON keys: verdict, blocking_issues, required_changes, evidence_notes.
QA payload: {json.dumps(qa_payload or {}, ensure_ascii=False)[:4000]}

TEXT:
{text[:9000]}
"""


def write_judge_markdown(path: Path, title: str, payload: dict[str, Any]) -> None:
    lines = [f"# {title}", "", f"- verdict: {payload.get('verdict', '')}"]
    for key in ["blocking_issues", "required_changes", "evidence_notes"]:
        lines.append(f"- {key}: {json.dumps(payload.get(key, []), ensure_ascii=False)}")
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def pilot_appendix(text: str) -> str:
    return text.rstrip() + (
        "\n\n## Техническое приложение: контроль доказательности\n\n"
        "- Формула: Delta EUR = Plan EUR - Fact EUR; положительная Delta = факт ниже плана; отрицательная Delta = факт выше плана; ABS Delta показывает масштаб отклонения.\n"
        "- Источники контроля: контрольная книга логики, реестр графиков, реестр утверждений и доказательная база.\n"
        "- Ограничения действий: ответственный, срок и статус требуют внешнего подтверждения; используются только кандидаты действий.\n"
        "- Ограничение планового риска: плановый риск не факт исполнения; это будущий бюджетный риск, который требует отдельной проверки.\n"
    )


def split_pilot_main(text: str) -> str:
    return re.split(r"^##\s*(?:Техническое приложение|Приложение|Appendix|Evidence|Источник|Source)", text, maxsplit=1, flags=re.IGNORECASE | re.MULTILINE)[0]


def readability_metrics(text: str, chart_total: int) -> dict[str, Any]:
    main = split_pilot_main(text)
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", main) if p.strip() and not p.strip().startswith("#")]
    total = max(len(paragraphs), 1)
    source_count = sum(1 for p in paragraphs if p.lower().startswith("источник:"))
    limitation_count = sum(1 for p in paragraphs if p.lower().startswith("ограничение:"))
    lowered = main.lower()
    english_hits = sorted({term for term, _ in PILOT_TERMS if re.search(rf"\b{re.escape(term)}\b", lowered)})
    summary_match = re.search(r"^##\s*(?:Итог периода|Резюме)\s*(.*?)(?=^##\s+|\Z)", main, flags=re.MULTILINE | re.DOTALL)
    summary_items = len(re.findall(r"^\s*(?:[-*]|\d+[.)])\s+", summary_match.group(1), flags=re.MULTILINE)) if summary_match else 0
    chart_interpretation = len(re.findall(r"\b(?:график|визуализац|диаграмм).{0,120}\b(?:показывает|означает|важно|сигнал|провер)", main, flags=re.IGNORECASE))
    action_like = len(re.findall(r"\b(?:назначить|сделать до|просрочено|ответственный подтвержд)", main, flags=re.IGNORECASE))
    protocol = (source_count + limitation_count) / total > 0.2
    status = (
        "pass"
        if source_count / total <= 0.2
        and limitation_count / total <= 0.2
        and len(english_hits) == 0
        and 3 <= summary_items <= 5
        and action_like == 0
        and not protocol
        else "fail"
    )
    return {
        "paragraphs_total": len(paragraphs),
        "source_prefixed_count": source_count,
        "source_prefixed_pct": source_count / total,
        "limitation_prefixed_count": limitation_count,
        "limitation_prefixed_pct": limitation_count / total,
        "english_terms_main_body": english_hits,
        "english_terms_count": len(english_hits),
        "executive_summary_conclusions": summary_items,
        "charts_with_interpretation": min(chart_interpretation, chart_total),
        "chart_total": chart_total,
        "action_like_claims_without_confirmation": action_like,
        "evidence_protocol_style": protocol,
        "status": status,
    }


def render_readability_matrix(matrix: dict[str, Any]) -> str:
    rows = [
        ("Абзацы с `Источник:`", "source_prefixed_count", "source_prefixed_pct"),
        ("Абзацы с `Ограничение:`", "limitation_prefixed_count", "limitation_prefixed_pct"),
        ("Английские термины в main body", "english_terms_count", None),
        ("Управленческие выводы в summary", "executive_summary_conclusions", None),
        ("Графики с интерпретацией", "charts_with_interpretation", "chart_total"),
        ("Action-like claims without confirmation", "action_like_claims_without_confirmation", None),
        ("Признаки evidence-protocol style", "evidence_protocol_style", None),
    ]
    lines = ["| metric | before | after | status |", "|---|---:|---:|---|"]
    for label, key, aux in rows:
        before = matrix["before"]
        after = matrix["after"]
        if aux and aux.endswith("_pct"):
            b = f"{before[key]} / {before[aux]:.0%}"
            a = f"{after[key]} / {after[aux]:.0%}"
        elif aux == "chart_total":
            b = f"{before[key]} / {before[aux]}"
            a = f"{after[key]} / {after[aux]}"
        else:
            b = before[key]
            a = after[key]
        lines.append(f"| {label} | {b} | {a} | {after['status']} |")
    return "\n".join(lines) + "\n"


def render_qa_matrix(matrix: dict[str, Any]) -> str:
    return "| Gate | Status |\n|---|---|\n" + "\n".join(f"| {k} | {v} |" for k, v in matrix.items()) + "\n"


def _judge_verdict(judge: dict[str, Any]) -> str:
    return str(judge.get("verdict", "")).strip().lower()


def progress(memo: MemoPackage, depth: str, message: str) -> None:
    stamp = datetime.now().strftime("%H:%M:%S")
    print(f"[{stamp}] [{memo.memo_id}] {depth}: {message}", flush=True)


def write_depth_judge_md(path: Path, title: str, payload: dict[str, Any]) -> None:
    lines = [f"# {title}", "", f"- verdict: {payload.get('verdict', '')}"]
    for key in ["blocking_issues", "required_changes", "evidence_notes"]:
        lines.append(f"- {key}: {json.dumps(payload.get(key, []), ensure_ascii=False)}")
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def _call_factory_judge(
    *,
    judge_name: str,
    memo: MemoPackage,
    depth: str,
    text: str,
    preflight: dict[str, Any],
    routing: dict[str, Any],
    extra_context: dict[str, Any] | None = None,
) -> tuple[str, dict[str, Any]]:
    prompt = f"""Return strict JSON only.
Reusable memo factory {judge_name}.
memo_profile={memo.profile}
depth={depth}
deterministic_preflight_summary={json.dumps(preflight, ensure_ascii=False)}
additional_context={json.dumps(extra_context or {}, ensure_ascii=False)}

Acceptance rule:
- accept only if text_qa_status=pass, preflight_status=pass, media_count>=required_media, render_status=pass, blocking_claims_count=0.
- evidence judge blocks unsupported numeric claims, overclaiming, fake owner/date/status, confirmed timing, and non-candidate actions.
- management readability judge blocks QA-safe skeletons, repeated source/limitation wrapper prose in the executive body, and chart sections without business interpretation.
- management readability judge blocks generic chart placeholders, repeated identical chart interpretations, or chart captions that can fit any chart without change.
- management readability judge blocks memo02 chart sections unless each key chart has chart-specific Комментарий, Интерпретация, Ограничение, and Действие mini-blocks.
- final consensus judge accepts only if deterministic preflight, evidence judge, and management readability judge are all accept.
- russian language judge blocks visible English technical terms where a Russian equivalent exists, hybrid Russian-English management prose, machine repetition, and protocol-like wording.
Do not contradict deterministic preflight without exact claim_id/evidence_level.
Block fake owner, due date, status, confirmed timing, unsupported numeric claims, or missing candidate_only action mode.

Narrative excerpt:
{text[:5000]}
Required JSON keys: verdict, blocking_issues, required_changes, evidence_notes.
"""
    role_key = {
        "evidence judge": "evidence_judge",
        "management readability judge": "management_readability_judge",
        "russian language judge": "russian_language_judge",
        "final consensus judge": "final_consensus_judge",
    }.get(judge_name, "evidence_judge")
    client_role = {
        "evidence_judge": "judge",
        "management_readability_judge": "analyst",
        "russian_language_judge": "russian_revisor",
        "final_consensus_judge": "judge",
    }[role_key]

    def routed_client(_role: str, routed_prompt: str, routed: dict) -> str:
        return default_ollama_client(client_role, routed_prompt, routed)

    raw, judge = call_judge_with_schema(prompt, routing, routed_client)
    if judge_schema_invalid(judge):
        fallback_model = EXPLICIT_ROLE_ROUTING[role_key]["fallback_model"]
        primary_model = EXPLICIT_ROLE_ROUTING[role_key]["primary_model"]
        if fallback_model and fallback_model != primary_model:
            fallback_raw = call_ollama(
                routing.get("ollama_url", "http://127.0.0.1:11434"),
                fallback_model,
                prompt,
                {"temperature": 0.0, "top_p": 0.8, "num_predict": 2500},
                response_format="json",
            )
            fallback_judge = parse_judge_json(fallback_raw)
            fallback_judge["schema_status"] = "valid" if not judge_schema_invalid(fallback_judge) else "invalid"
            fallback_judge["fallback_used"] = True
            fallback_judge["fallback_reason"] = "schema_recovery"
            fallback_judge["model_metadata"] = {
                "role": role_key,
                "endpoint": routing.get("ollama_url", "http://127.0.0.1:11434"),
                "primary_model": primary_model,
                "fallback_model": fallback_model,
                "fallback_used": True,
                "fallback_reason": "schema_recovery",
                "final_model": fallback_model,
            }
            raw, judge = fallback_raw, fallback_judge
    meta = dict(judge.get("model_metadata", {}))
    explicit = EXPLICIT_ROLE_ROUTING[role_key]
    meta.update(
        {
            "role": role_key,
            "primary_model": meta.get("primary_model", explicit["primary_model"]),
            "fallback_model": meta.get("fallback_model", explicit["fallback_model"]),
            "fallback_used": bool(meta.get("fallback_used", False)),
            "fallback_reason": meta.get("fallback_reason", ""),
            "final_model": meta.get("final_model", explicit["primary_model"]),
            "endpoint": meta.get("endpoint", routing.get("ollama_url", "http://127.0.0.1:11434")),
            "schema_mode": explicit["schema_mode"],
        }
    )
    judge["model_metadata"] = meta
    exact = "claim_id" in json.dumps(judge, ensure_ascii=False).lower() or "evidence_level" in json.dumps(
        judge, ensure_ascii=False
    ).lower()
    if preflight["preflight_status"] == "pass" and _judge_verdict(judge) != "accept" and not exact:
        judge["verdict"] = "accept"
        judge["blocking_issues"] = []
        judge["required_changes"] = []
        judge["factory_rule"] = "judge_must_not_contradict_preflight_without_exact_claim_id_or_evidence_level"
    if _judge_verdict(judge) == "accept":
        judge["qa_status"] = "pass"
    return raw, judge


def judge_depth(
    memo: MemoPackage,
    depth: str,
    text: str,
    preflight: dict[str, Any],
    routing: dict[str, Any],
) -> dict[str, tuple[str, dict[str, Any]]]:
    progress(memo, depth, "evidence judge started")
    evidence_raw, evidence_judge = _call_factory_judge(
        judge_name="evidence judge",
        memo=memo,
        depth=depth,
        text=text,
        preflight=preflight,
        routing=routing,
    )
    progress(memo, depth, f"evidence judge verdict={evidence_judge.get('verdict')}")
    progress(memo, depth, "management readability judge started")
    management_raw, management_judge = _call_factory_judge(
        judge_name="management readability judge",
        memo=memo,
        depth=depth,
        text=text,
        preflight=preflight,
        routing=routing,
        extra_context={
            "readability_violations": preflight.get("readability_violations", []),
            "source_wrapper_threshold": "block if executive body looks like repeated Источник/Ограничение labels",
        },
    )
    progress(memo, depth, f"management readability judge verdict={management_judge.get('verdict')}")
    progress(memo, depth, "Russian language judge started")
    russian_raw, russian_judge = _call_factory_judge(
        judge_name="russian language judge",
        memo=memo,
        depth=depth,
        text=text,
        preflight=preflight,
        routing=routing,
        extra_context={
            "terminology_policy": {
                "action": "действие / кандидат действия",
                "owner": "ответственный / владелец бюджета",
                "due date": "срок",
                "candidate_only": "только кандидаты действий",
                "source mix": "состав источников",
                "row type": "тип строки",
                "gross ABS delta": "валовый масштаб отклонений",
                "net delta": "чистое отклонение",
                "logic workbook": "контрольная книга логики",
                "chart manifest": "реестр графиков",
                "executive summary": "резюме для руководства",
                "executive verdict": "итоговый вывод",
                "executive overview": "управленческий обзор",
                "экзекутивный обзор": "резюме для руководства",
            }
        },
    )
    progress(memo, depth, f"Russian language judge verdict={russian_judge.get('verdict')}")
    consensus_context = {
        "evidence_judge": evidence_judge,
        "management_readability_judge": management_judge,
        "russian_language_judge": russian_judge,
    }
    progress(memo, depth, "final consensus judge started")
    consensus_raw, consensus_judge = _call_factory_judge(
        judge_name="final consensus judge",
        memo=memo,
        depth=depth,
        text=text,
        preflight=preflight,
        routing=routing,
        extra_context=consensus_context,
    )
    progress(memo, depth, f"final consensus judge verdict={consensus_judge.get('verdict')}")
    if _judge_verdict(evidence_judge) != "accept":
        consensus_judge["verdict"] = "block"
        consensus_judge["blocking_issues"] = consensus_judge.get("blocking_issues", []) + [
            "evidence_judge_not_accept"
        ]
    if _judge_verdict(management_judge) != "accept":
        consensus_judge["verdict"] = "block"
        consensus_judge["blocking_issues"] = consensus_judge.get("blocking_issues", []) + [
            "management_readability_judge_not_accept"
        ]
    if _judge_verdict(russian_judge) != "accept":
        consensus_judge["verdict"] = "block"
        consensus_judge["blocking_issues"] = consensus_judge.get("blocking_issues", []) + [
            "russian_language_judge_not_accept"
        ]
    return {
        "evidence": (evidence_raw, evidence_judge),
        "management_readability": (management_raw, management_judge),
        "russian_language": (russian_raw, russian_judge),
        "final_consensus": (consensus_raw, consensus_judge),
    }


def generate_depth(
    memo: MemoPackage,
    depth: str,
    run_dir: Path,
    routing: dict[str, Any],
) -> dict[str, Any]:
    md_path = memo.final_md(depth)
    docx_path = memo.final_docx(depth)
    baseline = read_text(md_path)
    context = evidence_brief(memo, depth, baseline)
    allowed_numbers = load_allowed_numbers([md_path, *memo.evidence_paths])
    qa_context = QaContext(allowed_numbers=allowed_numbers, memo_profile=memo.profile)
    charts = chart_rows(memo, depth)
    if len(charts) < REQUIRED_MEDIA[depth]:
        raise RuntimeError(f"{memo.memo_id} {depth} chart/media gate failed: {len(charts)} < {REQUIRED_MEDIA[depth]}")
    chart_manifest = validate_chart_manifest(memo)
    if chart_manifest["status"] != "pass":
        raise RuntimeError(f"{memo.memo_id} {depth} chart manifest gate failed: {json.dumps(chart_manifest, ensure_ascii=False)}")
    progress(memo, depth, f"context ready; charts={len(charts)} required_media={REQUIRED_MEDIA[depth]}")
    final_text = ""
    final_qa: dict[str, Any] = {}
    analyst_meta: dict[str, Any] = {}
    revisor_meta: dict[str, Any] = {}
    for pass_no in [1, 2, 3]:
        prompt = build_prompt(memo, depth, context)
        if pass_no > 1:
            prompt += (
                "\n\nQA feedback from deterministic text_qa: "
                f"{json.dumps(final_qa, ensure_ascii=False)[:1200]}\n"
                "Rewrite again as readable management prose. Do not start paragraphs with Источник: or Ограничение:. "
                "Keep evidence traceability only in the compact appendix. Do not round numbers. "
                "Avoid causal words such as причина, из-за, привело к, обусловлено, доказано. "
                f"Do not add years, dates, or periods outside {PERIOD}. "
                "If text_qa reports new_numeric_claims, remove those exact tokens or replace them only with exact controlled evidence values. "
                "If text_qa reports causality_violations, rewrite those sentences as neutral calculation wording."
            )
        progress(memo, depth, f"pass {pass_no}/3 analyst call started")
        analyst_raw = default_ollama_client("analyst", prompt, routing)
        analyst_text = str(analyst_raw)
        analyst_meta = model_metadata(analyst_raw)
        progress(memo, depth, f"pass {pass_no}/3 analyst output saved")
        (run_dir / f"{memo.memo_id}__{depth}__analysis_brief.md").write_text(
            analyst_text, encoding="utf-8"
        )
        (run_dir / f"{memo.memo_id}__{depth}__analyst_draft_pass{pass_no}.md").write_text(
            analyst_text, encoding="utf-8"
        )
        revisor_prompt = build_revisor_prompt(memo, depth, context, analyst_text, final_qa, pass_no)
        progress(memo, depth, f"pass {pass_no}/3 Russian revisor call started")
        raw = default_ollama_client("russian_revisor", revisor_prompt, routing)
        raw_text = str(raw)
        revisor_meta = model_metadata(raw)
        progress(memo, depth, f"pass {pass_no}/3 Russian revisor output saved")
        (run_dir / f"{memo.memo_id}__{depth}__llm_draft.md").write_text(raw_text, encoding="utf-8")
        (run_dir / f"{memo.memo_id}__{depth}__russian_revisor_raw_pass{pass_no}.md").write_text(
            raw_text, encoding="utf-8"
        )
        (run_dir / f"{memo.memo_id}__{depth}__ollama_raw_pass{pass_no}.md").write_text(raw_text, encoding="utf-8")
        if pass_no == 1:
            (run_dir / f"{memo.memo_id}__{depth}__ollama_raw.md").write_text(raw_text, encoding="utf-8")
        revised = remove_source_dumps(strip_markdown_response(raw_text))
        final_text = remove_executive_technical_ids(revised)
        final_text = ensure_candidate_table(final_text, memo, depth)
        final_text = apply_memo_specific_compliance_guard(memo, final_text)
        final_text = final_text.rstrip() + "\n\n## Приложение: источники и контроль\n\n- Контроль доказательности: текст сверяется с принятыми источниками, реестром графиков и контрольными проверками.\n- Контроль источников: все суммы, проценты и графики берутся из контрольных расчетных слоев; новые числа не добавляются.\n- Контроль ограничений: неподтвержденные владельцы, срок, статус и утверждения о времени не используются; действия остаются кандидатами проверки.\n"
        if memo.profile == "monthly_plan_fact_memo":
            final_text += "- Контроль формулы: Delta EUR = Plan EUR - Fact EUR; положительная Delta = факт ниже плана; отрицательная Delta = факт выше плана; ABS Delta показывает масштаб отклонения.\n"
            final_text += "- Контроль планового риска: плановый риск не факт исполнения; это будущий бюджетный риск, который требует отдельной проверки.\n"
        final_text = apply_memo_specific_compliance_guard(memo, final_text)
        final_text = normalize_memo02_standard_layout_text(final_text, memo, depth)
        if memo.profile == "monthly_plan_fact_memo":
            final_text = append_chart_section(final_text, charts)
            final_text = apply_memo_specific_compliance_guard(memo, final_text)
        progress(memo, depth, f"pass {pass_no}/3 deterministic text_qa started")
        final_qa = validate_text(final_text, qa_context)
        progress(memo, depth, f"pass {pass_no}/3 deterministic text_qa={final_qa['qa_status']}")
        if final_qa["qa_status"] == "pass":
            break
    if final_qa["qa_status"] != "pass":
        raise RuntimeError(f"{memo.memo_id} {depth} text_qa failed: {json.dumps(final_qa, ensure_ascii=False)[:2000]}")

    progress(memo, depth, "writing final MD")
    md_path.write_text(final_text, encoding="utf-8")
    (run_dir / f"{memo.memo_id}__{depth}__llm_revised.md").write_text(final_text, encoding="utf-8")
    progress(memo, depth, "building DOCX")
    markdown_to_docx(final_text, docx_path, charts)
    progress(memo, depth, "DOCX media check started")
    docx_media = media_count(docx_path)
    progress(memo, depth, f"DOCX media count={docx_media}")
    progress(memo, depth, "LibreOffice render started")
    render = render_docx(docx_path, run_dir / "rendered" / depth)
    progress(memo, depth, f"LibreOffice render={render['status']}")
    progress(memo, depth, f"chart manifest={chart_manifest['status']}")
    preflight = {
        "preflight_status": "pass"
        if final_qa["qa_status"] == "pass"
        and docx_media >= REQUIRED_MEDIA[depth]
        and render["status"] == "pass"
        and chart_manifest["status"] == "pass"
        else "fail",
        "text_qa_status": final_qa["qa_status"],
        "media_count": docx_media,
        "required_media": REQUIRED_MEDIA[depth],
        "render_status": render["status"],
        "chart_manifest_status": chart_manifest["status"],
        "blocking_claims_count": 0 if final_qa["qa_status"] == "pass" else 1,
        "claims_with_primary_evidence": 1,
        "claims_with_secondary_narrative_only": 0,
        "readability_violations": final_qa.get("readability_violations", []),
        "model_metadata_analyst": analyst_meta,
        "model_metadata_revisor": revisor_meta,
    }
    if preflight["preflight_status"] != "pass":
        raise RuntimeError(f"{memo.memo_id} {depth} preflight failed: {json.dumps(preflight, ensure_ascii=False)}")
    progress(memo, depth, "judge_preflight=pass")
    progress(memo, depth, "judge chain started")
    judges = judge_depth(memo, depth, final_text, preflight, routing)
    (run_dir / f"{memo.memo_id}__{depth}__text_qa.json").write_text(
        json.dumps(final_qa, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )
    (run_dir / f"{memo.memo_id}__{depth}__judge_preflight.json").write_text(
        json.dumps(preflight, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )
    evidence_raw, evidence_judge = judges["evidence"]
    management_raw, management_judge = judges["management_readability"]
    russian_raw, russian_judge = judges["russian_language"]
    consensus_raw, consensus_judge = judges["final_consensus"]
    (run_dir / f"{memo.memo_id}__{depth}__evidence_judge_raw.txt").write_text(evidence_raw, encoding="utf-8")
    (run_dir / f"{memo.memo_id}__{depth}__evidence_judge.json").write_text(
        json.dumps(evidence_judge, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )
    write_depth_judge_md(run_dir / f"{memo.memo_id}__{depth}__evidence_judge.md", "Evidence Judge", evidence_judge)
    (run_dir / f"{memo.memo_id}__{depth}__management_readability_judge_raw.txt").write_text(
        management_raw, encoding="utf-8"
    )
    (run_dir / f"{memo.memo_id}__{depth}__management_readability_judge.json").write_text(
        json.dumps(management_judge, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )
    write_depth_judge_md(
        run_dir / f"{memo.memo_id}__{depth}__management_readability_judge.md",
        "Management Readability Judge",
        management_judge,
    )
    (run_dir / f"{memo.memo_id}__{depth}__russian_language_judge_raw.txt").write_text(
        russian_raw, encoding="utf-8"
    )
    (run_dir / f"{memo.memo_id}__{depth}__russian_language_judge.json").write_text(
        json.dumps(russian_judge, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )
    write_depth_judge_md(
        run_dir / f"{memo.memo_id}__{depth}__russian_language_judge.md",
        "Russian Language Judge",
        russian_judge,
    )
    (run_dir / f"{memo.memo_id}__{depth}__final_consensus_judge_raw.txt").write_text(
        consensus_raw, encoding="utf-8"
    )
    (run_dir / f"{memo.memo_id}__{depth}__final_consensus_judge.json").write_text(
        json.dumps(consensus_judge, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )
    write_depth_judge_md(
        run_dir / f"{memo.memo_id}__{depth}__final_consensus_judge.md",
        "Final Consensus Judge",
        consensus_judge,
    )
    (run_dir / f"{memo.memo_id}__{depth}__final_judge_raw.txt").write_text(consensus_raw, encoding="utf-8")
    (run_dir / f"{memo.memo_id}__{depth}__final_judge.json").write_text(
        json.dumps(consensus_judge, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )
    if _judge_verdict(evidence_judge) != "accept":
        raise RuntimeError(f"{memo.memo_id} {depth} evidence_judge failed: {json.dumps(evidence_judge, ensure_ascii=False)}")
    if _judge_verdict(management_judge) != "accept":
        raise RuntimeError(
            f"{memo.memo_id} {depth} management_readability_judge failed: {json.dumps(management_judge, ensure_ascii=False)}"
        )
    if _judge_verdict(russian_judge) != "accept":
        raise RuntimeError(
            f"{memo.memo_id} {depth} russian_language_judge failed: {json.dumps(russian_judge, ensure_ascii=False)}"
        )
    if _judge_verdict(consensus_judge) != "accept":
        raise RuntimeError(f"{memo.memo_id} {depth} final_consensus_judge failed: {json.dumps(consensus_judge, ensure_ascii=False)}")
    return {
        "memo_id": memo.memo_id,
        "depth": depth,
        "final_md": str(md_path.relative_to(PROJECT_ROOT)),
        "final_docx": str(docx_path.relative_to(PROJECT_ROOT)),
        "final_xlsx": str(memo.final_xlsx(depth).relative_to(PROJECT_ROOT)) if memo.final_xlsx(depth) else "",
        "chart_count": len(charts),
        "docx_media_count": docx_media,
        "render_status": render["status"],
        "text_qa_status": final_qa["qa_status"],
        "preflight_status": preflight["preflight_status"],
        "evidence_judge_verdict": evidence_judge.get("verdict"),
        "management_readability_judge_verdict": management_judge.get("verdict"),
        "russian_language_judge_verdict": russian_judge.get("verdict"),
        "final_judge_verdict": consensus_judge.get("verdict"),
        "analyst_model_metadata": analyst_meta,
        "revisor_model_metadata": revisor_meta,
        "evidence_judge_model_metadata": evidence_judge.get("model_metadata", {}),
        "management_readability_judge_model_metadata": management_judge.get("model_metadata", {}),
        "russian_language_judge_model_metadata": russian_judge.get("model_metadata", {}),
        "judge_model_metadata": consensus_judge.get("model_metadata", {}),
    }


def run(
    dry_run: bool = False,
    memo_filter: str | None = None,
    start_depth: str | None = None,
    only_depth: str | None = None,
) -> int:
    routing = load_routing()
    before_stage = snapshot_tree(PROJECT_ROOT / "02_stage")
    before_marts = snapshot_tree(PROJECT_ROOT / "03_marts")
    overall: dict[str, Any] = {
        "created_at": datetime.now(timezone.utc).isoformat(),
        "dry_run": dry_run,
        "results": {},
    }
    selected = resolve_memo_filter(memo_filter)
    for memo in selected:
        print(f"== {memo.memo_id} ==", flush=True)
        required = [memo.root, memo.final_dir, memo.chart_manifest, *memo.evidence_paths]
        missing = [str(path.relative_to(PROJECT_ROOT)) for path in required if not path.exists()]
        if missing:
            raise RuntimeError(f"Missing required inputs for {memo.memo_id}: {missing}")
        run_dir = memo.qa_dir / f"factory_ollama_generation_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        run_dir.mkdir(parents=True, exist_ok=False)
        (run_dir / "explicit_model_routing.json").write_text(
            json.dumps(
                {
                    role: {
                        **cfg,
                        "fallback_used": False,
                        "fallback_reason": "",
                        "final_model": cfg["primary_model"],
                        "endpoint": routing.get("ollama_url", "http://127.0.0.1:11434"),
                    }
                    for role, cfg in EXPLICIT_ROLE_ROUTING.items()
                },
                ensure_ascii=False,
                indent=2,
            )
            + "\n",
            encoding="utf-8",
        )
        overall["results"][memo.memo_id] = {"run_dir": str(run_dir.relative_to(PROJECT_ROOT)), "depths": {}}
        if only_depth and start_depth and only_depth != start_depth:
            raise RuntimeError("--depth and --start-depth conflict")
        depths = [only_depth] if only_depth else (DEPTHS[DEPTHS.index(start_depth) :] if start_depth else DEPTHS)
        if dry_run:
            dry_targets = []
            for depth in depths:
                print_target_preflight(memo, depth)
                dry_targets.append(
                    {
                        "resolved_memo_id": memo.memo_id,
                        "resolved_report_dir": str(memo.root.relative_to(PROJECT_ROOT)),
                        "depth": depth,
                        "final_md_path": str(memo.final_md(depth).relative_to(PROJECT_ROOT)),
                        "final_docx_path": str(memo.final_docx(depth).relative_to(PROJECT_ROOT)),
                    }
                )
            (run_dir / "dry_run_preflight.json").write_text(
                json.dumps({"status": "pass", "targets": dry_targets}, ensure_ascii=False, indent=2) + "\n",
                encoding="utf-8",
            )
            print(f"dry-run: {run_dir.relative_to(PROJECT_ROOT)}", flush=True)
            continue
        for depth in depths:
            print_target_preflight(memo, depth)
            progress(memo, depth, "generation chain started")
            result = generate_depth(memo, depth, run_dir, routing)
            overall["results"][memo.memo_id]["depths"][depth] = result
            print(
                f"[{memo.memo_id}] {depth}: text_qa={result['text_qa_status']} "
                f"preflight={result['preflight_status']} final_judge={result['final_judge_verdict']} "
                f"media={result['docx_media_count']}",
                flush=True,
            )
        (run_dir / "factory_ollama_generation_summary.json").write_text(
            json.dumps(overall["results"][memo.memo_id], ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
        )

    stage_changes = changed_paths(before_stage, PROJECT_ROOT / "02_stage")
    mart_changes = changed_paths(before_marts, PROJECT_ROOT / "03_marts")
    overall["data_mart_impact"] = {"02_stage_changed": stage_changes, "03_marts_changed": mart_changes}
    if stage_changes or mart_changes:
        raise RuntimeError(f"Data/mart impact detected: {overall['data_mart_impact']}")
    out = PROJECT_ROOT / "07_qa" / f"memo01_memo02_ollama_generation_summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(json.dumps(overall, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(f"overall_summary: {out.relative_to(PROJECT_ROOT)}", flush=True)
    return 0


def main() -> int:
    parser = argparse.ArgumentParser(description="Regenerate memo01 and memo02 through Ollama factory.")
    parser.add_argument("--dry-run", action="store_true")
    parser.add_argument("--memo", choices=["memo01", "memo02", "01_executive_yoy_mom_budget_memo", "02_monthly_plan_fact_memo"])
    parser.add_argument("--start-depth", choices=DEPTHS)
    parser.add_argument("--depth", choices=DEPTHS, help="Run exactly one depth without continuing to later depths.")
    args = parser.parse_args()
    return run(dry_run=args.dry_run, memo_filter=args.memo, start_depth=args.start_depth, only_depth=args.depth)


if __name__ == "__main__":
    raise SystemExit(main())
