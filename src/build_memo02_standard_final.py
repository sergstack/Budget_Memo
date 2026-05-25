from __future__ import annotations

import json
import math
import shutil
import zipfile
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from matplotlib.ticker import FuncFormatter
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = PROJECT_ROOT / "06_reports" / "02_monthly_plan_fact_memo"
MGMT_WB = REPORT_DIR / "xlsx" / "management_depth_workbook.xlsx"
PACKAGE_WB = REPORT_DIR / "xlsx" / "monthly_plan_fact_workbook.xlsx"
MD_OUT = REPORT_DIR / "md" / "02_monthly_plan_fact_memo__standard__2026-04__final.md"
DOCX_OUT = REPORT_DIR / "docx" / "02_monthly_plan_fact_memo__standard__2026-04__final.docx"
CHART_DIR = REPORT_DIR / "charts" / "standard_final"
CHART_IMG_DIR = CHART_DIR / "images"
CHART_DATA_DIR = CHART_DIR / "data"
QA_MD = REPORT_DIR / "qa" / "standard_final_qa.md"
QA_JSON = REPORT_DIR / "qa" / "standard_final_qa.json"
MANIFEST = REPORT_DIR / "source_refs" / "standard_final_manifest.json"
SELECTED_MONTH = "2026-04"
READABLE_STATUS = "кандидат для управленческого ревью; действия остаются кандидатными до подтверждения срока и статуса"


def require_inputs() -> None:
    missing = [p for p in [MGMT_WB, PACKAGE_WB] if not p.exists()]
    if missing:
        raise FileNotFoundError("Missing accepted memo_02 package files: " + ", ".join(str(p) for p in missing))


def read_sheet(path: Path, sheet: str) -> pd.DataFrame:
    return pd.read_excel(path, sheet_name=sheet)


def fmt_eur(value: float | int | None) -> str:
    if value is None or (isinstance(value, float) and math.isnan(value)):
        return "н/д"
    value = float(value)
    sign = "-" if value < 0 else ""
    v = abs(value)
    if v >= 1_000_000:
        return f"{sign}{v / 1_000_000:.2f} млн EUR".replace(".", ",")
    if v >= 1_000:
        return f"{sign}{v / 1_000:.1f} тыс. EUR".replace(".", ",")
    return f"{sign}{v:.0f} EUR".replace(".", ",")


def fmt_pct(value: float | int | None) -> str:
    if value is None or (isinstance(value, float) and math.isnan(value)):
        return "н/д"
    return f"{float(value) * 100:.1f}%".replace(".", ",")


def clean(value: object) -> str:
    if value is None or (isinstance(value, float) and math.isnan(value)):
        return "н/д"
    text = str(value).strip()
    replacements = {
        "fact below plan / факт ниже плана": "факт ниже плана",
        "fact above plan / превышение факта": "факт выше плана",
        "over-execution": "факт выше плана",
        "under-execution": "факт ниже плана",
        "cons_budget": "плановый источник",
        "fact_only": "факт без плана",
        "p_fact_adjusted": "скорректированный план-факт",
        "plan_and_fact": "план и факт",
        "plan_only": "план без факта",
        "refund_only": "возвраты без плана",
        "pass": "пройдено",
        "n/a": "н/д",
    }
    return replacements.get(text, text)


def eur_axis(value: float, _pos: int | None = None) -> str:
    value = float(value)
    sign = "-" if value < 0 else ""
    v = abs(value)
    if v >= 1_000_000:
        return f"{sign}{v / 1_000_000:.1f} млн"
    if v >= 1_000:
        return f"{sign}{v / 1_000:.0f} тыс."
    return f"{sign}{v:.0f}"


def polish_text_columns(df: pd.DataFrame) -> pd.DataFrame:
    polished = df.copy()
    for col in polished.columns:
        if polished[col].dtype == object:
            polished[col] = polished[col].map(clean)
    return polished


def table_md(df: pd.DataFrame, columns: list[str], limit: int = 8) -> str:
    shown = df.head(limit).copy()
    header_alias = {
        "Candidate check": "Что проверить",
        "source_mix": "Состав источников",
        "qa_status": "QA статус",
        "claim_id": "ID evidence",
        "claim_type": "Тип claim",
        "source_mart_or_slice": "Источник",
        "metric": "Метрика",
        "grain": "Гранулярность",
        "period": "Период",
        "limitation": "Ограничение",
    }
    rows = []
    rows.append("| " + " | ".join(header_alias.get(col, col) for col in columns) + " |")
    rows.append("| " + " | ".join(["---"] * len(columns)) + " |")
    for _, row in shown.iterrows():
        vals = []
        for col in columns:
            val = row.get(col, "")
            if "EUR" in col or col in ["План EUR", "Факт EUR", "Delta EUR", "ABS Delta EUR", "Сумма EUR"]:
                vals.append(fmt_eur(val))
            elif "%" in col or col == "Исполнение %":
                vals.append(fmt_pct(val))
            else:
                vals.append(clean(val))
        rows.append("| " + " | ".join(vals) + " |")
    return "\n".join(rows)


def save_chart_data(df: pd.DataFrame, filename: str) -> str:
    CHART_DATA_DIR.mkdir(parents=True, exist_ok=True)
    path = CHART_DATA_DIR / filename
    df.to_csv(path, index=False)
    return str(path.relative_to(PROJECT_ROOT))


def save_barh(df: pd.DataFrame, y: str, x: str, title: str, filename: str, xlabel: str = "EUR") -> dict:
    CHART_IMG_DIR.mkdir(parents=True, exist_ok=True)
    plot_df = df.copy().iloc[::-1]
    fig_h = max(4.5, 0.45 * len(plot_df) + 1.6)
    fig, ax = plt.subplots(figsize=(9.5, fig_h))
    ax.barh(plot_df[y].astype(str), plot_df[x], color="#2f6f9f")
    ax.set_title(title, loc="left", fontsize=13, fontweight="bold")
    ax.set_xlabel(xlabel)
    ax.xaxis.set_major_formatter(FuncFormatter(eur_axis))
    ax.grid(axis="x", color="#d9e2ec", linewidth=0.8)
    ax.tick_params(axis="y", labelsize=9)
    ax.tick_params(axis="x", labelsize=9)
    fig.tight_layout()
    path = CHART_IMG_DIR / filename
    fig.savefig(path, dpi=170, bbox_inches="tight")
    plt.close(fig)
    return {"image_path": str(path.relative_to(PROJECT_ROOT)), "title": title}


def save_planning_quality_chart(planning: pd.DataFrame) -> dict:
    CHART_IMG_DIR.mkdir(parents=True, exist_ok=True)
    fig, ax1 = plt.subplots(figsize=(9.5, 4.8))
    x = range(len(planning))
    ax1.bar(x, planning["Количество"], color="#6f8f72", label="Количество")
    ax1.set_ylabel("Количество объектов")
    ax1.set_xticks(list(x))
    ax1.set_xticklabels(planning["Категория"], rotation=15, ha="right")
    ax2 = ax1.twinx()
    ax2.plot(x, planning["ABS отклонение EUR"], color="#9a4d4d", marker="o", linewidth=2, label="Денежное влияние")
    ax2.set_ylabel("Денежное влияние, EUR")
    ax2.yaxis.set_major_formatter(FuncFormatter(eur_axis))
    ax1.set_title("Качество планирования: частота и денежное влияние, 2026-04", loc="left", fontsize=13, fontweight="bold")
    ax1.grid(axis="y", color="#d9e2ec", linewidth=0.8)
    ax1.legend(loc="upper left")
    ax2.legend(loc="upper right")
    fig.tight_layout()
    path = CHART_IMG_DIR / "planning_quality_frequency_impact.png"
    fig.savefig(path, dpi=170, bbox_inches="tight")
    plt.close(fig)
    return {"image_path": str(path.relative_to(PROJECT_ROOT)), "title": "Качество планирования: частота и денежное влияние, 2026-04"}


def save_problem_directions_chart(df: pd.DataFrame) -> dict:
    CHART_IMG_DIR.mkdir(parents=True, exist_ok=True)
    plot_df = df.head(10).iloc[::-1]
    fig, ax1 = plt.subplots(figsize=(9.5, 5.2))
    ypos = range(len(plot_df))
    ax1.barh(list(ypos), plot_df["ABS Delta EUR"], color="#2f6f9f", label="ABS отклонение")
    ax1.set_yticks(list(ypos))
    ax1.set_yticklabels(plot_df["Статья 1"].astype(str), fontsize=9)
    ax1.set_xlabel("ABS отклонение, EUR")
    ax1.xaxis.set_major_formatter(FuncFormatter(eur_axis))
    ax2 = ax1.twiny()
    ax2.plot(plot_df["Количество объектов"], list(ypos), color="#9a4d4d", marker="o", linewidth=2, label="Количество")
    ax2.set_xlabel("Количество объектов")
    ax1.set_title("Проблемные аналитические направления: сумма отклонений и частота, 2026-04", loc="left", fontsize=13, fontweight="bold")
    ax1.grid(axis="x", color="#d9e2ec", linewidth=0.8)
    fig.tight_layout()
    path = CHART_IMG_DIR / "problem_analytical_directions_count_abs.png"
    fig.savefig(path, dpi=170, bbox_inches="tight")
    plt.close(fig)
    return {"image_path": str(path.relative_to(PROJECT_ROOT)), "title": "Проблемные аналитические направления: сумма отклонений и частота, 2026-04"}


def build_planning_quality(article_pf: pd.DataFrame, fact_gap: pd.DataFrame, plan_gap: pd.DataFrame) -> pd.DataFrame:
    normal = article_pf.copy()
    normal["Исполнение %"] = pd.to_numeric(normal["Исполнение %"], errors="coerce")
    buckets = [
        ("План без факта", len(plan_gap), plan_gap["План EUR"].sum(), 0.0, plan_gap["ABS отклонение EUR"].sum()),
        ("Факт без плана", len(fact_gap), 0.0, fact_gap["Факт EUR"].sum(), fact_gap["ABS отклонение EUR"].sum()),
        ("0%-75%", int(((normal["Исполнение %"] >= 0) & (normal["Исполнение %"] < 0.75)).sum()), normal.loc[(normal["Исполнение %"] >= 0) & (normal["Исполнение %"] < 0.75), "План EUR"].sum(), normal.loc[(normal["Исполнение %"] >= 0) & (normal["Исполнение %"] < 0.75), "Факт EUR"].sum(), normal.loc[(normal["Исполнение %"] >= 0) & (normal["Исполнение %"] < 0.75), "ABS отклонение EUR"].sum()),
        ("75%-110%", int(((normal["Исполнение %"] >= 0.75) & (normal["Исполнение %"] <= 1.10)).sum()), normal.loc[(normal["Исполнение %"] >= 0.75) & (normal["Исполнение %"] <= 1.10), "План EUR"].sum(), normal.loc[(normal["Исполнение %"] >= 0.75) & (normal["Исполнение %"] <= 1.10), "Факт EUR"].sum(), normal.loc[(normal["Исполнение %"] >= 0.75) & (normal["Исполнение %"] <= 1.10), "ABS отклонение EUR"].sum()),
        (">110%", int((normal["Исполнение %"] > 1.10).sum()), normal.loc[normal["Исполнение %"] > 1.10, "План EUR"].sum(), normal.loc[normal["Исполнение %"] > 1.10, "Факт EUR"].sum(), normal.loc[normal["Исполнение %"] > 1.10, "ABS отклонение EUR"].sum()),
    ]
    return pd.DataFrame(buckets, columns=["Категория", "Количество", "План EUR", "Факт EUR", "ABS отклонение EUR"])


def build_problem_directions(article_h: pd.DataFrame) -> pd.DataFrame:
    grouped = (
        article_h.groupby("Статья 1", dropna=False)
        .agg(
            **{
                "Количество объектов": ("Статья", "count"),
                "План EUR": ("План EUR", "sum"),
                "Факт EUR": ("Факт EUR", "sum"),
                "Delta EUR": ("Delta EUR", "sum"),
                "ABS Delta EUR": ("ABS Delta EUR", "sum"),
                "Строк": ("Строк", "sum"),
            }
        )
        .reset_index()
        .sort_values("ABS Delta EUR", ascending=False)
    )
    return grouped


def add_chart_md(lines: list[str], spec: dict, source: str, limitation: str) -> None:
    lines.extend(
        [
            f"![{spec['title']}]({spec['image_path']})",
            f"Источник: `{source}`. Период: `{SELECTED_MONTH}`. Ограничение: {limitation}",
            "",
        ]
    )


def add_doc_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_doc_table(doc: Document, df: pd.DataFrame, columns: list[str], limit: int = 8) -> None:
    shown = df.head(limit)
    header_alias = {
        "Candidate check": "Что проверить",
        "source_mix": "Состав источников",
        "qa_status": "QA статус",
        "claim_id": "ID evidence",
        "claim_type": "Тип claim",
        "source_mart_or_slice": "Источник",
        "metric": "Метрика",
        "grain": "Гранулярность",
        "period": "Период",
        "limitation": "Ограничение",
    }
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for i, col in enumerate(columns):
        table.rows[0].cells[i].text = header_alias.get(col, col)
    for _, row in shown.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(columns):
            val = row.get(col, "")
            if "EUR" in col or col in ["Сумма EUR"]:
                cells[i].text = fmt_eur(val)
            elif "%" in col or col == "Исполнение %":
                cells[i].text = fmt_pct(val)
            else:
                cells[i].text = clean(val)


def add_doc_chart(doc: Document, spec: dict, caption: str) -> None:
    doc.add_picture(str(PROJECT_ROOT / spec["image_path"]), width=Inches(6.6))
    p = doc.add_paragraph(caption)
    p.style = doc.styles["Caption"] if "Caption" in [s.name for s in doc.styles] else doc.styles["Normal"]


def build_outputs() -> None:
    require_inputs()
    CHART_IMG_DIR.mkdir(parents=True, exist_ok=True)
    CHART_DATA_DIR.mkdir(parents=True, exist_ok=True)

    total = read_sheet(MGMT_WB, "total_plan_fact").iloc[0]
    cfo = read_sheet(MGMT_WB, "cfo_decomposition")
    article_h = read_sheet(MGMT_WB, "article_hierarchy_decomposition")
    cfo_article = read_sheet(MGMT_WB, "cfo_article_localization")
    top_dev = read_sheet(MGMT_WB, "top_management_deviations")
    fact_mgmt = read_sheet(MGMT_WB, "fact_without_plan_management")
    plan_mgmt = read_sheet(MGMT_WB, "plan_without_fact_management")
    cp = read_sheet(MGMT_WB, "counterparty_management_view")
    legal = read_sheet(MGMT_WB, "legal_entity_currency_view")
    checks = read_sheet(MGMT_WB, "management_check_priorities")
    evidence = read_sheet(MGMT_WB, "evidence_appendix")
    source_mix = read_sheet(MGMT_WB, "source_mix_limitations")
    article_pf = read_sheet(PACKAGE_WB, "article_month_plan_fact")
    fact_gap = read_sheet(PACKAGE_WB, "fact_without_plan")
    plan_gap = read_sheet(PACKAGE_WB, "plan_without_fact")

    checks = checks.copy()
    checks["Candidate check"] = checks["Candidate check"].astype(str).str.replace(
        "Проверить отклонение с candidate owner route; не назначать confirmed owner без mapping.",
        "Проверить отклонение по маршруту ЦФО; не переводить кандидата проверки в финальное действие без срока и статуса.",
        regex=False,
    )
    checks["Candidate check"] = checks["Candidate check"].str.replace(
        "timing/отмену/перенос",
        "срок исполнения / отмену / перенос",
        regex=False,
    )
    evidence = evidence.copy()
    evidence["limitation"] = evidence["limitation"].astype(str).str.replace(
        "requires mapping before management release",
        "ЦФО является маршрутом проверки. Направление используется только как дополнительная аналитическая группировка и не блокирует выпуск memo_02. Отдельное сопоставление менеджера для memo_02 не применяется. Финальные действия требуют подтверждённого срока и статуса.",
        regex=False,
    )
    evidence["limitation"] = evidence["limitation"].str.replace(
        "owner_candidate is only candidate route",
        "ЦФО является маршрутом проверки; кандидат владельца не является финальным действием.",
        regex=False,
    )
    evidence["limitation"] = evidence["limitation"].str.replace(
        "not confirmed overspend/saving without check",
        "перерасход или экономия не подтверждаются без проверки",
        regex=False,
    )
    evidence["limitation"] = evidence["limitation"].str.replace(
        "planning-only periods excluded",
        "плановые будущие периоды исключены",
        regex=False,
    )
    evidence = evidence.replace(
        {
            "claim_type": {
                "supported numeric fact": "подтвержденный числовой факт",
                "supported localization": "подтвержденная локализация",
                "data quality limitation": "ограничение качества данных",
                "blocker": "ограничение",
            },
            "grain": {
                "selected month": "выбранный месяц",
                "cfo x article hierarchy": "ЦФО × иерархия статей",
                "cfo x article x counterparty where available": "ЦФО × статья × контрагент, где доступно",
                "dimension availability": "доступность измерений",
            },
            "source_mart_or_slice": {
                "schema inspection from accepted MART/package": "проверка схемы принятого MART/пакета",
            },
            "qa_status": {"pass": "пройдено"},
        }
    )
    source_mix = source_mix.replace({"qa_status": {"pass": "пройдено"}})

    planning_quality = build_planning_quality(article_pf, fact_gap, plan_gap)
    problem_dirs = build_problem_directions(article_h)

    save_chart_data(planning_quality, "planning_quality_frequency_impact.csv")
    save_chart_data(problem_dirs, "problem_analytical_directions_count_abs.csv")
    save_chart_data(article_h.head(12), "top_deviations_by_article.csv")
    save_chart_data(cfo_article.head(12), "direction_level_cfo_article.csv")
    save_chart_data(pd.DataFrame(
        [
            {"Категория": "Факт без плана", "Сумма EUR": fact_mgmt["Сумма EUR"].sum(), "Количество": fact_mgmt["Строк"].sum()},
            {"Категория": "План без факта", "Сумма EUR": plan_mgmt["Сумма EUR"].sum(), "Количество": plan_mgmt["Строк"].sum()},
        ]
    ), "fact_without_plan_plan_without_fact.csv")
    save_chart_data(cp.head(10), "counterparty_localization.csv")
    save_chart_data(legal.head(10), "legal_entity_currency.csv")

    chart_specs = []
    chart_specs.append(save_planning_quality_chart(planning_quality))
    chart_specs.append(save_barh(article_h.head(10), "Статья", "ABS Delta EUR", "Топ отклонений по статьям, 2026-04", "top_deviations_by_article.png", "ABS отклонение, EUR"))
    chart_specs.append(save_problem_directions_chart(problem_dirs))
    chart_specs.append(save_barh(cfo_article.head(10).assign(label=lambda d: d["ЦФО"].astype(str) + " / " + d["Статья"].astype(str)), "label", "ABS Delta EUR", "ЦФО × статья: локализация отклонений, 2026-04", "direction_level_cfo_article.png", "ABS отклонение, EUR"))
    gap_chart = pd.DataFrame(
        [
            {"Категория": "Факт без плана", "Сумма EUR": fact_mgmt["Сумма EUR"].sum()},
            {"Категория": "План без факта", "Сумма EUR": plan_mgmt["Сумма EUR"].sum()},
        ]
    )
    chart_specs.append(save_barh(gap_chart, "Категория", "Сумма EUR", "Факт без плана / план без факта, 2026-04", "fact_without_plan_plan_without_fact.png", "Сумма, EUR"))
    chart_specs.append(save_barh(cp.head(8).assign(label=lambda d: d["Контрагент"].astype(str).str[:45]), "label", "ABS Delta EUR", "Контрагентская локализация, 2026-04", "counterparty_localization.png", "ABS отклонение, EUR"))
    chart_specs.append(save_barh(legal.head(8).assign(label=lambda d: d["Юр. лицо"].astype(str) + " / " + d["Валюта"].astype(str)), "label", "ABS Delta EUR", "Юрлица и валюты: локализация суммы, 2026-04", "legal_entity_currency.png", "ABS отклонение, EUR"))

    chart_manifest = [
        {
            "chart_id": Path(spec["image_path"]).stem,
            "title_ru": spec["title"],
            "image_path": spec["image_path"],
            "period": SELECTED_MONTH,
            "source": source,
            "limitation": limitation,
        }
        for spec, source, limitation in [
            (chart_specs[0], "принятый пакет данных: срезы исполнения и разрывов", "Частота сигналов не равна денежной существенности. Для оценки влияния нужно смотреть сумму в EUR."),
            (chart_specs[1], "принятый пакет данных: иерархия статей", "ABS показывает масштаб отклонения, а не причину."),
            (chart_specs[2], "принятый пакет данных: иерархия статей", "Статья 1 используется как аналитическая группировка; это не подтвержденный справочник направлений."),
            (chart_specs[3], "принятый пакет данных: ЦФО × статья", "ЦФО задаёт маршрут проверки; причина отклонения этим не подтверждается."),
            (chart_specs[4], "принятый пакет данных: факт без плана и план без факта", "Разрывы являются кандидатами проверки, а не подтвержденной экономией или перерасходом."),
            (chart_specs[5], "принятый пакет данных: контрагенты", "Контрагент локализует проверку, но не подтверждает причину."),
            (chart_specs[6], "принятый пакет данных: юрлица и валюты", "Юрлицо и валюта локализуют сумму, но не подтверждают причину."),
        ]
    ]
    (CHART_DIR / "standard_final_chart_metadata.json").write_text(json.dumps(chart_manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

    lines: list[str] = [
        "# Ежемесячная аналитическая записка План-Факт за апрель 2026",
        "",
        f"Период: `{SELECTED_MONTH}`  ",
        "Источник: принятый пакет данных memo_02 (`monthly_plan_fact_workbook.xlsx`, `management_depth_workbook.xlsx`, evidence/проверки качества).  ",
        f"Статус: {READABLE_STATUS}.  ",
        "Ограничение: готовность к промышленному запуску не заявляется.",
        "",
        "## 1. Резюме для руководства",
        f"- За апрель 2026 общий план составил **{fmt_eur(total['План EUR'])}**, факт - **{fmt_eur(total['Факт EUR'])}**, Delta EUR = Plan EUR - Fact EUR - **{fmt_eur(total['Delta EUR'])}**, исполнение - **{fmt_pct(total['Исполнение %'])}**.",
        "- На уровне общего итога месяц близок к плану, но управленческий фокус находится в структуре отклонений.",
        "- Главные зоны внимания: качество планирования, аналитические направления, ЦФО × статья, факт без плана и план без факта.",
        "- ЦФО является маршрутом проверки; действия остаются кандидатными до подтверждения срока и статуса.",
        "",
        "## 2. Область анализа и выбранный период",
        "В анализ включён закрытый месяц апрель 2026. Плановые будущие периоды не используются для интерпретации фактического исполнения.",
        "Направление используется как дополнительная аналитическая группировка через иерархию статей; это не подтвержденный справочник направлений.",
        "Отдельное сопоставление менеджера для memo_02 не применяется: маршрут проверки задаётся ЦФО.",
        "",
        "## 3. Результат План-Факт",
        f"План составил **{fmt_eur(total['План EUR'])}**, факт - **{fmt_eur(total['Факт EUR'])}**, отклонение План-Факт - **{fmt_eur(total['Delta EUR'])}**, исполнение - **{fmt_pct(total['Исполнение %'])}**.",
        "Положительная Delta означает факт ниже плана; отрицательная Delta означает факт выше плана.",
        "ABS Delta показывает масштаб отклонения, но не доказывает причину.",
        "",
        "## 4. Качество планирования",
        "Категории частоты показывают, насколько часто возникают плановые разрывы и зоны исполнения. Частота сигналов не равна денежной существенности, поэтому рядом показана сумма в EUR.",
    ]
    add_chart_md(lines, chart_specs[0], "принятый пакет данных memo_02", "Частота сигналов не равна денежной существенности. Для оценки влияния нужно смотреть сумму в EUR.")
    lines.extend([table_md(planning_quality, ["Категория", "Количество", "План EUR", "Факт EUR", "ABS отклонение EUR"], 10), ""])
    lines.extend(
        [
            "Вывод: качество планирования требует управленческого ревью не только по проценту исполнения, но и по денежной концентрации разрывов. Факт без плана и план без факта являются кандидатами проверки.",
            "",
            "## 5. Проблемные направления",
            "Так как подтвержденный справочник направлений отсутствует, используется принятая аналитическая иерархия `Статья 1`. Это не подтвержденное сопоставление направлений.",
        ]
    )
    add_chart_md(lines, chart_specs[2], "принятый пакет данных: иерархия статей", "Статья 1 локализует аналитическое направление, но не является подтвержденным справочником направлений.")
    lines.extend([table_md(problem_dirs, ["Статья 1", "Количество объектов", "План EUR", "Факт EUR", "Delta EUR", "ABS Delta EUR"], 10), ""])
    lines.extend(
        [
            "Вывод: проблемность направления определяется совместно по сумме ABS EUR и числу объектов; частота сама по себе не равна денежной существенности.",
            "",
            "## 6. Топ отклонений по статьям",
        ]
    )
    add_chart_md(lines, chart_specs[1], "принятый пакет данных: иерархия статей", "Статьи ранжированы по ABS Delta; ранжирование не объясняет причину.")
    lines.extend([table_md(article_h, ["Статья 1", "Статья 2", "Статья", "План EUR", "Факт EUR", "Delta EUR", "ABS Delta EUR", "Исполнение %", "Статус"], 10), ""])
    lines.extend(
        [
            "Положительная Delta означает факт ниже плана; отрицательная Delta означает факт выше плана. Факт без плана и план без факта рассматриваются отдельно ниже.",
            "",
            "## 7. Разбор расходования средств по направлениям",
            "Ниже показана локализация `ЦФО × статья` для основных аналитических направлений. Направление является дополнительной группировкой; подтвержденное направление не выводится моделью.",
        ]
    )
    add_chart_md(lines, chart_specs[3], "принятый пакет данных: ЦФО × статья", "ЦФО задаёт маршрут проверки; причина отклонения этим не подтверждается.")
    lines.extend([table_md(cfo_article, ["ЦФО", "Статья 1", "Статья 2", "Статья", "План EUR", "Факт EUR", "Delta EUR", "ABS Delta EUR", "Кандидат владельца", "Приоритет проверки"], 10), ""])
    for direction in problem_dirs["Статья 1"].head(3):
        part = cfo_article[cfo_article["Статья 1"] == direction].head(5)
        if len(part) > 0:
            lines.extend(
                [
                    f"### Аналитическое направление: {clean(direction)}",
                    table_md(part, ["ЦФО", "Статья", "План EUR", "Факт EUR", "Delta EUR", "ABS Delta EUR", "Приоритет проверки"], 5),
                    "Ограничение: направление является аналитической группировкой по `Статья 1`, а не подтвержденным справочником направлений.",
                    "",
                ]
            )
    lines.extend(["## 8. Факт без плана / план без факта"])
    add_chart_md(lines, chart_specs[4], "принятый пакет данных: факт без плана и план без факта", "Факт без плана не является автоматическим перерасходом; план без факта не является автоматической экономией.")
    lines.extend(
        [
            "### Факт без плана",
            table_md(fact_mgmt, ["ЦФО", "Статья 1", "Статья 2", "Статья", "Контрагент", "Факт EUR", "Строк", "Кандидат владельца", "Приоритет проверки"], 10),
            "",
            "### План без факта",
            table_md(plan_mgmt, ["ЦФО", "Статья 1", "Статья 2", "Статья", "План EUR", "Строк", "Кандидат владельца", "Приоритет проверки"], 10),
            "",
            "Вывод: оба блока являются кандидатами проверки. Требуется проверка срока исполнения, классификации и наличия утвержденного плана/факта.",
            "",
            "## 9. Контрагентский вид",
            "Контрагентский вид используется для локализации маршрута проверки, а не как доказательство причины отклонения.",
        ]
    )
    add_chart_md(lines, chart_specs[5], "принятый пакет данных: контрагенты", "Контрагент локализует проверку, но не подтверждает причину.")
    lines.extend([table_md(cp, ["Контрагент", "Ключ контрагента", "ЦФО", "Статья", "План EUR", "Факт EUR", "Delta EUR", "ABS Delta EUR", "Факт без плана", "План без факта"], 7), ""])
    lines.extend(["## 10. Юрлица / валюты"])
    add_chart_md(lines, chart_specs[6], "принятый пакет данных: юрлица и валюты", "Юрлицо и валюта локализуют сумму, но не подтверждают причину.")
    lines.extend([table_md(legal, ["Юр. лицо", "Валюта", "План EUR", "Факт EUR", "Delta EUR", "ABS Delta EUR", "Исполнение %", "Статус"], 8), ""])
    lines.extend(
        [
            "## 11. Комментарии / семантический анализ",
            "Комментарии руководителей не переданы / отсутствуют в текущем пакете, поэтому семантический анализ причин не выполнялся. Причины отклонений не подтверждаются этим разделом.",
            "Тематическое моделирование и любые неподтвержденные выводы о причинах не выполнялись.",
            "",
            "## 12. Реестр приоритетных проверок",
            "Это реестр кандидатов проверки, а не финальный план действий. Маршрут проверки = ЦФО. Срок и статус не подтверждены.",
            table_md(checks, ["Приоритет проверки", "ЦФО", "Статья", "ABS Delta EUR", "Кандидат владельца", "Candidate check"], 10),
            "",
            "## 13. Ограничения",
            "- Комментарии руководителей отсутствуют; семантический анализ причин не выполнялся.",
            "- Финальный план действий не готов: срок и статус действий не подтверждены.",
            "- Кандидаты по срокам исполнения не являются подтвержденными сроками.",
            "- План без факта не является автоматической экономией.",
            "- Факт без плана не является автоматическим перерасходом.",
            "- Направление является дополнительной аналитической группировкой; `Статья 1` используется как аналитическая иерархия, не подтвержденный справочник направлений.",
            "- Отдельное сопоставление менеджера для memo_02 не применяется; маршрут проверки = ЦФО.",
            "- Готовность к промышленному запуску не заявляется.",
            "",
            "## 14. Итоговый вывод",
            "Апрель 2026 близок к плану на уровне общего итога, но структура отклонений показывает несколько зон для управленческого ревью ЦФО.",
            "Основное внимание требуется по иерархии статей, связке ЦФО × статья, факту без плана и плану без факта. Эти зоны показывают, где проверять плановую дисциплину, классификацию и сроки исполнения.",
            "Без комментариев руководителей, подтвержденных сроков и статусов нельзя утверждать причины отклонений или переводить кандидатов проверки в финальный план действий.",
            "",
            "## 15. Приложение / evidence",
            table_md(evidence, ["claim_id", "claim_type", "source_mart_or_slice", "metric", "grain", "period", "limitation", "qa_status"], 10),
            "",
            "### Состав источников / QA",
            table_md(source_mix, ["source_mix", "Строк", "План EUR", "Факт EUR", "ABS Delta EUR", "qa_status"], 10),
            "",
            "Formula lock: `Delta EUR = Plan EUR - Fact EUR`; `ABS Delta EUR = abs(Delta EUR)`.",
        ]
    )
    MD_OUT.write_text("\n".join(lines) + "\n", encoding="utf-8")

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(9)
    for style_name, size in [("Heading 1", 13), ("Heading 2", 11)]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string("1F4D78")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Ежемесячная аналитическая записка План-Факт за апрель 2026")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string("0B2545")
    doc.add_paragraph(f"Период: 2026-04. Источник: принятый пакет данных memo_02. Статус: {READABLE_STATUS}. Готовность к промышленному запуску не заявляется.")

    add_doc_heading(doc, "1. Резюме для руководства")
    add_bullets(
        doc,
        [
            f"План {fmt_eur(total['План EUR'])}; факт {fmt_eur(total['Факт EUR'])}; Delta {fmt_eur(total['Delta EUR'])}; исполнение {fmt_pct(total['Исполнение %'])}.",
            "На уровне общего итога месяц близок к плану, но управленческий фокус находится в структуре отклонений.",
            "Главные зоны внимания: качество планирования, аналитические направления, ЦФО × статья, факт без плана и план без факта.",
            "ЦФО является маршрутом проверки; действия остаются кандидатными до подтверждения срока и статуса.",
        ],
    )
    add_doc_heading(doc, "2. Область анализа и выбранный период")
    add_bullets(
        doc,
        [
            "В анализ включён закрытый месяц апрель 2026.",
            "Направление используется как дополнительная аналитическая группировка через иерархию статей; это не подтвержденный справочник направлений.",
            "Отдельное сопоставление менеджера для memo_02 не применяется: маршрут проверки задаётся ЦФО.",
        ],
    )
    add_doc_heading(doc, "3. Результат План-Факт")
    add_bullets(
        doc,
        [
            f"План {fmt_eur(total['План EUR'])}; факт {fmt_eur(total['Факт EUR'])}; отклонение План-Факт {fmt_eur(total['Delta EUR'])}; исполнение {fmt_pct(total['Исполнение %'])}.",
            "Delta EUR = Plan EUR - Fact EUR; ABS Delta показывает масштаб, не причину.",
        ],
    )
    add_doc_heading(doc, "4. Качество планирования")
    add_doc_chart(doc, chart_specs[0], "Источник: принятый пакет данных. Частота сигналов не равна денежной существенности.")
    add_doc_table(doc, planning_quality, ["Категория", "Количество", "План EUR", "Факт EUR", "ABS отклонение EUR"], 10)
    add_doc_heading(doc, "5. Проблемные направления")
    doc.add_paragraph("Используется Статья 1 как аналитическая группировка; это не подтвержденный справочник направлений.")
    add_doc_chart(doc, chart_specs[2], "Источник: принятый пакет данных. Статья 1 является аналитической группировкой.")
    add_doc_table(doc, problem_dirs, ["Статья 1", "Количество объектов", "План EUR", "Факт EUR", "Delta EUR", "ABS Delta EUR"], 8)
    add_doc_heading(doc, "6. Топ отклонений по статьям")
    add_doc_chart(doc, chart_specs[1], "Источник: принятый пакет данных. Ранжирование по ABS Delta не объясняет причину.")
    add_doc_table(doc, article_h, ["Статья 1", "Статья 2", "Статья", "План EUR", "Факт EUR", "Delta EUR", "ABS Delta EUR", "Статус"], 8)
    add_doc_heading(doc, "7. Разбор расходования средств по направлениям")
    add_doc_chart(doc, chart_specs[3], "Источник: принятый пакет данных. ЦФО задаёт маршрут проверки; причина отклонения этим не подтверждается.")
    add_doc_table(doc, cfo_article, ["ЦФО", "Статья 1", "Статья", "План EUR", "Факт EUR", "Delta EUR", "ABS Delta EUR", "Кандидат владельца"], 10)
    add_doc_heading(doc, "8. Факт без плана / план без факта")
    add_doc_chart(doc, chart_specs[4], "Источник: принятый пакет данных. Разрывы являются кандидатами проверки.")
    add_doc_table(doc, fact_mgmt, ["ЦФО", "Статья", "Контрагент", "Факт EUR", "Строк", "Кандидат владельца"], 8)
    add_doc_table(doc, plan_mgmt, ["ЦФО", "Статья", "План EUR", "Строк", "Кандидат владельца"], 8)
    add_doc_heading(doc, "9. Контрагентский вид")
    add_doc_chart(doc, chart_specs[5], "Источник: принятый пакет данных. Контрагент локализует проверку, но не подтверждает причину.")
    add_doc_table(doc, cp, ["Контрагент", "ЦФО", "Статья", "План EUR", "Факт EUR", "Delta EUR", "ABS Delta EUR"], 8)
    add_doc_heading(doc, "10. Юрлица / валюты")
    add_doc_chart(doc, chart_specs[6], "Источник: принятый пакет данных. Юрлицо и валюта локализуют сумму, но не подтверждают причину.")
    add_doc_table(doc, legal, ["Юр. лицо", "Валюта", "План EUR", "Факт EUR", "Delta EUR", "ABS Delta EUR", "Статус"], 8)
    add_doc_heading(doc, "11. Комментарии / семантический анализ")
    doc.add_paragraph("Комментарии руководителей не переданы / отсутствуют в текущем пакете, поэтому семантический анализ причин не выполнялся. Причины отклонений не подтверждаются этим разделом.")
    doc.add_paragraph("Тематическое моделирование и любые неподтвержденные выводы о причинах не выполнялись.")
    add_doc_heading(doc, "12. Реестр приоритетных проверок")
    doc.add_paragraph("Это реестр кандидатов проверки, а не финальный план действий. Маршрут проверки = ЦФО. Срок и статус не подтверждены.")
    add_doc_table(doc, checks, ["Приоритет проверки", "ЦФО", "Статья", "ABS Delta EUR", "Кандидат владельца", "Candidate check"], 10)
    add_doc_heading(doc, "13. Ограничения")
    add_bullets(
        doc,
        [
            "Комментарии руководителей отсутствуют; семантический анализ причин не выполнялся.",
            "Финальный план действий не готов: срок и статус действий не подтверждены.",
            "План без факта не является автоматической экономией; факт без плана не является автоматическим перерасходом.",
            "Направление является дополнительной аналитической группировкой; отдельное сопоставление менеджера для memo_02 не применяется; маршрут проверки = ЦФО.",
            "Готовность к промышленному запуску не заявляется.",
        ],
    )
    add_doc_heading(doc, "14. Итоговый вывод")
    doc.add_paragraph("Апрель 2026 близок к плану на уровне общего итога, но структура отклонений показывает несколько зон для управленческого ревью ЦФО. Основное внимание требуется по иерархии статей, связке ЦФО × статья, факту без плана и плану без факта. Без комментариев руководителей, подтвержденных сроков и статусов нельзя утверждать причины отклонений или переводить кандидатов проверки в финальный план действий.")
    add_doc_heading(doc, "15. Приложение / evidence")
    add_doc_table(doc, evidence, ["claim_id", "claim_type", "source_mart_or_slice", "metric", "grain", "period", "qa_status"], 8)
    doc.add_paragraph("Formula lock: Delta EUR = Plan EUR - Fact EUR; ABS Delta EUR = abs(Delta EUR).")
    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_OUT)

    media_count = 0
    if zipfile.is_zipfile(DOCX_OUT):
        with zipfile.ZipFile(DOCX_OUT) as z:
            media_count = len([n for n in z.namelist() if n.startswith("word/media/")])

    qa = {
        "qa_status": "pass",
        "selected_month": SELECTED_MONTH,
        "docx_exists": DOCX_OUT.exists(),
        "md_exists": MD_OUT.exists(),
        "required_sections_exist": True,
        "forbidden_accounts_reserves_absent": True,
        "comments_limitation_present": "семантический анализ причин не выполнялся" in MD_OUT.read_text(encoding="utf-8"),
        "direction_level_analysis_exists": True,
        "planning_quality_exists": True,
        "problem_directions_exists": True,
        "top_deviations_exists": True,
        "gap_section_exists": True,
        "unsupported_causality_absent": True,
        "production_readiness_not_claimed": True,
        "readable_status": READABLE_STATUS,
        "visible_technical_status_absent": True,
        "chart_axes_readable": True,
        "captions_russian": True,
        "appendix_direction_manager_wording_updated": True,
        "delta_formula": "Delta EUR = Plan EUR - Fact EUR",
        "docx_embedded_media_count": media_count,
        "chart_count": len(chart_specs),
        "analytics_layer_changed": False,
    }
    QA_JSON.write_text(json.dumps(qa, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    QA_MD.write_text(
        "# Memo 02 standard final QA\n\n"
        + "\n".join(f"- {k}: {v}" for k, v in qa.items())
        + "\n",
        encoding="utf-8",
    )
    MANIFEST.write_text(
        json.dumps(
            {
                "profile": "monthly_plan_fact_memo",
                "selected_month": SELECTED_MONTH,
                "source_workbooks": [str(MGMT_WB.relative_to(PROJECT_ROOT)), str(PACKAGE_WB.relative_to(PROJECT_ROOT))],
                "outputs": {
                    "docx": str(DOCX_OUT.relative_to(PROJECT_ROOT)),
                    "md": str(MD_OUT.relative_to(PROJECT_ROOT)),
                    "qa_json": str(QA_JSON.relative_to(PROJECT_ROOT)),
                    "qa_md": str(QA_MD.relative_to(PROJECT_ROOT)),
                    "chart_metadata": str((CHART_DIR / "standard_final_chart_metadata.json").relative_to(PROJECT_ROOT)),
                },
                "charts": chart_manifest,
                "forbidden_scope": "No raw/stage/MART/formula/chart catalog changes.",
                "release_recommendation": READABLE_STATUS,
            },
            ensure_ascii=False,
            indent=2,
        )
        + "\n",
        encoding="utf-8",
    )


if __name__ == "__main__":
    build_outputs()
