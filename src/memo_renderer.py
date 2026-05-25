from __future__ import annotations

from pathlib import Path

from docx import Document

from src.memo_display_contract import (
    MemoBlock,
    MemoDisplayContract,
    MemoSection,
    MemoTable,
    validate_display_contract,
)


def render_memo_contract_to_docx(contract: MemoDisplayContract, output_path: Path) -> Path:
    errors = validate_display_contract(contract)
    if errors:
        raise ValueError("Invalid MemoDisplayContract: " + "; ".join(errors))

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading(contract.title, level=0)
    if contract.subtitle:
        doc.add_paragraph(contract.subtitle)
    if contract.status_line:
        doc.add_paragraph(contract.status_line)
    doc.add_paragraph(f"Период: {contract.period}")
    doc.add_paragraph(f"Аудитория: {contract.audience}")

    for section in contract.sections:
        _render_section(doc, section)

    if contract.appendix:
        doc.add_heading("Приложение / подтверждения", level=1)
        for block in contract.appendix:
            _render_block(doc, block)

    doc.save(output_path)
    return output_path


def _render_section(doc: Document, section: MemoSection) -> None:
    doc.add_heading(section.title, level=1)
    for block in section.blocks:
        _render_block(doc, block)


def _render_block(doc: Document, block: MemoBlock) -> None:
    if block.block_type == "paragraph":
        doc.add_paragraph(block.text)
    elif block.block_type == "bullet_list":
        for item in block.bullets:
            doc.add_paragraph(item, style="List Bullet")
    elif block.block_type == "kpi_cards":
        for card in block.kpi_cards:
            parts = [card.label, card.value]
            if card.status:
                parts.append(card.status)
            if card.source_ref:
                parts.append(f"источник: {card.source_ref}")
            doc.add_paragraph(" | ".join(parts))
    elif block.block_type == "table":
        _render_table(doc, block.table)
    elif block.block_type == "chart" and block.chart is not None:
        doc.add_paragraph(f"График: {block.chart.title}")
        if block.chart.caption:
            doc.add_paragraph(block.chart.caption)
    elif block.block_type == "limitation_box" and block.limitation is not None:
        doc.add_heading(block.limitation.title, level=2)
        text = block.limitation.text
        if block.limitation.severity:
            text = f"{text} Уровень: {block.limitation.severity}"
        doc.add_paragraph(text)
    elif block.block_type == "action_table":
        rows = [
            [item.action, item.owner, item.status, _render_marker(item.marker), item.due_date, item.evidence_ref]
            for item in block.action_items
        ]
        _render_table(doc, MemoTable(headers=["Действие", "Владелец", "Статус", "Маркер", "Срок", "Подтверждение"], rows=rows))
    elif block.block_type == "evidence_appendix":
        if block.text:
            doc.add_paragraph(block.text)
        for evidence_ref in block.evidence_refs:
            doc.add_paragraph(evidence_ref, style="List Bullet")


def _render_table(doc: Document, memo_table: MemoTable | None) -> None:
    if memo_table is None:
        return
    table = doc.add_table(rows=1, cols=len(memo_table.headers))
    table.style = "Table Grid"
    for index, header in enumerate(memo_table.headers):
        table.rows[0].cells[index].text = header
    for row in memo_table.rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    if memo_table.caption:
        doc.add_paragraph(memo_table.caption)


def _render_marker(marker: str) -> str:
    return {"candidate": "кандидат", "final": "финал"}.get(marker, marker)
