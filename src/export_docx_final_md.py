from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


def escape_cell(text: str) -> str:
    return " ".join(text.split()).replace("|", "\\|")


def paragraph_to_markdown(text: str, style_name: str) -> str:
    text = text.strip()
    if not text:
        return ""
    lowered = style_name.lower()
    if "heading 1" in lowered or lowered == "title":
        return f"# {text}"
    if "heading 2" in lowered:
        return f"## {text}"
    if "heading 3" in lowered:
        return f"### {text}"
    if "list bullet" in lowered:
        return f"- {text}"
    if "list number" in lowered:
        return f"1. {text}"
    return text


def table_to_markdown(table) -> list[str]:
    rows = [[escape_cell(cell.text) for cell in row.cells] for row in table.rows]
    rows = [row for row in rows if any(cell for cell in row)]
    if not rows:
        return []
    width = max(len(row) for row in rows)
    rows = [row + [""] * (width - len(row)) for row in rows]
    header = rows[0]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(["---"] * width) + " |",
    ]
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return lines


def export_docx_to_markdown(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    lines: list[str] = []
    table_iter = iter(doc.tables)
    for block in doc.element.body:
        tag = block.tag.rsplit("}", 1)[-1]
        if tag == "p":
            paragraph = next(p for p in doc.paragraphs if p._p is block)
            line = paragraph_to_markdown(paragraph.text, paragraph.style.name if paragraph.style else "")
            if line:
                lines.append(line)
                lines.append("")
        elif tag == "tbl":
            table = next(table_iter)
            table_lines = table_to_markdown(table)
            if table_lines:
                lines.extend(table_lines)
                lines.append("")
    return "\n".join(lines).strip() + "\n"


def main() -> None:
    parser = argparse.ArgumentParser(description="Export accepted DOCX final narrative to canonical Markdown.")
    parser.add_argument("docx", type=Path)
    parser.add_argument("output_md", type=Path)
    args = parser.parse_args()
    markdown = export_docx_to_markdown(args.docx)
    args.output_md.parent.mkdir(parents=True, exist_ok=True)
    args.output_md.write_text(markdown, encoding="utf-8")
    print(f"wrote {args.output_md}")


if __name__ == "__main__":
    main()
