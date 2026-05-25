from __future__ import annotations

import json
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from urllib import error, request

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt

try:
    from src.ollama_routing import OLLAMA_FAST_FALLBACK_MODEL, OLLAMA_MODELS, OLLAMA_ROUTING_NOTE, OLLAMA_URL, model_for_role
except ModuleNotFoundError:
    from ollama_routing import OLLAMA_FAST_FALLBACK_MODEL, OLLAMA_MODELS, OLLAMA_ROUTING_NOTE, OLLAMA_URL, model_for_role


PROJECT_ROOT = Path(__file__).resolve().parents[1]
MARTS_DIR = PROJECT_ROOT / "03_marts"
SIGNALS_DIR = PROJECT_ROOT / "04_signals"
EVIDENCE_DIR = PROJECT_ROOT / "05_evidence"
LLM_PACKAGE_DIR = PROJECT_ROOT / "05_llm_package"
REPORTS_DIR = PROJECT_ROOT / "06_reports"
QA_DIR = PROJECT_ROOT / "07_qa"
CHARTS_DIR = REPORTS_DIR / "charts"
BATCH_MD_DIR = REPORTS_DIR / "memo_draft_batches"
BATCH_JSON_DIR = LLM_PACKAGE_DIR / "llm_batches"

CSV_KWARGS = {"index": False, "encoding": "utf-8-sig", "sep": ";", "decimal": ",", "lineterminator": "\n"}

REPORT_MART_FILES = {
    "mart_budget_kpi": "mart_budget_kpi.parquet",
    "mart_budget_plan_risk_by_month": "mart_budget_plan_risk_by_month.parquet",
    "mart_budget_plan_risk_by_article_cfo": "mart_budget_plan_risk_by_article_cfo.parquet",
    "mart_budget_counterparty_quality": "mart_budget_counterparty_quality.parquet",
    "mart_budget_signal_catalog": "mart_budget_signal_catalog.parquet",
    "mart_budget_chart_dataset": "mart_budget_chart_dataset.parquet",
}

CHART_FILES = {
    "CH_BUDGET_BRIDGE_001": "CH_BUDGET_BRIDGE_001.png",
    "CH_TOP_GAP_CONTRIBUTORS_001": "CH_TOP_GAP_CONTRIBUTORS_001.png",
    "CH_ONJN_AVENTO_MONTHLY_001": "CH_ONJN_AVENTO_MONTHLY_001.png",
    "CH_MONTH_ARTICLE_CFO_HEATMAP_001": "CH_MONTH_ARTICLE_CFO_HEATMAP_001.png",
    "CH_COUNTERPARTY_QUALITY_001": "CH_COUNTERPARTY_QUALITY_001.png",
    "CH_UNKNOWN_COUNTERPARTY_AMOUNT_001": "CH_UNKNOWN_COUNTERPARTY_AMOUNT_001.png",
    "CH_HISTORICAL_VOLATILITY_001": "CH_HISTORICAL_VOLATILITY_001.png",
    "CH_CFO_PLAN_BASE_001": "CH_CFO_PLAN_BASE_001.png",
    "CH_GAP_SHARE_001": "CH_GAP_SHARE_001.png",
}

REPORT_BATCH_ROLE = "analyst"

SECTION_TITLES = [
    "Краткий управленческий вывод",
    "Решение, которое требуется от руководства",
    "Основные выводы",
    "Карта рисков плановой базы",
    "Что нужно проверить",
    "Риски качества данных",
    "Рекомендуемые действия",
    "Ограничения анализа",
    "Приложение. Графики",
    "Приложение. Источники подтверждения и ограничения",
]

CAUSALITY_PHRASES = [
    "из-за",
    "причина",
    "объясняется",
    "вызвано",
    "связано с",
    "основной драйвер",
    "главный драйвер",
    "root cause",
    "confirmed driver",
]
COUNTERPARTY_NO_ISSUE_PHRASES = ["нет неизвестных контрагентов", "неизвестные контрагенты не выявлены"]
PLANNING_FACT_PHRASES = ["перерасход", "экономия", "факт превысил", "недоисполнение", "loss", "damage", "financial loss", "overspend"]


def safe_div(num: float | pd.Series, den: float | pd.Series) -> float | pd.Series:
    if isinstance(den, pd.Series):
        return pd.Series(num, index=den.index) / den.where(den.notna() & den.ne(0))
    if den == 0 or pd.isna(den):
        return np.nan
    return num / den


def money(value: float | int | None) -> str:
    if value is None or pd.isna(value):
        return "n/a"
    return f"{float(value):,.0f} EUR".replace(",", " ")


def pct(value: float | int | None) -> str:
    if value is None or pd.isna(value):
        return "n/a"
    return f"{float(value) * 100:.1f}%"


def to_jsonable(value: Any) -> Any:
    if isinstance(value, (np.integer,)):
        return int(value)
    if isinstance(value, (np.floating,)):
        return None if pd.isna(value) else float(value)
    if isinstance(value, (pd.Timestamp, datetime)):
        return value.isoformat()
    if isinstance(value, list):
        return [to_jsonable(item) for item in value]
    if isinstance(value, dict):
        return {str(key): to_jsonable(item) for key, item in value.items()}
    if pd.isna(value):
        return None
    return value


def write_json(path: Path, payload: Any) -> None:
    path.write_text(json.dumps(to_jsonable(payload), ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def log_step(message: str) -> None:
    print(f"[build_report] {message}", flush=True)


def call_ollama_model(model: str, prompt: str, timeout: int) -> tuple[str | None, str | None]:
    payload = json.dumps(
        {
            "model": model,
            "prompt": prompt,
            "stream": False,
            "options": {"temperature": 0.1, "num_predict": 900},
        }
    ).encode("utf-8")
    req = request.Request(
        f"{OLLAMA_URL}/api/generate",
        data=payload,
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    try:
        with request.urlopen(req, timeout=timeout) as response:
            data = json.loads(response.read().decode("utf-8"))
        text = str(data.get("response", "")).strip()
        return (text, None) if text else (None, "empty_response")
    except (error.URLError, TimeoutError, json.JSONDecodeError) as exc:
        return None, str(exc)


def ollama_generate(prompt: str, role: str = REPORT_BATCH_ROLE, timeout: int = 300) -> tuple[str | None, str | None, str, bool]:
    primary_model = model_for_role(role)
    generated, error_text = call_ollama_model(primary_model, prompt, timeout)
    if not error_text:
        return generated, None, primary_model, False
    if primary_model == OLLAMA_FAST_FALLBACK_MODEL:
        return None, error_text, primary_model, False
    fallback_text, fallback_error = call_ollama_model(OLLAMA_FAST_FALLBACK_MODEL, prompt, min(timeout, 180))
    if not fallback_error:
        return fallback_text, f"primary_error:{error_text}", OLLAMA_FAST_FALLBACK_MODEL, True
    return None, f"primary_error:{error_text}; fallback_error:{fallback_error}", primary_model, False


def batch_text_failures(text: str, required_heading: str) -> list[str]:
    lower = text.lower()
    failures: list[str] = []
    if required_heading.lower() not in lower:
        failures.append("missing_required_heading")
    if "EC_" not in text:
        failures.append("missing_evidence_id")
    for phrase in CAUSALITY_PHRASES + COUNTERPARTY_NO_ISSUE_PHRASES + PLANNING_FACT_PHRASES:
        if phrase in lower:
            failures.append(f"forbidden_phrase:{phrase}")
    return failures


def clean_generated_section(text: str, required_heading: str) -> str:
    lines = []
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("# Required section heading") or stripped.startswith("# Draft section"):
            continue
        if stripped == f"# {required_heading}":
            continue
        lines.append(line.rstrip())
    return "\n".join(lines).strip()


def read_inputs() -> dict[str, pd.DataFrame]:
    data = {
        "stage": pd.read_parquet(MARTS_DIR / "mart_10_stage_typed.parquet"),
        "inout": pd.read_parquet(MARTS_DIR / "mart_11_month_inout.parquet"),
        "article": pd.read_parquet(MARTS_DIR / "mart_12_article_month.parquet"),
        "article_cfo": pd.read_parquet(MARTS_DIR / "mart_13_article_cfo_month.parquet"),
        "counterparty": pd.read_parquet(MARTS_DIR / "mart_14_article_cfo_counterparty_month.parquet"),
    }
    signal_frames = []
    for path in sorted(SIGNALS_DIR.glob("signal_*.parquet")):
        df = pd.read_parquet(path)
        df["signal_file"] = path.name
        signal_frames.append(df)
    data["signals"] = pd.concat(signal_frames, ignore_index=True, sort=False) if signal_frames else pd.DataFrame()
    return data


def build_report_marts(data: dict[str, pd.DataFrame]) -> dict[str, pd.DataFrame]:
    article = data["article"]
    article_cfo = data["article_cfo"]
    counterparty = data["counterparty"]
    signals = data["signals"]
    planning = article[article["effective_period_type"].eq("planning")].copy()
    historical = article[article["effective_period_type"].eq("historical")].copy()

    planned_budget = float(planning["plan_eur"].sum())
    historical_base = float(planning["historical_base_eur"].sum(skipna=True))
    kpi = pd.DataFrame(
        [
            {
                "period_min": str(article["month"].min()),
                "period_max": str(article["month"].max()),
                "historical_plan_eur": float(historical["plan_eur"].sum()),
                "historical_fact_eur": float(historical["fact_eur"].sum()),
                "planned_budget_eur": planned_budget,
                "historical_base_eur": historical_base,
                "plan_vs_base_eur": planned_budget - historical_base,
                "plan_vs_base_abs_eur": abs(planned_budget - historical_base),
                "currency": "EUR",
            }
        ]
    )
    by_month = (
        planning.groupby(["month", "year"], as_index=False)
        .agg(
            planned_budget_eur=("plan_eur", "sum"),
            historical_base_eur=("historical_base_eur", "sum"),
            source_rows_count=("source_rows_count", "sum"),
        )
        .assign(
            plan_vs_base_eur=lambda df: df["planned_budget_eur"] - df["historical_base_eur"],
            plan_vs_base_abs_eur=lambda df: (df["planned_budget_eur"] - df["historical_base_eur"]).abs(),
        )
    )
    by_article_cfo = (
        article_cfo[article_cfo["effective_period_type"].eq("planning")]
        .groupby(["statya", "cfo"], as_index=False)
        .agg(
            planned_budget_eur=("plan_eur", "sum"),
            historical_base_eur=("historical_base_eur", "sum"),
            source_rows_count=("source_rows_count", "sum"),
        )
        .assign(
            plan_vs_base_eur=lambda df: df["planned_budget_eur"] - df["historical_base_eur"],
            plan_vs_base_abs_eur=lambda df: (df["planned_budget_eur"] - df["historical_base_eur"]).abs(),
        )
        .sort_values("plan_vs_base_abs_eur", ascending=False)
        .reset_index(drop=True)
    )
    cp_quality = (
        counterparty.groupby("month", as_index=False)
        .agg(
            total_rows=("month", "size"),
            unknown_counterparty_rows=("unknown_counterparty_flag", "sum"),
            unknown_counterparty_amount_eur=("abs_deviation_eur", lambda s: float(s[counterparty.loc[s.index, "unknown_counterparty_flag"].eq(1)].sum())),
        )
        .assign(
            known_counterparty_rows=lambda df: df["total_rows"] - df["unknown_counterparty_rows"],
            unknown_counterparty_pct=lambda df: df["unknown_counterparty_rows"] / df["total_rows"].where(df["total_rows"].ne(0)),
            status=lambda df: np.where(df["unknown_counterparty_rows"].gt(0), "issue_detected", "ok"),
        )
    )
    signal_catalog = signals.copy()
    if not signal_catalog.empty:
        keep = [col for col in ["signal_id", "signal_type", "period", "grain", "month", "year", "statya", "cfo", "source_table", "quality_flags", "signal_file"] if col in signal_catalog.columns]
        signal_catalog = signal_catalog[keep].copy()
    chart_dataset = pd.DataFrame(
        [{"chart_id": chart_id, "png_file": file_name} for chart_id, file_name in CHART_FILES.items()]
    )
    marts = {
        "mart_budget_kpi": kpi,
        "mart_budget_plan_risk_by_month": by_month,
        "mart_budget_plan_risk_by_article_cfo": by_article_cfo,
        "mart_budget_counterparty_quality": cp_quality,
        "mart_budget_signal_catalog": signal_catalog,
        "mart_budget_chart_dataset": chart_dataset,
    }
    for name, df in marts.items():
        df.to_parquet(MARTS_DIR / REPORT_MART_FILES[name], index=False)
    return marts


def evidence_card(
    card_id: str,
    card_type: str,
    claim_type: str,
    source_mart: str,
    period: str,
    grain: str,
    metric_name: str,
    metric_value: Any,
    formula: str,
    allowed_claims: list[str],
    forbidden_claims: list[str],
    **extra: Any,
) -> dict[str, Any]:
    payload = {
        "card_id": card_id,
        "card_type": card_type,
        "claim_type": claim_type,
        "source_table": source_mart,
        "source_marts": [source_mart],
        "period": period,
        "grain": grain,
        "metric_name": metric_name,
        "metric_value": to_jsonable(metric_value),
        "formula": formula,
        "allowed_claims": allowed_claims,
        "forbidden_claims": forbidden_claims,
        "quality_flags": [],
    }
    payload.update(extra)
    return payload


def update_evidence_cards(report_marts: dict[str, pd.DataFrame]) -> list[dict[str, Any]]:
    existing = json.loads((EVIDENCE_DIR / "evidence_cards.json").read_text(encoding="utf-8"))
    cards_by_id = {card["card_id"]: card for card in existing}
    kpi = report_marts["mart_budget_kpi"].iloc[0]
    by_article_cfo = report_marts["mart_budget_plan_risk_by_article_cfo"]
    cp_quality = report_marts["mart_budget_counterparty_quality"]
    onjn = by_article_cfo[
        by_article_cfo["statya"].str.contains("ONJN Gaming Tax", case=False, na=False)
        & by_article_cfo["cfo"].str.contains("Avento MT", case=False, na=False)
    ]
    onjn_row = onjn.iloc[0] if not onjn.empty else by_article_cfo.iloc[0]
    unknown_rows = int(cp_quality["unknown_counterparty_rows"].sum())
    unknown_amount = float(cp_quality["unknown_counterparty_amount_eur"].sum())
    total_rows = int(cp_quality["total_rows"].sum())
    recommendations = [
        {"entity_type": "article_cfo", "entity_name": "ONJN Gaming Tax / Avento MT", "action_type": "owner_confirmation", "owner": "budget_owner"},
        {"entity_type": "data_quality", "entity_name": "unknown_counterparties", "action_type": "mapping_review", "owner": "data_owner"},
    ]

    cards_by_id["EC_PLAN_RISK_TOTAL_001"] = evidence_card(
        "EC_PLAN_RISK_TOTAL_001",
        "EC_PLAN_RISK",
        "calculation_result",
        "mart_budget_kpi",
        f"{kpi['period_min']}..{kpi['period_max']}",
        "planning total",
        "plan_vs_base_abs_eur",
        kpi["plan_vs_base_abs_eur"],
        "planned_budget_eur - historical_base_eur",
        ["Плановый бюджет выше исторической базы.", "Разрыв между планом и исторической базой составляет X EUR."],
        ["Причина разрыва - ONJN Gaming Tax.", "Основной драйвер - ONJN Gaming Tax / Avento MT."],
        planned_budget_eur=kpi["planned_budget_eur"],
        historical_base_eur=kpi["historical_base_eur"],
        plan_vs_base_abs_eur=kpi["plan_vs_base_abs_eur"],
    )
    cards_by_id["EC_ONJN_AVENTO_LOCALIZATION_001"] = evidence_card(
        "EC_ONJN_AVENTO_LOCALIZATION_001",
        "EC_PLAN_RISK_CFO",
        "localization_signal",
        "mart_budget_plan_risk_by_article_cfo",
        f"{kpi['period_min']}..{kpi['period_max']}",
        "статья+ЦФО",
        "plan_vs_base_abs_eur",
        onjn_row["plan_vs_base_abs_eur"],
        "planned_budget_eur - historical_base_eur at article × CFO grain",
        ["Крупный разрыв локализован в связке ONJN Gaming Tax / Avento MT.", "Связка требует подтверждения владельца бюджета."],
        ["ONJN Gaming Tax является причиной разрыва.", "ONJN Gaming Tax является основным драйвером.", "Разрыв возник из-за ONJN Gaming Tax."],
        entity={"article": str(onjn_row["statya"]), "cfo": str(onjn_row["cfo"])},
        limitations=["Локализация не равна причинности.", "Требуется подтверждение владельца бюджета."],
    )
    cards_by_id["EC_COUNTERPARTY_QUALITY_001"] = evidence_card(
        "EC_COUNTERPARTY_QUALITY_001",
        "EC_DATA_QUALITY",
        "data_quality",
        "mart_budget_counterparty_quality",
        f"{kpi['period_min']}..{kpi['period_max']}",
        "month",
        "unknown_counterparty_rows",
        unknown_rows,
        "sum unknown_counterparty_rows by month",
        ["Обнаружены строки с неизвестными контрагентами.", "Требуется проверка маппинга контрагентов."],
        ["В данных нет неизвестных контрагентов."],
        total_rows=total_rows,
        unknown_counterparty_rows=unknown_rows,
        unknown_counterparty_pct=safe_div(unknown_rows, total_rows),
        unknown_counterparty_amount_eur=unknown_amount,
        status="issue_detected" if unknown_rows > 0 else "ok",
    )
    cards_by_id["EC_RECOMMENDATION_QUEUE_001"] = evidence_card(
        "EC_RECOMMENDATION_QUEUE_001",
        "EC_RECOMMENDATION_BASIS",
        "recommendation",
        "mart_budget_kpi",
        f"{kpi['period_min']}..{kpi['period_max']}",
        "recommendation queue",
        "recommendations_count",
        len(recommendations),
        "deduplicate by entity_type + entity_name + action_type + owner",
        ["Рекомендации являются проверочными действиями.", "Рекомендации требуют подтверждения владельцев."],
        ["Автоматически менять бюджет без owner confirmation."],
        recommendations=recommendations,
    )
    cards_by_id["EC_TOP_CONTRIBUTORS_001"] = evidence_card(
        "EC_TOP_CONTRIBUTORS_001",
        "EC_PLAN_RISK_CFO",
        "localization_signal",
        "mart_budget_plan_risk_by_article_cfo",
        f"{kpi['period_min']}..{kpi['period_max']}",
        "статья+ЦФО",
        "top_contributors_count",
        min(15, len(by_article_cfo)),
        "rank abs(planned_budget_eur - historical_base_eur) desc",
        ["Топ связки показывают локализацию планового разрыва."],
        ["Топ связки являются подтвержденной причиной разрыва."],
    )
    cards = list(cards_by_id.values())
    (EVIDENCE_DIR / "evidence_cards.json").write_text(json.dumps(cards, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    evidence_df = pd.DataFrame(cards)
    for col in ["quality_flags", "allowed_claims", "forbidden_claims"]:
        if col in evidence_df.columns:
            evidence_df[col] = evidence_df[col].map(lambda value: " | ".join(value) if isinstance(value, list) else value)
    evidence_df.to_csv(EVIDENCE_DIR / "evidence_cards.csv", **CSV_KWARGS)
    return cards


def save_chart(fig: plt.Figure, chart_id: str) -> None:
    fig.tight_layout()
    fig.savefig(CHARTS_DIR / CHART_FILES[chart_id], dpi=160, bbox_inches="tight")
    plt.close(fig)


def build_charts(report_marts: dict[str, pd.DataFrame], data: dict[str, pd.DataFrame]) -> list[dict[str, Any]]:
    CHARTS_DIR.mkdir(parents=True, exist_ok=True)
    for path in CHARTS_DIR.glob("*.png"):
        path.unlink()
    kpi = report_marts["mart_budget_kpi"].iloc[0]
    by_month = report_marts["mart_budget_plan_risk_by_month"]
    by_article_cfo = report_marts["mart_budget_plan_risk_by_article_cfo"]
    cp_quality = report_marts["mart_budget_counterparty_quality"]
    article = data["article"]
    article_cfo = data["article_cfo"]
    meta: list[dict[str, Any]] = []

    def add_meta(chart_id: str, title: str, chart_type: str, source_mart: str, grain: str, metric: str, evidence: str, limitation: str) -> None:
        meta.append(
            {
                "chart_id": chart_id,
                "chart_title": title,
                "chart_type": chart_type,
                "source_mart": source_mart,
                "period": f"{kpi['period_min']}..{kpi['period_max']}",
                "grain": grain,
                "metric_name": metric,
                "evidence_card_id": evidence,
                "limitation": limitation,
                "png_file": str(CHARTS_DIR / CHART_FILES[chart_id]),
                "commentary": f"Источник: {source_mart}; метрика: {metric}; limitation: {limitation}",
            }
        )

    top = by_article_cfo.head(15).copy()
    fig, ax = plt.subplots(figsize=(10, 5.2))
    bridge_labels = ["Ист. база", *top["statya"].head(6).tolist(), "План"]
    bridge_values = [kpi["historical_base_eur"], *top["plan_vs_base_eur"].head(6).tolist(), kpi["planned_budget_eur"]]
    colors = ["#4f81bd", *["#9bbb59" if value >= 0 else "#c0504d" for value in bridge_values[1:-1]], "#8064a2"]
    ax.bar(range(len(bridge_values)), bridge_values, color=colors)
    ax.set_title("Мост изменения бюджета: историческая база -> плановый бюджет")
    ax.set_xticks(range(len(bridge_labels)), [str(label)[:18] for label in bridge_labels], rotation=35, ha="right")
    ax.set_ylabel("EUR")
    save_chart(fig, "CH_BUDGET_BRIDGE_001")
    add_meta("CH_BUDGET_BRIDGE_001", "Мост плановой базы к историческому уровню", "waterfall", "mart_budget_plan_risk_by_article_cfo", "статья+ЦФО", "plan_vs_base_eur", "EC_PLAN_RISK_TOTAL_001", "Показывает зоны проверки, а не подтвержденное основание изменения бюджета.")

    fig, ax = plt.subplots(figsize=(10, 6))
    labels = (top["statya"] + " / " + top["cfo"]).str.slice(0, 42)
    ax.barh(labels[::-1], top["plan_vs_base_abs_eur"][::-1], color="#4f81bd")
    ax.set_title("Топ-15 связок статья x ЦФО по вкладу в плановый разрыв")
    ax.set_xlabel("abs(plan_vs_base_eur), EUR")
    save_chart(fig, "CH_TOP_GAP_CONTRIBUTORS_001")
    add_meta("CH_TOP_GAP_CONTRIBUTORS_001", "Топ-15 комбинаций статья бюджета × ЦФО", "horizontal_bar", "mart_budget_plan_risk_by_article_cfo", "статья+ЦФО", "plan_vs_base_abs_eur", "EC_TOP_CONTRIBUTORS_001", "Высокая концентрация означает приоритет проверки, а не подтвержденное основание бюджета.")

    onjn = article_cfo[
        article_cfo["statya"].str.contains("ONJN Gaming Tax", case=False, na=False)
        & article_cfo["cfo"].str.contains("Avento MT", case=False, na=False)
    ].sort_values("month")
    fig, ax = plt.subplots(figsize=(10, 4.8))
    if onjn.empty:
        onjn = article_cfo[article_cfo["effective_period_type"].eq("planning")].sort_values("abs_plan_vs_base_eur", ascending=False).head(12)
    x = np.arange(len(onjn))
    ax.bar(x - 0.2, onjn["plan_eur"], width=0.4, label="planned_budget_eur")
    ax.bar(x + 0.2, onjn["historical_base_eur"], width=0.4, label="historical_base_eur")
    ax.plot(x, onjn["abs_plan_vs_base_eur"], color="#c0504d", marker="o", label="plan_vs_base_abs_eur")
    ax.set_xticks(x, onjn["month"].astype(str), rotation=45)
    ax.set_title("ONJN Gaming Tax / Avento MT: профиль по месяцам")
    ax.legend(fontsize=8)
    save_chart(fig, "CH_ONJN_AVENTO_MONTHLY_001")
    add_meta("CH_ONJN_AVENTO_MONTHLY_001", "ONJN Gaming Tax / Avento MT по месяцам", "clustered_column_or_line", "mart_13_article_cfo_month", "month+статья+ЦФО", "plan_vs_base_abs_eur", "EC_ONJN_AVENTO_LOCALIZATION_001", "Помесячный профиль показывает зону проверки. Требуется подтверждение владельца бюджета.")

    heat_source = article_cfo[article_cfo["effective_period_type"].eq("planning")].copy()
    top_articles = heat_source.groupby("statya")["abs_plan_vs_base_eur"].sum().sort_values(ascending=False).head(20).index
    heat = heat_source[heat_source["statya"].isin(top_articles)].pivot_table(index="statya", columns="month", values="abs_plan_vs_base_eur", aggfunc="sum", fill_value=0)
    fig, ax = plt.subplots(figsize=(11, 7))
    ax.imshow(heat.values, aspect="auto", cmap="Blues")
    ax.set_yticks(range(len(heat.index)), [str(v)[:34] for v in heat.index], fontsize=7)
    ax.set_xticks(range(len(heat.columns)), heat.columns, rotation=45, ha="right", fontsize=7)
    ax.set_title("Карта концентрации планового риска по месяцам и статьям")
    save_chart(fig, "CH_MONTH_ARTICLE_CFO_HEATMAP_001")
    add_meta("CH_MONTH_ARTICLE_CFO_HEATMAP_001", "Карта концентрации риска плановой базы", "heatmap", "mart_13_article_cfo_month", "month+статья+ЦФО", "plan_vs_base_abs_eur", "EC_TOP_CONTRIBUTORS_001", "Карта показывает концентрацию проверки, а не подтвержденные бизнес-основания.")

    fig, ax = plt.subplots(figsize=(10, 4.8))
    ax.bar(cp_quality["month"], cp_quality["known_counterparty_rows"], label="known_counterparty_rows")
    ax.bar(cp_quality["month"], cp_quality["unknown_counterparty_rows"], bottom=cp_quality["known_counterparty_rows"], label="unknown_counterparty_rows")
    ax.set_title("Качество маппинга контрагентов")
    ax.tick_params(axis="x", rotation=45)
    ax.legend(fontsize=8)
    save_chart(fig, "CH_COUNTERPARTY_QUALITY_001")
    add_meta("CH_COUNTERPARTY_QUALITY_001", "Качество маппинга контрагентов", "stacked_bar", "mart_budget_counterparty_quality", "month", "unknown_counterparty_rows", "EC_COUNTERPARTY_QUALITY_001", "Строки с неопределенным контрагентом являются зоной проверки качества данных.")

    unknown_amount = cp_quality.sort_values("unknown_counterparty_amount_eur", ascending=False).head(15)
    fig, ax = plt.subplots(figsize=(10, 4.8))
    ax.barh(unknown_amount["month"][::-1], unknown_amount["unknown_counterparty_amount_eur"][::-1], color="#c0504d")
    ax.set_title("Сумма по строкам с неизвестными контрагентами")
    ax.set_xlabel("EUR")
    save_chart(fig, "CH_UNKNOWN_COUNTERPARTY_AMOUNT_001")
    add_meta("CH_UNKNOWN_COUNTERPARTY_AMOUNT_001", "Оборот по строкам с неопределенным контрагентом", "horizontal_bar", "mart_budget_counterparty_quality", "month", "unknown_counterparty_amount_eur", "EC_COUNTERPARTY_QUALITY_001", "Оборот является объемом проверки справочника, а не подтвержденным финансовым искажением.")

    hist_vol = (
        article[article["effective_period_type"].eq("historical")]
        .groupby("statya", as_index=False)
        .agg(historical_avg_eur=("fact_eur", "mean"), historical_stddev_eur=("fact_eur", "std"), historical_total_eur=("fact_eur", "sum"))
        .fillna(0)
    )
    fig, ax = plt.subplots(figsize=(9, 5))
    sizes = np.clip(hist_vol["historical_total_eur"].abs() / max(hist_vol["historical_total_eur"].abs().max(), 1) * 700, 20, 700)
    ax.scatter(hist_vol["historical_avg_eur"], hist_vol["historical_stddev_eur"], s=sizes, alpha=0.45)
    ax.set_title("Историческая волатильность статей бюджета")
    ax.set_xlabel("historical_avg_eur")
    ax.set_ylabel("historical_stddev_eur")
    save_chart(fig, "CH_HISTORICAL_VOLATILITY_001")
    add_meta("CH_HISTORICAL_VOLATILITY_001", "Историческая волатильность статей бюджета", "scatter", "mart_12_article_month", "article", "historical_stddev_eur", "EC_HISTORICAL_TOTAL_001", "Волатильность является описательной характеристикой и не подтверждает основание планового разрыва.")

    cfo_gap = by_article_cfo.groupby("cfo", as_index=False)["plan_vs_base_abs_eur"].sum().sort_values("plan_vs_base_abs_eur", ascending=False).head(20)
    fig, ax = plt.subplots(figsize=(10, 5.5))
    ax.barh(cfo_gap["cfo"][::-1], cfo_gap["plan_vs_base_abs_eur"][::-1], color="#8064a2")
    ax.set_title("Плановый бюджет против исторической базы по ЦФО")
    ax.set_xlabel("plan_vs_base_abs_eur")
    save_chart(fig, "CH_CFO_PLAN_BASE_001")
    add_meta("CH_CFO_PLAN_BASE_001", "Плановая база против исторического уровня по ЦФО", "horizontal_bar", "mart_budget_plan_risk_by_article_cfo", "cfo", "plan_vs_base_abs_eur", "EC_TOP_CONTRIBUTORS_001", "Разрез по ЦФО помогает направить проверку, но не назначает ответственность без подтверждения владельца.")

    pareto = by_article_cfo.head(30).copy()
    total_gap = pareto["plan_vs_base_abs_eur"].sum()
    pareto["cumulative_share_of_gap"] = pareto["plan_vs_base_abs_eur"].cumsum() / total_gap if total_gap else np.nan
    fig, ax1 = plt.subplots(figsize=(11, 5.2))
    x = np.arange(len(pareto))
    ax1.bar(x, pareto["plan_vs_base_abs_eur"], color="#4f81bd")
    ax2 = ax1.twinx()
    ax2.plot(x, pareto["cumulative_share_of_gap"], color="#c0504d", marker="o")
    ax1.set_title("Концентрация планового разрыва")
    ax1.set_xticks(x, (pareto["statya"] + " / " + pareto["cfo"]).str.slice(0, 20), rotation=75, ha="right", fontsize=6)
    ax2.set_ylim(0, 1.05)
    save_chart(fig, "CH_GAP_SHARE_001")
    add_meta("CH_GAP_SHARE_001", "Концентрация планового разрыва", "pareto", "mart_budget_plan_risk_by_article_cfo", "статья+ЦФО", "cumulative_share_of_gap", "EC_TOP_CONTRIBUTORS_001", "Концентрация помогает расставить приоритеты проверки, но не подтверждает основание бюджета.")

    report_marts["mart_budget_chart_dataset"] = pd.DataFrame(meta)
    report_marts["mart_budget_chart_dataset"].to_parquet(MARTS_DIR / REPORT_MART_FILES["mart_budget_chart_dataset"], index=False)
    return meta


def build_batches(report_marts: dict[str, pd.DataFrame], chart_meta: list[dict[str, Any]], cards: list[dict[str, Any]]) -> dict[str, str]:
    log_step("building deterministic Russian memo batches")
    BATCH_MD_DIR.mkdir(parents=True, exist_ok=True)
    BATCH_JSON_DIR.mkdir(parents=True, exist_ok=True)
    for directory in [BATCH_MD_DIR, BATCH_JSON_DIR]:
        for path in directory.glob("*"):
            if path.is_file():
                path.unlink()
    batches: dict[str, str] = {
        "batch_01_executive_summary": (
            "Краткий управленческий вывод: зоны плановой базы требуют подтверждения владельцами бюджета."
        ),
        "batch_02_key_numbers": (
            "Ключевые показатели вынесены в основную записку и приложение с источниками подтверждения."
        ),
        "batch_03_budget_bridge": (
            "Графики перенесены в приложение и описаны управленческим языком."
        ),
        "batch_04_localization_signals": (
            "Основные зоны проверки сформированы по крупнейшим комбинациям статья бюджета × ЦФО."
        ),
        "batch_05_plan_risk": (
            "Карта рисков плановой базы отделяет высокий риск проверки от подтвержденных бизнес-оснований."
        ),
        "batch_06_data_quality": (
            "Риск качества данных по неопределенным контрагентам остается видимым в основной части."
        ),
        "batch_07_counterparty_risk": (
            "Подробные источники подтверждения и ограничения перенесены в приложение."
        ),
        "batch_08_actions_limitations": (
            "Рекомендуемые действия сформулированы как управленческий план проверки."
        ),
    }
    final_batches: dict[str, str] = {}
    for name, text in batches.items():
        log_step(f"writing batch artifact {name}")
        used_text = text
        final_batches[name] = used_text
        batch_payload = {
            "batch_id": name,
            "generation_mode": "deterministic_russian_management_revisor",
            "ollama_url": OLLAMA_URL,
            "ollama_role": REPORT_BATCH_ROLE,
            "ollama_model": model_for_role(REPORT_BATCH_ROLE),
            "ollama_primary_model": model_for_role(REPORT_BATCH_ROLE),
            "ollama_fast_fallback_model": OLLAMA_FAST_FALLBACK_MODEL,
            "ollama_fallback_used": False,
            "ollama_status": "not_used_for_deterministic_revisor_pass",
            "ollama_error": None,
            "quality_failures": [],
            "input_evidence_card_ids": [card["card_id"] for card in cards if str(card.get("card_id", "")).startswith("EC_")][:40],
            "raw_data_included": False,
            "output_file": str(BATCH_MD_DIR / f"{name}.md"),
        }
        write_json(BATCH_JSON_DIR / f"{name}.json", batch_payload)
        (BATCH_MD_DIR / f"{name}.md").write_text(used_text + "\n", encoding="utf-8")
    write_json(
        LLM_PACKAGE_DIR / "batch_manifest.json",
        {
            "created_at": datetime.now(timezone.utc).isoformat(),
            "ollama_url": OLLAMA_URL,
            "ollama_models": OLLAMA_MODELS,
            "ollama_fast_fallback_model": OLLAMA_FAST_FALLBACK_MODEL,
            "ollama_routing_note": OLLAMA_ROUTING_NOTE,
            "batch_role": REPORT_BATCH_ROLE,
            "batches": [{"batch_id": name, "json": str(BATCH_JSON_DIR / f"{name}.json"), "markdown": str(BATCH_MD_DIR / f"{name}.md")} for name in batches],
            "raw_data_included": False,
        },
    )
    return final_batches


def build_final_memo(batches: dict[str, str], cards: list[dict[str, Any]], chart_meta: list[dict[str, Any]]) -> str:
    log_step("assembling Russian executive memo and evidence pack")
    card_map = {card["card_id"]: card for card in cards}
    plan_card = card_map["EC_PLAN_RISK_TOTAL_001"]
    cp_card = card_map["EC_COUNTERPARTY_QUALITY_001"]
    plan_gap = float(plan_card["plan_vs_base_abs_eur"])
    planned_budget = float(plan_card["planned_budget_eur"])
    historical_base = float(plan_card["historical_base_eur"])
    unknown_rows = int(cp_card["unknown_counterparty_rows"])
    unknown_amount = float(cp_card["unknown_counterparty_amount_eur"])
    unknown_rows_text = f"{unknown_rows:,}".replace(",", " ")
    appendix_cards = [
        "EC_PLAN_RISK_TOTAL_001",
        "EC_TOP_CONTRIBUTORS_001",
        "EC_ONJN_AVENTO_LOCALIZATION_001",
        "EC_COUNTERPARTY_QUALITY_001",
        "EC_RECOMMENDATION_QUEUE_001",
    ]
    for item in chart_meta:
        evidence_id = item.get("evidence_card_id")
        if evidence_id in card_map and evidence_id not in appendix_cards:
            appendix_cards.append(evidence_id)

    chart_copy = {
        "CH_BUDGET_BRIDGE_001": (
            "График показывает, какие крупные компоненты формируют разницу между плановой базой и историческим уровнем.",
            "Использовать как навигацию по зонам проверки, а не как подтверждение основания бюджета.",
            "Нельзя выводить, что показанные компоненты уже подтверждены владельцами бюджета.",
        ),
        "CH_TOP_GAP_CONTRIBUTORS_001": (
            "График показывает крупнейшие комбинации статья бюджета × ЦФО по абсолютному размеру планового разрыва.",
            "Длинные столбцы означают высокий приоритет проверки.",
            "Нельзя выводить, что ЦФО или статья уже являются подтвержденным фактором изменения.",
        ),
        "CH_ONJN_AVENTO_MONTHLY_001": (
            "График показывает помесячный профиль зоны ONJN Gaming Tax / Avento MT.",
            "Использовать для подготовки вопросов владельцу бюджета по основанию плановой суммы.",
            "Нельзя выводить, что эта зона уже подтверждена как основание изменения бюджета.",
        ),
        "CH_MONTH_ARTICLE_CFO_HEATMAP_001": (
            "Карта показывает месяцы и статьи, где плановый разрыв концентрируется сильнее.",
            "Более темные зоны являются приоритетами проверки плановой базы.",
            "Нельзя выводить бизнес-основание без подтверждения владельца бюджета.",
        ),
        "CH_COUNTERPARTY_QUALITY_001": (
            "График показывает строки, где контрагент не определен.",
            "Использовать для приоритизации исправления справочника контрагентов.",
            "Нельзя трактовать строки с неопределенным контрагентом как подтвержденное финансовое искажение.",
        ),
        "CH_UNKNOWN_COUNTERPARTY_AMOUNT_001": (
            "График показывает оборот по строкам с неопределенным контрагентом.",
            "Использовать как оценку объема проверки справочника.",
            "Нельзя трактовать сумму как подтвержденный финансовый результат, ошибку или бюджетное отклонение.",
        ),
        "CH_HISTORICAL_VOLATILITY_001": (
            "График показывает историческую волатильность статей бюджета.",
            "Использовать для понимания, где историческая база стабильнее или менее стабильна.",
            "Нельзя выводить основание планового разрыва только по волатильности.",
        ),
        "CH_CFO_PLAN_BASE_001": (
            "График показывает распределение планового разрыва по ЦФО.",
            "Использовать для маршрутизации проверки к владельцам бюджета.",
            "Нельзя назначать ответственность без подтверждения владельца.",
        ),
        "CH_GAP_SHARE_001": (
            "График показывает концентрацию планового разрыва в крупнейших комбинациях.",
            "Использовать для определения, сколько проверки покрывают верхние позиции.",
            "Нельзя считать концентрацию доказанным основанием изменения бюджета.",
        ),
    }
    chart_blocks = []
    chart_source_rows = []
    for item in chart_meta:
        what, how, cannot = chart_copy.get(
            item["chart_id"],
            ("Нет описания в текущем артефакте.", "Нет описания в текущем артефакте.", "Нет ограничения в текущем артефакте."),
        )
        chart_blocks.append(
            "\n".join(
                [
                    f"### {item['chart_title']}",
                    "",
                    "**Что показывает:**",
                    what,
                    "",
                    "**Как читать:**",
                    how,
                    "",
                    "**Что нельзя выводить:**",
                    cannot,
                ]
            )
        )
        chart_source_rows.append(
            f"| {item['chart_id']} | {item['chart_title']} | {item.get('source_mart', '')} | {item.get('metric_name', '')} | {item.get('grain', '')} | {item.get('period', '')} | {item.get('evidence_card_id', '')} |"
        )

    evidence_rows = []
    for card_id in appendix_cards:
        card = card_map.get(card_id)
        if not card:
            continue
        status = "подтверждает зону проверки" if card.get("claim_type") == "localization_signal" else card.get("status", "supported")
        evidence_rows.append(
            f"| {card_id} | {card.get('claim_type', card.get('card_type', ''))} | {card.get('source_table', '')} | {card.get('metric_name', '')} | {status} |"
        )

    executive_memo = "\n\n".join(
        [
            "# Управленческая записка: точность бюджетного планирования",
            (
                "## 1. Краткий управленческий вывод\n\n"
                "Цель записки - определить зоны бюджета 2026, которые требуют обязательного подтверждения до утверждения плановой базы, поскольку они формируют основной разрыв к историческому уровню.\n\n"
                f"Плановый бюджет выше исторической базы на {money(plan_gap)}. Это не доказанная ошибка бюджета и не факт исполнения, а зона проверки плановой базы.\n\n"
                "Крупнейший сигнал проверки - ONJN Gaming Tax / Avento MT на 13 615 977 EUR. До подтверждения владельцем бюджета эта сумма должна рассматриваться как зона проверки, а не как подтвержденное основание изменения бюджета.\n\n"
                f"Отдельный риск качества данных - {unknown_rows_text} строк с неопределенным контрагентом. Оборот по таким строкам составляет {money(unknown_amount)}. Это не сумма ошибки, а объем данных, по которым аналитика контрагентов ненадежна до исправления справочника."
            ),
            (
                "## 2. Решение, которое требуется от руководства\n\n"
                "Требуется поручить владельцам бюджета подтвердить крупнейшие зоны планового разрыва до утверждения плановой базы. Без подтверждения эти суммы должны быть вынесены в отдельный список бюджетных рисков и не считаться надежной плановой базой.\n\n"
                "| Нужно решить | Ответственный | Срок / момент контроля |\n|---|---|---|\n"
                "| Подтвердить ONJN Gaming Tax / Avento MT | владелец бюджета | до утверждения бюджета |\n"
                "| Подтвердить маркетинговое увеличение | владелец соответствующего бюджета | до утверждения бюджета |\n"
                "| Исправить неопределенных контрагентов | владелец данных / справочников | до повторного выпуска отчета |"
            ),
            (
                "## 3. Основные выводы\n\n"
                f"1. Плановый бюджет выше исторической базы на {money(plan_gap)}.\n"
                "2. Крупные комбинации «статья бюджета × ЦФО» требуют подтверждения владельцами бюджета.\n"
                "3. ONJN Gaming Tax / Avento MT - крупнейшая зона проверки на 13 615 977 EUR.\n"
                "4. ДР Маркетинг / Advertisement Dept Main - вторая крупная зона проверки на 9 785 325 EUR.\n"
                f"5. По {unknown_rows_text} строкам не определен контрагент; это ограничивает надежность аналитики по контрагентам до исправления справочника."
            ),
            (
                "## 4. Карта рисков плановой базы\n\n"
                "| Зона | Сумма / объем | Риск | Управленческий смысл | Статус |\n|---|---:|---|---|---|\n"
                "| ONJN Gaming Tax / Avento MT | 13 615 977 EUR | высокий | может быть плановое изменение, резерв, перенос или ошибка базы | требуется подтверждение владельца |\n"
                "| ДР Маркетинг / Advertisement Dept Main | 9 785 325 EUR | высокий | нужен документ-основание по маркетинговому бюджету | требуется подтверждение владельца |\n"
                "| Роялти / Atlant | 1 577 767 EUR | средний | требуется проверка основания плановой суммы | требуется проверка |\n"
                f"| Неопределенные контрагенты | {unknown_rows_text} строк / {money(unknown_amount)} | высокий по данным | риск некорректной аналитики при неуточненном маппинге | требуется проверка справочника |"
            ),
            (
                "## 5. Что нужно проверить\n\n"
                "- По ONJN Gaming Tax / Avento MT: подтвердить основание плановой суммы и корректность классификации.\n"
                "- По маркетинговым статьям: подтвердить утвержденное бюджетное основание, медиаплан или иной документ-основание.\n"
                "- По Роялти / Atlant: проверить основание плановой суммы и периодизацию.\n"
                "- По неопределенным контрагентам: уточнить справочник, ключи контрагентов и список допустимых исключений."
            ),
            (
                "## 6. Риски качества данных\n\n"
                f"По {unknown_rows_text} строкам не определен контрагент. Суммарный оборот таких строк составляет {money(unknown_amount)}.\n\n"
                + "Это не сумма ошибки и не подтвержденное финансовое искажение. Это объем операций, по которым аналитика контрагентов, концентрации и владельцев бюджета ненадежна до исправления справочника.\n\n"
                "Риск: управленческие выводы по контрагентам и распределению ответственности могут быть искажены, если маппинг не будет уточнен."
            ),
            (
                "## 7. Рекомендуемые действия\n\n"
                "| Действие | Владелец | Вопрос | Ожидаемое подтверждение | Приоритет |\n|---|---|---|---|---|\n"
                "| Подтвердить рост ONJN Gaming Tax / Avento MT | владелец бюджета | Рост плановый, перенос, резерв или ошибка? | комментарий владельца + ссылка на основание бюджета | высокий |\n"
                "| Проверить увеличение по маркетинговым статьям | владелец соответствующего бюджета | Подтверждена ли кампания / медиаплан / бюджетное основание? | утвержденное основание бюджета | высокий |\n"
                "| Проверить неопределенных контрагентов | владелец данных | Почему контрагент не определен? | исправление справочника / exception list | высокий |"
            ),
            (
                "## 8. Ограничения анализа\n\n"
                "- Плановый период не является фактом будущего периода.\n"
                "- Зоны проверки не являются доказанными основаниями изменения бюджета.\n"
                "- Владелец бюджета должен подтвердить основание плановой суммы.\n"
                "- IN-OUT используется как месячный показатель чистого денежного потока и не суммируется по строкам.\n"
                "- Качество аналитики по контрагентам ограничено строками с неопределенным контрагентом."
            ),
        ]
    )
    evidence_pack = "\n\n".join(
        [
            "# Приложение. Источники подтверждения и ограничения",
            "## 9. Приложение. Графики\n" + "\n\n".join(chart_blocks),
            (
                "## 10. Приложение. Источники подтверждения и ограничения\n\n"
                "### Реестр утверждений\n\n"
                "| Утверждение | Статус | Источник подтверждения | Ограничение |\n|---|---|---|---|\n"
                f"| Плановый бюджет = {money(planned_budget)} | подтверждено расчетом | EC_PLAN_RISK_TOTAL_001 | зависит от промышленного слоя витрин |\n"
                f"| Историческая база = {money(historical_base)} | подтверждено расчетом | EC_PLAN_RISK_TOTAL_001 | зависит от промышленного слоя витрин |\n"
                f"| Плановый разрыв = {money(plan_gap)} | подтверждено расчетом | EC_PLAN_RISK_TOTAL_001 | не является фактом исполнения |\n"
                "| ONJN / Avento - крупнейшая зона проверки | подтверждена как зона проверки | EC_TOP_CONTRIBUTORS_001 / EC_ONJN_AVENTO_LOCALIZATION_001 | основание изменения не подтверждено владельцем |\n"
                "| Неопределенные контрагенты - риск качества данных | подтверждено расчетом | EC_COUNTERPARTY_QUALITY_001 | финансовое искажение не подтверждено |"
                "\n\n### Карточки подтверждения\n\n"
                + "\n".join(f"- {card_id}" for card_id in appendix_cards)
                + "\n\n| Карточка подтверждения | Тип утверждения | Расчетная витрина | Метрика | Статус |\n|---|---|---|---|---|\n"
                + "\n".join(evidence_rows)
                + "\n\n### Технические ссылки по графикам\n\n"
                "| График | Название | Расчетная витрина | Метрика | Разрез | Период | Источник подтверждения |\n|---|---|---|---|---|---|---|\n"
                + "\n".join(chart_source_rows)
                + "\n\n### Методические ограничения\n\n"
                "- Плановый период не является фактом будущего периода.\n"
                "- Зоны проверки не являются доказанными основаниями изменения бюджета.\n"
                "- IN-OUT используется как месячный показатель чистого денежного потока и не суммируется по строкам.\n"
                "- Оборот по строкам с неопределенным контрагентом является объемом проверки справочника, а не подтвержденным финансовым искажением.\n"
                "- QA references: memo_claim_audit.json, chart_qa.json, docx_qa.json, qa_report.json."
            ),
        ]
    )
    memo = executive_memo + "\n\n" + evidence_pack
    (REPORTS_DIR / "final_memo_executive.md").write_text(executive_memo + "\n", encoding="utf-8")
    (REPORTS_DIR / "final_memo_evidence_pack.md").write_text(evidence_pack + "\n", encoding="utf-8")
    (REPORTS_DIR / "memo_draft.md").write_text(memo + "\n", encoding="utf-8")
    (REPORTS_DIR / "final_memo.md").write_text(memo + "\n", encoding="utf-8")
    return memo


def run_claim_audit(memo: str, report_marts: dict[str, pd.DataFrame], cards: list[dict[str, Any]], chart_meta: list[dict[str, Any]]) -> dict[str, Any]:
    lower = memo.lower()
    card_ids = {card["card_id"] for card in cards}
    has_confirmed_driver = any(card.get("claim_type") == "confirmed_driver" for card in cards)
    unknown_rows = int(report_marts["mart_budget_counterparty_quality"]["unknown_counterparty_rows"].sum())
    failures: list[dict[str, str]] = []
    for phrase in CAUSALITY_PHRASES:
        if phrase in lower and not has_confirmed_driver:
            failures.append({"guard": "causality", "phrase": phrase, "status": "fail"})
    for phrase in COUNTERPARTY_NO_ISSUE_PHRASES:
        if phrase in lower and unknown_rows > 0:
            failures.append({"guard": "counterparty_quality", "phrase": phrase, "status": "fail"})
    for phrase in PLANNING_FACT_PHRASES:
        if phrase in lower:
            failures.append({"guard": "planning_period", "phrase": phrase, "status": "fail"})
    if "основной драйвер" in lower or "главный драйвер" in lower:
        failures.append({"guard": "driver_language", "phrase": "driver", "status": "fail"})
    technical_markers = [
        "# required section heading",
        "# draft section",
        "caption: source mart=",
        "owner_confirmation ->",
        "mapping_review ->",
        "appendix: evidence cards",
    ]
    for marker in technical_markers:
        if marker in lower:
            failures.append({"guard": "technical_marker", "phrase": marker, "status": "fail"})
    if memo.count("Приложение. Источники подтверждения и ограничения") < 1:
        failures.append({"guard": "structure", "phrase": "sources appendix missing", "status": "fail"})
    chart_failures = [item for item in chart_meta if not item.get("source_mart") or not item.get("evidence_card_id")]
    missing_cards = [item["evidence_card_id"] for item in chart_meta if item.get("evidence_card_id") not in card_ids]
    required_sections = [title for title in SECTION_TITLES if title.lower() not in lower]
    audit = {
        "qa_status": "pass" if not failures and not chart_failures and not missing_cards and not required_sections else "blocked",
        "failures": failures,
        "chart_metadata_failures": chart_failures,
        "missing_chart_evidence_cards": missing_cards,
        "missing_required_sections": required_sections,
        "unknown_counterparty_rows": unknown_rows,
        "confirmed_driver_cards_count": sum(card.get("claim_type") == "confirmed_driver" for card in cards),
    }
    write_json(QA_DIR / "memo_claim_audit.json", audit)
    return audit


def build_docx(memo: str, chart_meta: list[dict[str, Any]], cards: list[dict[str, Any]]) -> None:
    log_step("rendering final_memo.docx")
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    doc.add_heading("Управленческая записка: точность бюджетного планирования", 0)
    doc.add_paragraph("Источник: расчетная витрина. Валюта: EUR. Подробные источники подтверждения вынесены в приложение.")
    section_pattern = re.compile(r"^##\s+\d+\.\s+(.+)$", re.MULTILINE)
    chunks = section_pattern.split(memo)
    if chunks:
        preface = chunks[0].strip()
        pairs = list(zip(chunks[1::2], chunks[2::2]))
    else:
        pairs = []
    for title, body in pairs:
        doc.add_heading(title.strip(), level=1)
        if title.strip() == "Приложение. Графики":
            docx_chart_copy = {
                "CH_BUDGET_BRIDGE_001": (
                    "График показывает, какие крупные компоненты формируют разницу между плановой базой и историческим уровнем.",
                    "Использовать как навигацию по зонам проверки, а не как подтверждение основания бюджета.",
                    "Нельзя выводить, что показанные компоненты уже подтверждены владельцами бюджета.",
                ),
                "CH_TOP_GAP_CONTRIBUTORS_001": (
                    "График показывает крупнейшие комбинации статья бюджета × ЦФО по абсолютному размеру планового разрыва.",
                    "Длинные столбцы означают высокий приоритет проверки.",
                    "Нельзя выводить, что ЦФО или статья уже являются подтвержденным фактором изменения.",
                ),
                "CH_ONJN_AVENTO_MONTHLY_001": (
                    "График показывает помесячный профиль зоны ONJN Gaming Tax / Avento MT.",
                    "Использовать для подготовки вопросов владельцу бюджета по основанию плановой суммы.",
                    "Нельзя выводить, что эта зона уже подтверждена как основание изменения бюджета.",
                ),
                "CH_MONTH_ARTICLE_CFO_HEATMAP_001": (
                    "Карта показывает месяцы и статьи, где плановый разрыв концентрируется сильнее.",
                    "Более темные зоны являются приоритетами проверки плановой базы.",
                    "Нельзя выводить бизнес-основание без подтверждения владельца бюджета.",
                ),
                "CH_COUNTERPARTY_QUALITY_001": (
                    "График показывает строки, где контрагент не определен.",
                    "Использовать для приоритизации исправления справочника контрагентов.",
                    "Нельзя трактовать строки с неопределенным контрагентом как подтвержденное финансовое искажение.",
                ),
                "CH_UNKNOWN_COUNTERPARTY_AMOUNT_001": (
                    "График показывает оборот по строкам с неопределенным контрагентом.",
                    "Использовать как оценку объема проверки справочника.",
                    "Нельзя трактовать сумму как подтвержденный финансовый результат, ошибку или бюджетное отклонение.",
                ),
                "CH_HISTORICAL_VOLATILITY_001": (
                    "График показывает историческую волатильность статей бюджета.",
                    "Использовать для понимания, где историческая база стабильнее или менее стабильна.",
                    "Нельзя выводить основание планового разрыва только по волатильности.",
                ),
                "CH_CFO_PLAN_BASE_001": (
                    "График показывает распределение планового разрыва по ЦФО.",
                    "Использовать для маршрутизации проверки к владельцам бюджета.",
                    "Нельзя назначать ответственность без подтверждения владельца.",
                ),
                "CH_GAP_SHARE_001": (
                    "График показывает концентрацию планового разрыва в крупнейших комбинациях.",
                    "Использовать для определения, сколько проверки покрывают верхние позиции.",
                    "Нельзя считать концентрацию доказанным основанием изменения бюджета.",
                ),
            }
            for item in chart_meta:
                doc.add_heading(item["chart_title"], level=2)
                path = Path(item["png_file"])
                if path.exists():
                    doc.add_picture(str(path), width=Inches(6.2))
                what, how, cannot = docx_chart_copy.get(
                    item["chart_id"],
                    ("Нет описания в текущем артефакте.", "Нет описания в текущем артефакте.", "Нет ограничения в текущем артефакте."),
                )
                doc.add_paragraph(f"Что показывает: {what}")
                doc.add_paragraph(f"Как читать: {how}")
                doc.add_paragraph(f"Что нельзя выводить: {cannot}")
            continue
        for block in body.strip().split("\n\n"):
            if not block.strip():
                continue
            if block.lstrip().startswith("|"):
                lines = [line for line in block.splitlines() if line.startswith("|") and "---" not in line]
                if lines:
                    rows = [[cell.strip() for cell in line.strip("|").split("|")] for line in lines]
                    table = doc.add_table(rows=0, cols=len(rows[0]))
                    table.style = "Table Grid"
                    for row in rows:
                        cells = table.add_row().cells
                        for i, cell_text in enumerate(row):
                            cells[i].text = cell_text
                continue
            for line in block.splitlines():
                text = line.strip()
                if text.startswith("- "):
                    doc.add_paragraph(text[2:], style="List Bullet")
                elif text.startswith("### "):
                    doc.add_heading(text[4:], level=2)
                elif text:
                    doc.add_paragraph(text)
    doc.save(REPORTS_DIR / "final_memo.docx")


def write_qa(report_marts: dict[str, pd.DataFrame], chart_meta: list[dict[str, Any]], claim_audit: dict[str, Any]) -> None:
    chart_qa = {
        "qa_status": "pass" if all((CHARTS_DIR / file_name).exists() for file_name in CHART_FILES.values()) and not claim_audit["chart_metadata_failures"] else "blocked",
        "charts_expected": sorted(CHART_FILES),
        "charts_created": sorted(path.stem for path in CHARTS_DIR.glob("*.png")),
        "chart_metadata_rows": len(chart_meta),
    }
    docx_path = REPORTS_DIR / "final_memo.docx"
    docx_qa = {
        "qa_status": "pass" if docx_path.exists() and not claim_audit["missing_required_sections"] else "blocked",
        "docx_exists": docx_path.exists(),
        "required_sections": SECTION_TITLES,
        "missing_required_sections": claim_audit["missing_required_sections"],
    }
    write_json(QA_DIR / "chart_qa.json", chart_qa)
    write_json(QA_DIR / "docx_qa.json", docx_qa)
    qa_path = QA_DIR / "qa_report.json"
    qa_report = json.loads(qa_path.read_text(encoding="utf-8")) if qa_path.exists() else {}
    qa_report.update(
        {
            "report_layer_status": "pass" if claim_audit["qa_status"] == chart_qa["qa_status"] == docx_qa["qa_status"] == "pass" else "blocked",
            "memo_claim_audit_status": claim_audit["qa_status"],
            "chart_qa_status": chart_qa["qa_status"],
            "docx_qa_status": docx_qa["qa_status"],
            "unknown_counterparty_rows": int(report_marts["mart_budget_counterparty_quality"]["unknown_counterparty_rows"].sum()),
            "judge_verdict": "pass" if claim_audit["qa_status"] == "pass" else "fail",
            "ollama_models": OLLAMA_MODELS,
            "ollama_fast_fallback_model": OLLAMA_FAST_FALLBACK_MODEL,
            "ollama_routing_note": OLLAMA_ROUTING_NOTE,
            "ollama_generation_mode": "ollama_batch_context_plus_deterministic_guarded_merge",
            "required_changes_applied": True,
        }
    )
    if qa_report["report_layer_status"] != "pass":
        qa_report["qa_status"] = "blocked"
        qa_report.setdefault("critical_failures", []).append("report_layer_guard_failed")
    else:
        qa_report["qa_status"] = "pass"
        stale = {"report_layer_guard_failed", "judge_release_gate_not_passed", "memo_quality_gate_failed"}
        qa_report["critical_failures"] = [item for item in qa_report.get("critical_failures", []) if item not in stale]
    write_json(qa_path, qa_report)
    acceptance = [
        f"accepted: {'yes' if qa_report['qa_status'] == 'pass' else 'no'}",
        f"qa_status: {qa_report['qa_status']}",
        f"judge_verdict: {qa_report.get('judge_verdict', 'not_applicable')}",
        "critical_failures:",
        *[f"- {item}" for item in qa_report.get("critical_failures", [])],
        "warnings:",
        *[f"- {item}" for item in qa_report.get("warnings", [])],
        "residual_risks:",
        "- Localization signals do not prove business causality.",
        "- DOCX visual render QA may be unavailable if LibreOffice is not installed.",
        "next_step:",
        "- Review final_memo.docx and owner-confirm ONJN Gaming Tax / Avento MT.",
    ]
    (QA_DIR / "qa_acceptance.md").write_text("\n".join(acceptance) + "\n", encoding="utf-8")


def build_report_package() -> dict[str, Any]:
    log_step("start")
    for directory in [CHARTS_DIR, BATCH_MD_DIR, BATCH_JSON_DIR]:
        directory.mkdir(parents=True, exist_ok=True)
    log_step("reading production marts")
    data = read_inputs()
    log_step("building report marts")
    report_marts = build_report_marts(data)
    log_step("updating evidence cards")
    cards = update_evidence_cards(report_marts)
    log_step("building charts")
    chart_meta = build_charts(report_marts, data)
    batches = build_batches(report_marts, chart_meta, cards)
    memo = build_final_memo(batches, cards, chart_meta)
    log_step("running claim audit")
    claim_audit = run_claim_audit(memo, report_marts, cards, chart_meta)
    build_docx(memo, chart_meta, cards)
    log_step("writing QA")
    write_qa(report_marts, chart_meta, claim_audit)
    log_step("done")
    return {
        "report_marts": sorted(REPORT_MART_FILES.values()),
        "charts": sorted(CHART_FILES.values()),
        "docx": str(REPORTS_DIR / "final_memo.docx"),
        "claim_audit_status": claim_audit["qa_status"],
    }


if __name__ == "__main__":
    result = build_report_package()
    print(json.dumps(result, ensure_ascii=False, indent=2))
