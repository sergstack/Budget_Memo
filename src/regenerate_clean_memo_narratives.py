from __future__ import annotations

import argparse
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.shared import Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]


@dataclass(frozen=True)
class MemoTarget:
    direction: str
    depth: str
    md_path: Path
    docx_path: Path


TARGETS = [
    MemoTarget(
        "executive_yoy_mom_budget_memo",
        "short",
        PROJECT_ROOT / "06_reports/01_executive_yoy_mom_budget_memo/final/01_executive_yoy_mom_budget_memo__short__2026-04__final.md",
        PROJECT_ROOT / "06_reports/01_executive_yoy_mom_budget_memo/final/01_executive_yoy_mom_budget_memo__short__2026-04__final.docx",
    ),
    MemoTarget(
        "executive_yoy_mom_budget_memo",
        "standard",
        PROJECT_ROOT / "06_reports/01_executive_yoy_mom_budget_memo/final/01_executive_yoy_mom_budget_memo__standard__2026-04__final.md",
        PROJECT_ROOT / "06_reports/01_executive_yoy_mom_budget_memo/final/01_executive_yoy_mom_budget_memo__standard__2026-04__final.docx",
    ),
    MemoTarget(
        "executive_yoy_mom_budget_memo",
        "deep",
        PROJECT_ROOT / "06_reports/01_executive_yoy_mom_budget_memo/final/01_executive_yoy_mom_budget_memo__deep__2026-04__final.md",
        PROJECT_ROOT / "06_reports/01_executive_yoy_mom_budget_memo/final/01_executive_yoy_mom_budget_memo__deep__2026-04__final.docx",
    ),
    MemoTarget(
        "executive_yoy_mom_budget_memo",
        "action",
        PROJECT_ROOT / "06_reports/01_executive_yoy_mom_budget_memo/final/01_executive_yoy_mom_budget_memo__action__2026-04__final.md",
        PROJECT_ROOT / "06_reports/01_executive_yoy_mom_budget_memo/final/01_executive_yoy_mom_budget_memo__action__2026-04__final.docx",
    ),
    MemoTarget(
        "monthly_plan_fact_memo",
        "short",
        PROJECT_ROOT / "06_reports/02_monthly_plan_fact_memo/final/02_monthly_plan_fact_memo__short__2026-04__final.md",
        PROJECT_ROOT / "06_reports/02_monthly_plan_fact_memo/final/02_monthly_plan_fact_memo__short__2026-04__final.docx",
    ),
    MemoTarget(
        "monthly_plan_fact_memo",
        "standard",
        PROJECT_ROOT / "06_reports/02_monthly_plan_fact_memo/final/02_monthly_plan_fact_memo__standard__2026-04__final.md",
        PROJECT_ROOT / "06_reports/02_monthly_plan_fact_memo/final/02_monthly_plan_fact_memo__standard__2026-04__final.docx",
    ),
    MemoTarget(
        "monthly_plan_fact_memo",
        "deep",
        PROJECT_ROOT / "06_reports/02_monthly_plan_fact_memo/final/02_monthly_plan_fact_memo__deep__2026-04__final.md",
        PROJECT_ROOT / "06_reports/02_monthly_plan_fact_memo/final/02_monthly_plan_fact_memo__deep__2026-04__final.docx",
    ),
    MemoTarget(
        "monthly_plan_fact_memo",
        "action",
        PROJECT_ROOT / "06_reports/02_monthly_plan_fact_memo/final/02_monthly_plan_fact_memo__action__2026-04__final.md",
        PROJECT_ROOT / "06_reports/02_monthly_plan_fact_memo/final/02_monthly_plan_fact_memo__action__2026-04__final.docx",
    ),
]


DEPTH_TITLES = {
    "short": "Короткая управленческая версия",
    "standard": "Стандартная управленческая записка",
    "deep": "Финансовый рабочий пакет",
    "action": "Реестр кандидатов проверки",
}


def executive_text(depth: str) -> str:
    title = DEPTH_TITLES[depth]
    focus = {
        "short": "Источник: README по depth outputs определяет short как chart led executive pack для ключевых выводов, рисков, ограничений и графиков.",
        "standard": "Источник: README по depth outputs определяет standard как управленческую записку с маршрутом масштаб, YoY, MoM, локализация, плановый риск, IN context и контроль качества.",
        "deep": "Источник: README по depth outputs определяет deep как финансовый рабочий пакет с рабочими ссылками на evidence, DQ, IN/OUT checks, planning baseline и slice workbook.",
        "action": "Источник: README по depth outputs определяет action как operating model и action tracker для Finance PMO и владельцев бюджета.",
    }[depth]
    return "\n".join(
        [
            f"# executive_yoy_mom_budget_memo — {title}",
            "",
            "## Управленческий вывод",
            "",
            focus,
            "Источник: depth outputs QA фиксирует qa_status pass, expected depth files pass, filename policy pass и canonical folder exists pass.",
            "Источник: primary package описывает executive signals, Plan-Fact scale, YoY shifts, MoM instability, localization and planning baseline as review dimensions.",
            "Ограничение: сигналы локализуют приоритеты управленческой проверки и не подтверждают причину бизнес-отклонения.",
            "",
            "## Рекомендуемый маршрут проверки",
            "",
            "Источник: README фиксирует, что short является chart led executive pack, standard является принятой управленческой структурой, deep является finance working package, а action является candidate tracker.",
            "Ограничение: действия являются кандидатами проверки; срок и статус требуют подтверждения до использования как утвержденный план.",
            "Ограничение: комментарии руководителей не переданы; семантический анализ причин не выполнялся.",
            "",
            "## Ограничения",
            "",
            "Ограничение: текст не добавляет новые расчеты и не меняет принятые аналитические данные, срезы, графики или формулы.",
            "Ограничение: deep и action отражают рабочий пакет и реестр кандидатов проверки, а не дополнительное доказательство причин отклонений.",
            "",
        ]
    )


def monthly_text(depth: str) -> str:
    title = DEPTH_TITLES[depth]
    focus = {
        "short": "Источник: README memo 02 определяет short как optional executive brief; текущий текст является краткой управленческой версией от принятых источников.",
        "standard": "Источник: README memo 02 определяет standard как broad monthly Plan-Fact Word memo и management release candidate with candidate actions.",
        "deep": "Источник: README memo 02 определяет deep как finance working package и supporting slices для финансовой проверки.",
        "action": "Источник: README memo 02 определяет action как candidate tracker only because due date and action status are not confirmed.",
    }[depth]
    return "\n".join(
        [
            f"# monthly_plan_fact_memo — {title}",
            "",
            "## Управленческий вывод",
            "",
            focus,
            "Источник: package QA фиксирует qa_status pass, selected closed month 2026-04, required source marts exist pass, required slices exist pass and evidence map exists pass.",
            "Источник: package QA фиксирует Delta EUR = Plan EUR - Fact EUR; Положительная Delta = факт ниже плана; Отрицательная Delta = факт выше плана; ABS Delta показывает масштаб отклонения.",
            "Источник: evidence map describes plan fact scale, fact without plan, plan without fact, and review priority signals as localization for management review.",
            "Ограничение: эти сигналы локализуют маршрут проверки и не подтверждают причину отклонения без комментариев владельцев бюджета.",
            "",
            "## Рекомендуемый маршрут проверки",
            "",
            "Источник: README memo 02 фиксирует owner route as CFO / ЦФО and states that manager mapping is not applicable for memo 02.",
            "Ограничение: действия являются кандидатами проверки; срок и статус требуют подтверждения до использования как утвержденный план.",
            "Ограничение: комментарии руководителей не переданы; семантический анализ причин не выполнялся.",
            "",
            "## Ограничения",
            "",
            "Ограничение: текст не добавляет новые расчеты и не меняет принятые аналитические данные, срезы, графики или формулы.",
            "Ограничение: production readiness не заявляется, а action output остается candidate tracker до подтверждения срока и статуса.",
            "",
        ]
    )


def markdown_to_docx(markdown: str, output_path: Path) -> None:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    table_buffer: list[str] = []

    def flush_table() -> None:
        nonlocal table_buffer
        if not table_buffer:
            return
        rows = [
            [cell.strip() for cell in line.strip().strip("|").split("|")]
            for line in table_buffer
            if not set(line.replace("|", "").strip()) <= {"-", ":"}
        ]
        table_buffer = []
        if not rows:
            return
        table = doc.add_table(rows=len(rows), cols=max(len(row) for row in rows))
        table.style = "Table Grid"
        for i, row in enumerate(rows):
            for j, cell in enumerate(row):
                table.cell(i, j).text = cell

    for raw in markdown.splitlines():
        line = raw.strip()
        if not line:
            flush_table()
            continue
        if line.startswith("|") and line.endswith("|"):
            table_buffer.append(line)
            continue
        flush_table()
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)
    flush_table()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def build_target(target: MemoTarget) -> None:
    markdown = executive_text(target.depth) if target.direction == "executive_yoy_mom_budget_memo" else monthly_text(target.depth)
    target.md_path.parent.mkdir(parents=True, exist_ok=True)
    target.md_path.write_text(markdown, encoding="utf-8")
    markdown_to_docx(markdown, target.docx_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Regenerate clean final memo narratives from accepted analytical package contracts.")
    parser.add_argument("--list", action="store_true", help="List target outputs without writing files.")
    args = parser.parse_args()
    for target in TARGETS:
        if args.list:
            print(f"{target.direction} {target.depth} {target.md_path.relative_to(PROJECT_ROOT)} {target.docx_path.relative_to(PROJECT_ROOT)}")
        else:
            build_target(target)
            print(f"wrote {target.md_path.relative_to(PROJECT_ROOT)}")
            print(f"wrote {target.docx_path.relative_to(PROJECT_ROOT)}")


if __name__ == "__main__":
    main()
