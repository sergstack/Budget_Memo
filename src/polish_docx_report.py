from __future__ import annotations

import json
import os
import re
import shutil
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    from src.progress import log_progress
except ImportError:  # pragma: no cover
    from progress import log_progress


PROJECT_ROOT = Path(__file__).resolve().parents[1]
RAW_DIR = PROJECT_ROOT / "01_raw"
STAGE_DIR = PROJECT_ROOT / "02_stage"
MARTS_DIR = PROJECT_ROOT / "03_marts"
CHARTS_DIR = PROJECT_ROOT / "04_charts"
REPORTS_DIR = PROJECT_ROOT / "06_reports"
MEMO_01_DIR = REPORTS_DIR / "01_executive_yoy_mom_budget_memo"
MEMO_01_FINAL_DIR = MEMO_01_DIR / "final"
MEMO_01_DRAFT_DIR = MEMO_01_DIR / "draft"
LLM_DIR = PROJECT_ROOT / "05_llm_package"
QA_DIR = PROJECT_ROOT / "07_qa"

DOCX_PATH = MEMO_01_FINAL_DIR / "memo_01__executive_yoy_mom_budget_memo__2026-05-21__final.docx"
FINAL_MD = MEMO_01_FINAL_DIR / "memo_01__executive_yoy_mom_budget_memo__2026-05-21__final.md"
REVISED_MD = MEMO_01_DRAFT_DIR / "memo_01__executive_yoy_mom_budget_memo__2026-05-21__review.md"
REVISED_QA = QA_DIR / "executive_yoy_mom_memo_revised_qa.json"
EVIDENCE_MAP = LLM_DIR / "executive_yoy_mom_evidence_map.xlsx"
CHART_CATALOG = CHARTS_DIR / "chart_catalog.xlsx"
DOCX_QA = QA_DIR / "executive_yoy_mom_docx_qa.json"
VISUAL_QA = QA_DIR / "executive_yoy_mom_visual_qa.json"
VISUAL_QA_SUMMARY = QA_DIR / "executive_yoy_mom_visual_qa_summary.md"
OLD_FINAL_MEMO = PROJECT_ROOT / "99_archive" / "report_cleanup_before_reorg_2026-05-21" / "final_memo.docx"
RENDER_DIR = QA_DIR / "rendered_docx_pages"
MEMO_PROFILE = "executive_yoy_mom_budget_memo"
DEPTH_MODE = "depth_2_management_memo"

SECTION_ORDER = [
    "Рамка анализа",
    "Executive Summary",
    "Как читать записку",
    "Исторический факт: масштаб Plan-Fact",
    "YoY: сдвиг уровня к прошлому году",
    "MoM: помесячная динамика и нестабильность",
    "Локализация: статья × ЦФО",
    "Плановый риск: план к исторической базе",
    "iGaming flow context: отклонения к IN",
    "QC и ограничения",
    "Реестр приоритетных проверок",
    "Итоговый вывод",
]
BODY_CHARTS = {
    "CH_EXEC_001_PLAN_FACT_TOP_ABS",
    "CH_EXEC_002_YOY_TOP_SHIFT",
    "CH_EXEC_003_MOM_INSTABILITY",
    "CH_EXEC_004_LOCALIZATION_ARTICLE_CFO",
    "CH_EXEC_005_PLANNING_RISK",
    "CH_EXEC_006_IN_CONTEXT",
    "CH_EXEC_007_FLOW_BASE",
    "CH_EXEC_008_QA_LIMITATIONS",
}
APPENDIX_CHARTS = {"CH_EXEC_009_COUNTERPARTY_QUALITY", "CH_EXEC_010_CURRENCY_EXPOSURE"}


def snapshot(path: Path) -> dict[str, tuple[int, int]]:
    if not path.exists():
        return {}
    return {
        str(item.relative_to(PROJECT_ROOT)): (item.stat().st_mtime_ns, item.stat().st_size)
        for item in path.rglob("*")
        if item.is_file()
    }


def file_stat(path: Path) -> tuple[int, int] | None:
    if not path.exists():
        return None
    st = path.stat()
    return st.st_mtime_ns, st.st_size


def split_sections(markdown: str) -> dict[str, list[str]]:
    sections: dict[str, list[str]] = {}
    current: str | None = None
    for line in markdown.splitlines():
        if line.startswith("## "):
            current = line[3:].strip()
            sections[current] = []
        elif current is not None:
            sections[current].append(line.rstrip())
    return sections


def remove_evidence_tables(lines: list[str]) -> list[str]:
    cleaned: list[str] = []
    idx = 0
    while idx < len(lines):
        if lines[idx].strip() == "Evidence table:":
            idx += 1
            while idx < len(lines) and (not lines[idx].strip() or lines[idx].startswith("|")):
                idx += 1
            continue
        cleaned.append(lines[idx])
        idx += 1
    return cleaned


def strip_inline_evidence(text: str) -> str:
    text = re.sub(r"\s*Evidence:\s*(?:EV-[A-Z_]+-\d+|CH_EXEC_\d{3}_[A-Z0-9_]+_DATA|report_contract|draft_data_package)\.", ".", text)
    text = re.sub(r"Chart `CH_EXEC_\d{3}_[A-Z0-9_]+`:", "График:", text)
    text = text.replace("source_quality: refund_only", "ограничение качества источников: возвраты игрокам")
    text = text.replace("Source mix", "Состав источника")
    text = text.replace("source mix", "состав источника")
    text = text.replace("Limitation candidate", "Кандидат ограничения")
    text = text.replace("headline conclusion", "основной вывод")
    text = text.replace("Unknown counterparty exposure", "Экспозиция неизвестных контрагентов")
    text = text.replace("non-EUR exposure", "экспозицию не-EUR")
    text = text.replace("plan_fact_scale:", "сигнал масштаба Plan-Fact:")
    text = text.replace("yoy_shift:", "YoY-сдвиг:")
    text = text.replace("mom_instability:", "MoM-нестабильность:")
    text = text.replace("localization_concentration:", "локализация сигнала:")
    text = text.replace("planning_risk:", "плановый риск:")
    text = text.replace("flow_pressure:", "нагрузка на IN:")
    text = text.replace("abs_delta_eur ranked 1 in slice_plan_fact_article_cfo", "абсолютное отклонение EUR имеет ранг 1 в срезе статья × ЦФО")
    text = text.replace("abs_delta_eur", "ABS отклонение, EUR")
    text = text.replace("abs_yoy_delta_eur", "ABS YoY отклонение, EUR")
    text = text.replace("abs_mom_delta_eur", "ABS MoM отклонение, EUR")
    text = text.replace("cfo_abs_delta_eur", "ABS отклонение ЦФО, EUR")
    text = text.replace("plan_vs_base_abs_delta_eur", "ABS план к базе, EUR")
    text = text.replace("out_eur", "OUT, EUR")
    text = text.replace("denominator status", "статус denominator IN")
    text = text.replace("full-period IN denominator", "IN за полный период")
    text = text.replace("flow grain", "помесячном flow-срезе")
    text = text.replace("candidate_only_incomplete_action_fields", "кандидат проверки, требует подтверждения owner / срока / статуса")
    text = text.replace("production-ready", "готовый к публикации")
    text = text.replace("memo profile, chart package и report contract", "профилю записки, пакету графиков и контракту отчёта")
    text = text.replace("Source quality", "Сигнал качества источников")
    text = text.replace("appendix charts", "графики приложения")
    text = text.replace("headline management conclusion", "основной управленческий вывод")
    text = text.replace("Scope ограничен", "Охват ограничен")
    text = text.replace("готовый к публикации выводом", "готовым к публикации выводом")
    text = text.replace("evidence package", "пакет подтверждений")
    text = text.replace("Сигнал качества источников сигналы", "Сигналы качества источников")
    text = text.replace("MoM classification", "Классификация MoM")
    text = text.replace("YoY evidence", "YoY-подтверждению")
    text = text.replace("flow rows", "сервисные flow-строки")
    text = text.replace("..", ".")
    return text


def chart_id(line: str) -> str | None:
    match = re.search(r"CH_EXEC_\d{3}_[A-Z0-9_]+", line)
    return match.group(0) if match else None


def clean_body_lines(lines: list[str]) -> list[str]:
    cleaned = []
    for line in remove_evidence_tables(lines):
        if line.startswith("!["):
            continue
        if line.startswith("Appendix candidate:"):
            continue
        cleaned.append(strip_inline_evidence(line))
    return cleaned


def build_executive_table_rows() -> list[list[str]]:
    return [
        ["Масштаб", "Сигнал масштаба Plan-Fact: ONJN Gaming Tax / Avento MT", "ABS отклонение, EUR", "Сигнал масштаба, не причина", "Проверить природу Plan-Fact отклонения"],
        ["YoY", "YoY-сдвиг: ONJN Gaming Tax", "ABS YoY отклонение, EUR", "Только при доступной базе прошлого года", "Проверить базу сравнения"],
        ["MoM", "MoM-нестабильность: ONJN Gaming Tax", "ABS MoM отклонение, EUR", "Классификация, не причина", "Проверить месячную динамику"],
        ["Локализация", "Локализация сигнала: ONJN Gaming Tax / Avento MT", "ABS отклонение ЦФО, EUR", "Кандидат владельца, не подтверждённая ответственность", "Подтвердить owner route"],
        ["Плановый риск", "Плановый риск: ONJN Gaming Tax / Avento MT", "ABS план к базе, EUR", "Будущий риск, не факт исполнения", "Проверить план к исторической базе"],
        ["IN context", "Нагрузка на IN: 2026-12", "OUT, EUR", "Только как пропорциональность к притоку", "Проверить denominator IN"],
        ["QC", "Ограничение качества источников: возвраты игрокам", "ABS отклонение, EUR", "Кандидат ограничения, не основной вывод", "Проверить ограничения качества данных"],
    ]


def markdown_table(header: list[str], rows: list[list[str]]) -> list[str]:
    return [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join("---" for _ in header) + " |",
        *["| " + " | ".join(row) + " |" for row in rows],
    ]


def build_final_markdown() -> str:
    source = REVISED_MD.read_text(encoding="utf-8")
    sections = split_sections(source)
    chart_catalog = pd.read_excel(CHART_CATALOG)
    chart_lookup = {row["ID графика"]: row for row in chart_catalog.to_dict("records")}
    evidence_map = pd.read_excel(EVIDENCE_MAP)
    lines = [
        "# Управленческая записка YoY/MoM по бюджету",
        "",
        "Статус: управленческий DOCX-черновик на основе принятого MART, пакета графиков и согласованной версии записки. Аналитика не переписывалась.",
        "",
    ]
    for section_name in SECTION_ORDER:
        lines.extend([f"## {section_name}", ""])
        source_body = remove_evidence_tables(sections.get(section_name, []))
        for source_line in source_body:
            if source_line.startswith("![") or source_line.startswith("Appendix candidate:"):
                continue
            line = strip_inline_evidence(source_line)
            if not line.strip():
                continue
            if section_name == "Executive Summary" and line.startswith("- "):
                lines.append(line)
                continue
            if line.startswith("|") and "Evidence ID" in line:
                line = line.replace("Evidence ID", "ID подтверждения")
            lines.append(line)
            cid = chart_id(source_line)
            if cid and cid in BODY_CHARTS:
                chart = chart_lookup[cid]
                lines.append(f"![{chart['Название графика']}]({chart['Путь к изображению']})")
        if section_name == "Executive Summary":
            lines.extend(["", "### Ключевые сигналы для проверки", ""])
            lines.extend(
                markdown_table(
                    ["Блок", "Главный сигнал", "Метрика", "Риск / ограничение", "Что проверить"],
                    build_executive_table_rows(),
                )
            )
        lines.append("")

    lines.extend(["## Appendix / evidence", ""])
    lines.append("Детальная трассировка вынесена из основного текста. Appendix поддерживает записку, но не заменяет executive memo.")
    lines.extend(["", "### Детальная evidence map", ""])
    evidence_rows = []
    for row in evidence_map.head(40).to_dict("records"):
        evidence_rows.append(
            [
                str(row.get("Evidence ID", "")),
                str(row.get("Источник среза", "")),
                str(row.get("Метрика", "")),
                str(row.get("Ограничение", "")),
            ]
        )
    lines.extend(markdown_table(["ID подтверждения", "Источник среза", "Метрика", "Ограничение"], evidence_rows))
    lines.extend(["", "### Appendix charts", ""])
    for chart_id_value in sorted(APPENDIX_CHARTS):
        chart = chart_lookup[chart_id_value]
        lines.append(f"- {chart['Название графика']} — {chart['Подпись']}")
    return "\n".join(lines).rstrip() + "\n"


def set_run_font(run: Any, size: float | None = None, color: str | None = None, bold: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell: Any) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": 80, "start": 120, "bottom": 80, "end": 120}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table: Any, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)
                row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(row.cells[idx])


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.05
    for style_name, size, color in [("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5"), ("Heading 3", 12, "1F4D78")]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(3)


def parse_table(lines: list[str], start: int) -> tuple[list[str], list[list[str]], int]:
    header = [part.strip() for part in lines[start].strip("|").split("|")]
    rows = []
    idx = start + 2
    while idx < len(lines) and lines[idx].startswith("|"):
        rows.append([part.strip() for part in lines[idx].strip("|").split("|")])
        idx += 1
    return header, rows, idx


def is_table_separator(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(cell and set(cell) <= {"-", ":"} for cell in cells)


def add_table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    widths = [1.0, 1.6, 1.2, 1.3, 1.3] if len(header) == 5 else [1.2, 1.7, 1.2, 2.1]
    if len(header) == 7:
        widths = [1.05, 0.9, 1.55, 1.0, 0.75, 1.05, 0.95]
    set_table_width(table, widths)
    for idx, value in enumerate(header):
        cell = table.rows[0].cells[idx]
        cell.text = value
        set_cell_shading(cell, "F2F4F7")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size=8.5, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row[: len(header)]):
            cells[idx].text = value
            for paragraph in cells[idx].paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    set_run_font(run, size=8)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("Управленческая записка YoY/MoM по бюджету")
    set_run_font(r, size=20, color="0B2545", bold=True)
    p = doc.add_paragraph()
    r = p.add_run("Управленческий DOCX-черновик на основе принятого MART, пакета графиков и согласованной версии записки. Аналитика не переписывалась.")
    set_run_font(r, size=10, color="555555")


def add_scope_table(doc: Document) -> None:
    doc.add_heading("Период / scope / статус данных", level=1)
    rows = [
        ["Профиль", "executive_yoy_mom_budget_memo"],
        ["Источник текста", "06_reports/executive_yoy_mom_memo_revised.md"],
        ["Статус данных", "Принятый MART, принятый пакет графиков, QA согласованной версии записки: pass"],
        ["Ограничение", "DOCX не заявляет готовность к публикации и не добавляет новые финансовые числа."],
    ]
    add_table(doc, ["Поле", "Значение"], rows)


def add_chart(doc: Document, chart: dict[str, Any]) -> None:
    image_path = PROJECT_ROOT / str(chart["Путь к изображению"])
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    p.add_run().add_picture(str(image_path), width=Inches(4.9))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(4)
    chart_order = chart.get("Порядок", "")
    prefix = f"График {int(chart_order)}" if pd.notna(chart_order) and str(chart_order) else "График"
    run = cap.add_run(f"{prefix}. {chart['Подпись']}")
    set_run_font(run, size=8.5, color="555555")


def build_docx(markdown: str) -> None:
    doc = Document()
    style_document(doc)
    add_title(doc)
    add_scope_table(doc)
    chart_catalog = pd.read_excel(CHART_CATALOG)
    chart_lookup = {row["ID графика"]: row for row in chart_catalog.to_dict("records")}
    lines = markdown.splitlines()
    idx = 0
    in_appendix = False
    while idx < len(lines):
        line = lines[idx].strip()
        if not line or line.startswith("# "):
            idx += 1
            continue
        if line.startswith("## "):
            heading = line[3:]
            in_appendix = heading == "Appendix / evidence"
            doc.add_heading(heading, level=1)
            idx += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
            idx += 1
            continue
        if line.startswith("|") and idx + 1 < len(lines) and is_table_separator(lines[idx + 1]):
            header, rows, idx = parse_table(lines, idx)
            add_table(doc, header, rows)
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            text = line[2:].replace("**", "")
            run = p.add_run(text)
            set_run_font(run)
            idx += 1
            continue
        if line.startswith("!["):
            cid = chart_id(line)
            if cid and cid in chart_lookup:
                add_chart(doc, chart_lookup[cid])
            idx += 1
            continue
        p = doc.add_paragraph()
        run = p.add_run(line.replace("**", ""))
        set_run_font(run)
        idx += 1
    for cid in sorted(APPENDIX_CHARTS):
        chart = chart_lookup[cid]
        doc.add_heading(str(chart["Название графика"]), level=2)
        add_chart(doc, chart)
    doc.save(DOCX_PATH)


def docx_text() -> str:
    doc = Document(str(DOCX_PATH))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def embedded_images() -> int:
    with zipfile.ZipFile(DOCX_PATH) as zf:
        return len([name for name in zf.namelist() if name.startswith("word/media/")])


def validate(raw_before: dict[str, tuple[int, int]], stage_before: dict[str, tuple[int, int]], marts_before: dict[str, tuple[int, int]], charts_before: dict[str, tuple[int, int]], old_final_before: tuple[int, int] | None) -> dict[str, Any]:
    text = docx_text()
    doc = Document(str(DOCX_PATH))
    forbidden_headers = {"Evidence", "Source slice", "Metric", "Evidence ID"}
    visible_cells = [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    main_text = text.split("Appendix / evidence")[0]
    checks = {
        "docx_exists": DOCX_PATH.exists() and DOCX_PATH.stat().st_size > 0,
        "final_md_exists": FINAL_MD.exists() and FINAL_MD.stat().st_size > 0,
        "english_service_subtitle_removed": "Executive draft from accepted MART" not in text,
        "technical_headers_translated": not any(header in visible_cells for header in forbidden_headers),
        "technical_action_status_humanized": "candidate_only_incomplete_action_fields" not in text,
        "main_body_not_overloaded_with_evidence_tables": "Evidence table:" not in main_text and main_text.count("ID подтверждения") <= 2,
        "appendix_contains_detailed_evidence": "Appendix / evidence" in text and "ID подтверждения" in text and "Источник среза" in text and "Метрика" in text,
        "executive_table_exists": all(header in text for header in ["Блок", "Главный сигнал", "Риск / ограничение", "Что проверить"]),
        "all_12_sections_exist": all(section in text for section in SECTION_ORDER),
        "body_charts_present": embedded_images() >= len(BODY_CHARTS),
        "appendix_charts_present": embedded_images() >= len(BODY_CHARTS) + len(APPENDIX_CHARTS),
        "planning_risk_future_not_execution": "не описывает фактическое исполнение" in text and "будущий бюджетный риск" in text,
        "yoy_mom_separate": text.find("YoY: сдвиг уровня к прошлому году") < text.find("MoM: помесячная динамика и нестабильность"),
        "old_final_memo_not_modified": old_final_before == file_stat(OLD_FINAL_MEMO),
        "old_final_content_not_used": "plan_base_risk" not in text.lower() and "CH_BUDGET_BRIDGE" not in text,
        "raw_untouched": raw_before == snapshot(RAW_DIR),
        "stage_untouched": stage_before == snapshot(STAGE_DIR),
        "mart_untouched": marts_before == snapshot(MARTS_DIR),
        "charts_untouched": charts_before == snapshot(CHARTS_DIR),
    }
    return {
        "created_at": datetime.now(timezone.utc).isoformat(),
        "qa_status": "pass" if all(checks.values()) else "fail",
        "checks": {key: bool(value) for key, value in checks.items()},
        "docx_path": str(DOCX_PATH.relative_to(PROJECT_ROOT)),
        "final_md_path": str(FINAL_MD.relative_to(PROJECT_ROOT)),
        "embedded_image_count": embedded_images(),
        "table_count": len(doc.tables),
        "visual_render_qa": "pending",
    }


def resolve_soffice() -> str | None:
    env_bin = os.environ.get("SOFFICE_BIN")
    if env_bin:
        return env_bin
    path_bin = shutil.which("soffice") or shutil.which("libreoffice")
    if path_bin:
        return path_bin
    macos_bin = Path("/Applications/LibreOffice.app/Contents/MacOS/soffice")
    if macos_bin.exists():
        return str(macos_bin)
    return None


def update_visual_blocked() -> None:
    soffice = resolve_soffice()
    rendered = list(RENDER_DIR.glob("*.png")) + list(RENDER_DIR.glob("*.pdf")) if RENDER_DIR.exists() else []
    if soffice or rendered:
        return
    qa = {
        "created_at": datetime.now(timezone.utc).isoformat(),
        "visual_qa_status": "blocked_soffice_unavailable",
        "qa_status": "blocked",
        "reason": "LibreOffice soffice executable is unavailable; polished DOCX could not be rendered to PNG/PDF.",
        "input_docx": str(DOCX_PATH.relative_to(PROJECT_ROOT)),
        "render_output_dir": str(RENDER_DIR.relative_to(PROJECT_ROOT)),
        "render_artifacts": [],
        "layout_findings": ["Visual layout could not be inspected because render conversion did not run."],
        "blocking_issues": ["blocked_soffice_unavailable"],
        "recommendation": "ready for internal structural review only; not ready for external sharing until visual render QA passes",
        "production_readiness_claimed": False,
    }
    VISUAL_QA.write_text(json.dumps(qa, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    VISUAL_QA_SUMMARY.write_text(
        "\n".join(
            [
                "# Executive YoY/MoM Visual QA Summary",
                "",
                "visual_qa_status: blocked_soffice_unavailable",
                "qa_status: blocked",
                "reason: LibreOffice soffice executable is unavailable; polished DOCX could not be rendered to PNG/PDF.",
                "",
                "## Recommendation",
                "- Ready for internal structural review only; not ready for external sharing until visual render QA passes.",
            ]
        )
        + "\n",
        encoding="utf-8",
    )


def main() -> None:
    log_progress(memo_profile=MEMO_PROFILE, depth_mode=DEPTH_MODE, stage="docx_presentation_polish", status="start")
    for path in [REVISED_MD, DOCX_QA, EVIDENCE_MAP, CHART_CATALOG]:
        if not path.exists():
            raise FileNotFoundError(f"Missing input: {path}")
    raw_before = snapshot(RAW_DIR)
    stage_before = snapshot(STAGE_DIR)
    marts_before = snapshot(MARTS_DIR)
    charts_before = snapshot(CHARTS_DIR)
    old_final_before = file_stat(OLD_FINAL_MEMO)
    log_progress(memo_profile=MEMO_PROFILE, depth_mode=DEPTH_MODE, stage="final_markdown_polish", status="start")
    markdown = build_final_markdown()
    FINAL_MD.write_text(markdown, encoding="utf-8")
    log_progress(memo_profile=MEMO_PROFILE, depth_mode=DEPTH_MODE, stage="final_markdown_polish", status="done")
    log_progress(memo_profile=MEMO_PROFILE, depth_mode=DEPTH_MODE, stage="docx_rebuild_after_polish", status="start")
    build_docx(markdown)
    log_progress(memo_profile=MEMO_PROFILE, depth_mode=DEPTH_MODE, stage="docx_rebuild_after_polish", status="done")
    log_progress(memo_profile=MEMO_PROFILE, depth_mode=DEPTH_MODE, stage="docx_polish_qa", status="start")
    qa = validate(raw_before, stage_before, marts_before, charts_before, old_final_before)
    DOCX_QA.write_text(json.dumps(qa, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    update_visual_blocked()
    log_progress(memo_profile=MEMO_PROFILE, depth_mode=DEPTH_MODE, stage="docx_presentation_polish", status=qa["qa_status"])
    print(json.dumps({"qa_status": qa["qa_status"], "images": qa["embedded_image_count"], "tables": qa["table_count"]}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
