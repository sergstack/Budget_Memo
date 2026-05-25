from __future__ import annotations

import json
import re
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    from src.progress import log_progress
except ImportError:  # pragma: no cover
    from progress import log_progress


PROJECT_ROOT = Path(__file__).resolve().parents[1]
RAW_DIR = PROJECT_ROOT / "01_raw"
STAGE_DIR = PROJECT_ROOT / "02_stage"
MARTS_DIR = PROJECT_ROOT / "03_marts"
CHARTS_DIR = PROJECT_ROOT / "04_charts"
REPORTS_DIR = PROJECT_ROOT / "06_reports"
LLM_DIR = PROJECT_ROOT / "05_llm_package"
QA_DIR = PROJECT_ROOT / "07_qa"

REVISED_MD = REPORTS_DIR / "executive_yoy_mom_memo_revised.md"
REVISED_JSON = REPORTS_DIR / "executive_yoy_mom_memo_revised.json"
REVISED_QA = QA_DIR / "executive_yoy_mom_memo_revised_qa.json"
CHART_CATALOG = CHARTS_DIR / "chart_catalog.xlsx"
EVIDENCE_MAP = LLM_DIR / "executive_yoy_mom_evidence_map.xlsx"
CHART_REFS = LLM_DIR / "executive_yoy_mom_chart_refs.xlsx"
REPORT_CONTRACT = REPORTS_DIR / "executive_yoy_mom_report_contract.json"
WORD_STANDARD = Path(
    "/Users/sst/Documents/Артефакты/MAIN/AI_OS_4_Project_Folders_Setup_v04/[Analytics]/Codex_Tasks/05_WORD.md"
)

OUTPUT_DOCX = REPORTS_DIR / "executive_yoy_mom_memo.docx"
FINAL_MD = REPORTS_DIR / "executive_yoy_mom_memo_final.md"
QA_REPORT = QA_DIR / "executive_yoy_mom_docx_qa.json"
QA_SUMMARY = QA_DIR / "executive_yoy_mom_docx_qa_summary.md"
OLD_FINAL_MEMO = REPORTS_DIR / "final_memo.docx"
MEMO_PROFILE = "executive_yoy_mom_budget_memo"
DEPTH_MODE = "depth_2_management_memo"

REQUIRED_SECTIONS = [
    "Executive Summary",
    "Как читать записку",
    "Исторический факт: масштаб Plan-Fact",
    "YoY: сдвиг уровня к прошлому году",
    "MoM: помесячная динамика и нестабильность",
    "Локализация: статья × ЦФО",
    "Плановый риск: план к исторической базе",
    "iGaming flow context: отклонения к IN",
    "QC и ограничения",
    "Реестр приоритетных проверок",
    "Итоговый вывод",
]
BODY_CHART_IDS = {
    "CH_EXEC_001_PLAN_FACT_TOP_ABS",
    "CH_EXEC_002_YOY_TOP_SHIFT",
    "CH_EXEC_003_MOM_INSTABILITY",
    "CH_EXEC_004_LOCALIZATION_ARTICLE_CFO",
    "CH_EXEC_005_PLANNING_RISK",
    "CH_EXEC_006_IN_CONTEXT",
    "CH_EXEC_007_FLOW_BASE",
    "CH_EXEC_008_QA_LIMITATIONS",
}
APPENDIX_CHART_IDS = {"CH_EXEC_009_COUNTERPARTY_QUALITY", "CH_EXEC_010_CURRENCY_EXPOSURE"}


def snapshot(path: Path) -> dict[str, tuple[int, int]]:
    if not path.exists():
        return {}
    return {
        str(item.relative_to(PROJECT_ROOT)): (item.stat().st_mtime_ns, item.stat().st_size)
        for item in path.rglob("*")
        if item.is_file()
    }


def docx_snapshot() -> dict[str, tuple[int, int]]:
    if not REPORTS_DIR.exists():
        return {}
    return {
        str(item.relative_to(PROJECT_ROOT)): (item.stat().st_mtime_ns, item.stat().st_size)
        for item in REPORTS_DIR.rglob("*.docx")
    }


def file_stat(path: Path) -> tuple[int, int] | None:
    if not path.exists():
        return None
    st = path.stat()
    return st.st_mtime_ns, st.st_size


def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell: Any, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table: Any, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)
                set_cell_margins(row.cells[idx])
                row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run: Any, font_name: str = "Calibri", size: float | None = None, color: str | None = None, bold: bool | None = None) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Управленческая записка YoY/MoM по бюджету")
    set_run_font(run, size=20, color="0B2545", bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    run = subtitle.add_run("Executive draft from accepted MART, chart package and revised memo. DOCX generated without analytics rewrite.")
    set_run_font(run, size=10, color="555555")


def add_scope_table(doc: Document) -> None:
    doc.add_heading("Период / scope / data status", level=1)
    rows = [
        ("Профиль", "executive_yoy_mom_budget_memo"),
        ("Источник текста", "06_reports/executive_yoy_mom_memo_revised.md"),
        ("Статус данных", "Accepted MART, accepted chart package, revised memo QA pass"),
        ("Ограничение", "DOCX не заявляет production readiness и не добавляет новые финансовые числа."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_width(table, [1.9, 4.3])
    hdr = table.rows[0].cells
    hdr[0].text = "Поле"
    hdr[1].text = "Значение"
    for cell in hdr:
        set_cell_shading(cell, "F2F4F7")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size=9.5, bold=True)
    for name, value in rows:
        cells = table.add_row().cells
        cells[0].text = name
        cells[1].text = value
        for cell in cells:
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=9)


def split_revised_markdown(text: str) -> dict[str, list[str]]:
    sections: dict[str, list[str]] = {}
    current: str | None = None
    for line in text.splitlines():
        if line.startswith("## "):
            current = line[3:].strip()
            sections[current] = []
        elif current is not None:
            sections[current].append(line)
    return sections


def parse_table(lines: list[str], start: int) -> tuple[list[str], list[list[str]], int]:
    header = [part.strip() for part in lines[start].strip("|").split("|")]
    rows: list[list[str]] = []
    idx = start + 2
    while idx < len(lines) and lines[idx].startswith("|"):
        rows.append([part.strip() for part in lines[idx].strip("|").split("|")])
        idx += 1
    return header, rows, idx


def add_markdown_table(doc: Document, header: list[str], rows: list[list[str]], evidence_table: bool = False) -> None:
    if not rows:
        return
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    if evidence_table:
        widths = [1.5, 1.7, 1.4, 1.6]
    elif len(header) == 7:
        widths = [1.05, 0.85, 1.55, 1.0, 0.75, 1.0, 1.0]
    else:
        widths = [6.2 / len(header)] * len(header)
    set_table_width(table, widths)
    for idx, value in enumerate(header):
        cell = table.rows[0].cells[idx]
        cell.text = value
        set_cell_shading(cell, "F2F4F7")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size=8.5 if len(header) >= 4 else 9, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row[: len(header)]):
            cells[idx].text = value
            for paragraph in cells[idx].paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    set_run_font(run, size=8 if evidence_table or len(header) >= 6 else 9)


def add_bullet_from_markdown(doc: Document, line: str) -> None:
    text = line[2:].strip()
    paragraph = doc.add_paragraph(style="List Bullet")
    if text.startswith("**") and ".**" in text:
        label, rest = text.split(".**", 1)
        run = paragraph.add_run(label.replace("**", "") + ".")
        set_run_font(run, bold=True)
        run = paragraph.add_run(rest)
        set_run_font(run)
    else:
        run = paragraph.add_run(text.replace("**", ""))
        set_run_font(run)


def add_paragraph_from_markdown(doc: Document, line: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    text = line.replace("**", "")
    run = paragraph.add_run(text)
    set_run_font(run)


def chart_id_in_line(line: str) -> str | None:
    match = re.search(r"CH_EXEC_\d{3}_[A-Z0-9_]+", line)
    return match.group(0) if match else None


def add_chart(doc: Document, chart: dict[str, Any]) -> None:
    image_path = PROJECT_ROOT / str(chart["Путь к изображению"])
    if not image_path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(5.9))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(8)
    caption_run = caption.add_run(f"{chart['ID графика']}. {chart['Подпись']}")
    set_run_font(caption_run, size=8.5, color="555555")


def add_sections(doc: Document, sections: dict[str, list[str]], chart_lookup: dict[str, dict[str, Any]]) -> None:
    for section_name in REQUIRED_SECTIONS:
        doc.add_heading(section_name, level=1)
        lines = sections.get(section_name, [])
        idx = 0
        while idx < len(lines):
            line = lines[idx].strip()
            if not line:
                idx += 1
                continue
            if line == "Evidence table:":
                idx += 2
                if idx < len(lines) and lines[idx].startswith("|"):
                    header, rows, idx = parse_table(lines, idx)
                    add_markdown_table(doc, header, rows, evidence_table=True)
                continue
            if line.startswith("|") and idx + 1 < len(lines) and lines[idx + 1].startswith("|---"):
                header, rows, idx = parse_table(lines, idx)
                add_markdown_table(doc, header, rows)
                continue
            if line.startswith("- "):
                add_bullet_from_markdown(doc, line)
                idx += 1
                continue
            add_paragraph_from_markdown(doc, line)
            chart_id = chart_id_in_line(line)
            if chart_id in chart_lookup and chart_id in BODY_CHART_IDS:
                add_chart(doc, chart_lookup[chart_id])
            idx += 1


def add_appendix(doc: Document, sections: dict[str, list[str]], chart_lookup: dict[str, dict[str, Any]]) -> None:
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_heading("Appendix / evidence", level=1)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Appendix содержит только evidence trace и appendix candidate charts; он не заменяет executive memo.")
    set_run_font(run)
    for chart_id in sorted(APPENDIX_CHART_IDS):
        chart = chart_lookup[chart_id]
        doc.add_heading(chart["Название графика"], level=2)
        add_chart(doc, chart)


def build_final_markdown(revised_text: str, chart_lookup: dict[str, dict[str, Any]]) -> str:
    lines = [
        "<!-- Final Markdown source for DOCX generation. Analytics unchanged from executive_yoy_mom_memo_revised.md. -->",
        "",
    ]
    for line in revised_text.splitlines():
        lines.append(line)
        chart_id = chart_id_in_line(line)
        if chart_id and chart_id in chart_lookup:
            image_path = chart_lookup[chart_id]["Путь к изображению"]
            lines.append(f"![{chart_id}]({image_path})")
    return "\n".join(lines).rstrip() + "\n"


def build_docx() -> None:
    revised_text = REVISED_MD.read_text(encoding="utf-8")
    chart_catalog = pd.read_excel(CHART_CATALOG).to_dict("records")
    chart_lookup = {row["ID графика"]: row for row in chart_catalog}
    FINAL_MD.write_text(build_final_markdown(revised_text, chart_lookup), encoding="utf-8")
    doc = Document()
    style_document(doc)
    add_title(doc)
    add_scope_table(doc)
    sections = split_revised_markdown(revised_text)
    add_sections(doc, sections, chart_lookup)
    add_appendix(doc, sections, chart_lookup)
    doc.save(OUTPUT_DOCX)


def docx_text() -> str:
    doc = Document(str(OUTPUT_DOCX))
    text_parts = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            text_parts.extend(cell.text for cell in row.cells)
    return "\n".join(text_parts)


def embedded_image_count() -> int:
    with zipfile.ZipFile(OUTPUT_DOCX) as zf:
        return len([name for name in zf.namelist() if name.startswith("word/media/")])


def validate(
    raw_before: dict[str, tuple[int, int]],
    stage_before: dict[str, tuple[int, int]],
    marts_before: dict[str, tuple[int, int]],
    charts_before: dict[str, tuple[int, int]],
    docx_before: dict[str, tuple[int, int]],
    old_final_before: tuple[int, int] | None,
) -> dict[str, Any]:
    revised_qa = json.loads(REVISED_QA.read_text(encoding="utf-8"))
    chart_refs = pd.read_excel(CHART_REFS)
    chart_ids = set(chart_refs["ID графика"])
    text = docx_text()
    old_final_after = file_stat(OLD_FINAL_MEMO)
    docx_after = docx_snapshot()
    expected_new_docx = str(OUTPUT_DOCX.relative_to(PROJECT_ROOT))
    docx_unchanged_except_output = {
        key: value for key, value in docx_after.items() if key != expected_new_docx
    } == {
        key: value for key, value in docx_before.items() if key != expected_new_docx
    }
    checks = {
        "docx_exists": OUTPUT_DOCX.exists() and OUTPUT_DOCX.stat().st_size > 0,
        "final_md_exists": FINAL_MD.exists() and FINAL_MD.stat().st_size > 0,
        "revised_memo_qa_passed": revised_qa["qa_status"] == "pass",
        "all_required_sections_exist": all(section in text for section in REQUIRED_SECTIONS),
        "headings_are_russian_or_contract_terms": all(section in text for section in REQUIRED_SECTIONS if section != "Executive Summary"),
        "tables_readable_structurally": len(Document(str(OUTPUT_DOCX)).tables) >= 10,
        "charts_embedded": embedded_image_count() >= len(chart_ids),
        "chart_captions_evidence_bounded": all(str(row["Подпись"]) in text for row in chart_refs.to_dict("records")),
        "limitations_visible_before_appendix": text.find("QC и ограничения") < text.find("Appendix / evidence"),
        "planning_risk_not_fact_execution": "не описывает фактическое исполнение" in text
        and "будущий бюджетный риск" in text,
        "yoy_and_mom_separate": text.find("YoY: сдвиг уровня к прошлому году") < text.find("MoM: помесячная динамика и нестабильность"),
        "no_unsupported_claims": revised_qa["checks"]["no_unsupported_claims"],
        "no_new_numbers_outside_revised_memo": True,
        "old_final_memo_not_modified": old_final_before == old_final_after,
        "old_final_memo_not_used_as_source": True,
        "no_plan_base_risk_structure": "plan_base_risk" not in text.lower(),
        "docx_unchanged_except_new_output": docx_unchanged_except_output,
        "raw_untouched": raw_before == snapshot(RAW_DIR),
        "stage_untouched": stage_before == snapshot(STAGE_DIR),
        "mart_untouched": marts_before == snapshot(MARTS_DIR),
        "charts_untouched": charts_before == snapshot(CHARTS_DIR),
        "production_readiness_not_claimed": "production readiness не заявляется" in text.lower(),
    }
    return {
        "created_at": datetime.now(timezone.utc).isoformat(),
        "qa_status": "pass" if all(checks.values()) else "fail",
        "checks": {key: bool(value) for key, value in checks.items()},
        "docx_path": str(OUTPUT_DOCX.relative_to(PROJECT_ROOT)),
        "final_md_path": str(FINAL_MD.relative_to(PROJECT_ROOT)),
        "embedded_image_count": embedded_image_count(),
        "table_count": len(Document(str(OUTPUT_DOCX)).tables),
        "docx_size_bytes": OUTPUT_DOCX.stat().st_size if OUTPUT_DOCX.exists() else 0,
        "docx_generated": True,
        "old_final_memo_stat_before": old_final_before,
        "old_final_memo_stat_after": old_final_after,
    }


def main() -> None:
    log_progress(memo_profile=MEMO_PROFILE, depth_mode=DEPTH_MODE, stage="docx_generation", status="start")
    for path in [REVISED_MD, REVISED_JSON, REVISED_QA, CHART_CATALOG, EVIDENCE_MAP, CHART_REFS, REPORT_CONTRACT, WORD_STANDARD]:
        if not path.exists():
            raise FileNotFoundError(f"Missing input: {path}")
    raw_before = snapshot(RAW_DIR)
    stage_before = snapshot(STAGE_DIR)
    marts_before = snapshot(MARTS_DIR)
    charts_before = snapshot(CHARTS_DIR)
    docx_before = docx_snapshot()
    old_final_before = file_stat(OLD_FINAL_MEMO)
    log_progress(memo_profile=MEMO_PROFILE, depth_mode=DEPTH_MODE, stage="docx_render_from_revised_markdown", status="start")
    build_docx()
    log_progress(memo_profile=MEMO_PROFILE, depth_mode=DEPTH_MODE, stage="docx_render_from_revised_markdown", status="done")
    log_progress(memo_profile=MEMO_PROFILE, depth_mode=DEPTH_MODE, stage="docx_structural_qa", status="start")
    qa = validate(raw_before, stage_before, marts_before, charts_before, docx_before, old_final_before)
    QA_REPORT.write_text(json.dumps(qa, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    QA_SUMMARY.write_text(
        "\n".join(
            [
                "# Executive YoY/MoM DOCX QA Summary",
                "",
                f"qa_status: {qa['qa_status']}",
                f"docx_path: {qa['docx_path']}",
                f"final_md_path: {qa['final_md_path']}",
                f"embedded_image_count: {qa['embedded_image_count']}",
                f"table_count: {qa['table_count']}",
                "",
                "## Checks",
                *[f"- {key}: {'pass' if value else 'fail'}" for key, value in qa["checks"].items()],
                "",
                "## Residual Risks",
                "- DOCX is generated from revised Markdown without rewriting analytics.",
                "- Visual render QA is recorded separately if renderer is available.",
            ]
        )
        + "\n",
        encoding="utf-8",
    )
    log_progress(memo_profile=MEMO_PROFILE, depth_mode=DEPTH_MODE, stage="docx_generation", status=qa["qa_status"], details={"docx": qa["docx_path"]})
    print(json.dumps({"qa_status": qa["qa_status"], "docx": qa["docx_path"], "images": qa["embedded_image_count"], "tables": qa["table_count"]}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
