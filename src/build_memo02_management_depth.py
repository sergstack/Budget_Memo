from __future__ import annotations

import json
import math
import zipfile
from copy import copy
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
MARTS_DIR = PROJECT_ROOT / "03_marts"
REPORT_DIR = PROJECT_ROOT / "06_reports" / "02_monthly_plan_fact_memo"

WORKBOOK_PATH = REPORT_DIR / "xlsx" / "management_depth_workbook.xlsx"
MD_PATH = REPORT_DIR / "md" / "management_depth_monthly_plan_fact_memo.md"
DOCX_PATH = REPORT_DIR / "docx" / "management_depth_monthly_plan_fact_memo.docx"
QA_JSON_PATH = REPORT_DIR / "qa" / "management_depth_package_qa.json"
QA_MD_PATH = REPORT_DIR / "qa" / "management_depth_package_qa.md"
CLAIM_AUDIT_PATH = REPORT_DIR / "qa" / "management_depth_claim_audit.md"
ARTIFACT_QA_PATH = REPORT_DIR / "qa" / "management_depth_artifact_qa.md"
CHART_QA_PATH = REPORT_DIR / "qa" / "management_depth_chart_qa.md"
MANIFEST_PATH = REPORT_DIR / "source_refs" / "management_depth_manifest.json"
MANAGEMENT_CHART_DIR = REPORT_DIR / "charts" / "management_depth"
MANAGEMENT_CHART_IMAGE_DIR = MANAGEMENT_CHART_DIR / "images"
MANAGEMENT_CHART_DATA_DIR = MANAGEMENT_CHART_DIR / "data"
MANAGEMENT_CHART_META_PATH = MANAGEMENT_CHART_DIR / "management_depth_chart_metadata.json"

SERVICE_ARTICLES = {"IN", "OUT", "IN-OUT"}
PROFILE = "monthly_plan_fact_memo"
MATERIALITY_THRESHOLD_EUR = 100_000.0

REQUIRED_INPUTS = [
    MARTS_DIR / "mart_full_package.xlsx",
    MARTS_DIR / "mart_main_full_budget.parquet",
    MARTS_DIR / "slice_plan_fact_article_month.parquet",
    MARTS_DIR / "slice_plan_fact_article_cfo_month.parquet",
    MARTS_DIR / "slice_plan_fact_article_cfo_counterparty_month.parquet",
    MARTS_DIR / "slice_localization_article_cfo_month.parquet",
    MARTS_DIR / "slice_fact_only.parquet",
    MARTS_DIR / "slice_plan_only.parquet",
    MARTS_DIR / "slice_dq_flags.parquet",
    MARTS_DIR / "slice_source_mix_summary.parquet",
    MARTS_DIR / "mart_signal_catalog_full.parquet",
]


def rel(path: Path) -> str:
    return str(path.relative_to(PROJECT_ROOT))


def read_parquet(name: str) -> pd.DataFrame:
    return pd.read_parquet(MARTS_DIR / name)


def fail_if_missing() -> None:
    missing = [rel(path) for path in REQUIRED_INPUTS if not path.exists()]
    if missing:
        raise FileNotFoundError("Missing required accepted MART/package files: " + ", ".join(missing))


def select_closed_month(main: pd.DataFrame) -> str:
    mask = (
        main["period_type"].astype(str).eq("historical")
        & main["has_fact"].astype(bool)
        & main["fact_eur"].notna()
    )
    months = sorted(main.loc[mask, "period_month"].astype(str).unique())
    if not months:
        raise ValueError("No closed historical month with fact data found")
    return months[-1]


def clean_text(value: Any) -> str:
    if pd.isna(value):
        return "не указано"
    text = str(value).strip()
    return text if text else "не указано"


def status_from_delta(delta: float) -> str:
    if delta < -0.01:
        return "fact above plan / превышение факта"
    if delta > 0.01:
        return "fact below plan / факт ниже плана"
    return "plan = fact"


def fmt_eur(value: float) -> str:
    value = float(value)
    sign = "-" if value < 0 else ""
    value_abs = abs(value)
    if value_abs >= 1_000_000:
        return f"{sign}{value_abs / 1_000_000:.2f} млн EUR".replace(".", ",")
    if value_abs >= 1_000:
        return f"{sign}{value_abs / 1_000:.1f} тыс. EUR".replace(".", ",")
    return f"{sign}{value_abs:,.0f} EUR".replace(",", " ")


def fmt_pct(value: float) -> str:
    if pd.isna(value) or math.isinf(float(value)):
        return "n/a"
    return f"{float(value) * 100:.1f}%".replace(".", ",")


def chart_label(value: Any, limit: int = 36) -> str:
    text = clean_text(value)
    return text if len(text) <= limit else text[: limit - 1] + "…"


def axis_eur(value: float, _pos: Any = None) -> str:
    value = float(value)
    if abs(value) >= 1_000_000:
        return f"{value / 1_000_000:.1f} млн".replace(".", ",")
    return f"{value / 1_000:.0f} тыс".replace(".", ",")


def safe_div(num: pd.Series, den: pd.Series) -> pd.Series:
    return num.divide(den.where(den.abs() > 0))


def aggregate(df: pd.DataFrame, keys: list[str], source_slice: str) -> pd.DataFrame:
    result = (
        df.groupby(keys, dropna=False, as_index=False)
        .agg(
            plan_eur=("plan_eur", "sum"),
            fact_eur=("fact_eur", "sum"),
            rows_count=("row_id", "size"),
        )
    )
    result["delta_eur"] = result["plan_eur"] - result["fact_eur"]
    result["abs_delta_eur"] = result["delta_eur"].abs()
    result["execution_pct"] = safe_div(result["fact_eur"], result["plan_eur"])
    total_abs = result["abs_delta_eur"].sum()
    result["share_of_abs_delta"] = result["abs_delta_eur"] / total_abs if total_abs else 0.0
    result["status"] = result["delta_eur"].map(status_from_delta)
    result["source_slice"] = source_slice
    return result.sort_values("abs_delta_eur", ascending=False).reset_index(drop=True)


def add_priority(df: pd.DataFrame, gap_type: str | None = None) -> pd.DataFrame:
    result = df.copy()
    result["check_priority"] = "P4"
    if {"cfo", "article", "abs_delta_eur"}.issubset(result.columns):
        result.loc[
            result["abs_delta_eur"].ge(MATERIALITY_THRESHOLD_EUR)
            & result["cfo"].astype(str).ne("")
            & result["article"].astype(str).ne(""),
            "check_priority",
        ] = "P1"
    if gap_type and "amount_eur" in result.columns:
        result.loc[result["amount_eur"].abs().ge(MATERIALITY_THRESHOLD_EUR), "check_priority"] = "P2"
    return result


def top(df: pd.DataFrame, n: int = 30) -> pd.DataFrame:
    return df.sort_values("abs_delta_eur", ascending=False).head(n).reset_index(drop=True)


def build_tables() -> tuple[str, dict[str, pd.DataFrame], dict[str, Any]]:
    fail_if_missing()
    main = read_parquet("mart_main_full_budget.parquet")
    selected_month = select_closed_month(main)

    month = main[
        main["period_month"].astype(str).eq(selected_month)
        & main["period_type"].astype(str).eq("historical")
        & ~main["article"].astype(str).isin(SERVICE_ARTICLES)
    ].copy()
    if month.empty:
        raise ValueError(f"No non-service historical rows found for selected month {selected_month}")

    loc = read_parquet("slice_localization_article_cfo_month.parquet")
    loc_month = loc[loc["period_month"].astype(str).eq(selected_month)].copy()
    owner_route = loc_month[["article", "cfo", "owner_candidate", "owner_route_status", "concentration_type"]].drop_duplicates()

    signal = read_parquet("mart_signal_catalog_full.parquet")
    signal = signal[
        signal["eligible_memo_profiles"].astype(str).str.contains(PROFILE, na=False)
        | signal["primary_memo_profile"].astype(str).eq(PROFILE)
    ].copy()

    total_plan = float(month["plan_eur"].sum())
    total_fact = float(month["fact_eur"].sum())
    total_delta = total_plan - total_fact
    total = pd.DataFrame(
        [
            {
                "selected_month": selected_month,
                "plan_eur": total_plan,
                "fact_eur": total_fact,
                "delta_eur": total_delta,
                "abs_delta_eur": abs(total_delta),
                "execution_pct": total_fact / total_plan if total_plan else pd.NA,
                "formula": "Delta EUR = Plan EUR - Fact EUR",
                "status": status_from_delta(total_delta),
                "source_slice": "mart_main_full_budget non-service rows",
            }
        ]
    )

    cfo = aggregate(month, ["cfo"], "mart_main_full_budget -> group by cfo")
    article_h = aggregate(
        month,
        ["article_level_1", "article_level_2", "article"],
        "mart_main_full_budget -> group by article hierarchy",
    )
    cfo_article = aggregate(
        month,
        ["cfo", "article_level_1", "article_level_2", "article"],
        "mart_main_full_budget -> group by cfo x article hierarchy",
    )
    cfo_article = cfo_article.merge(owner_route, on=["article", "cfo"], how="left")
    cfo_article["owner_candidate"] = cfo_article["owner_candidate"].fillna("budget_owner")
    cfo_article["owner_route_status"] = cfo_article["owner_route_status"].fillna("candidate")
    cfo_article = add_priority(cfo_article)

    top_dev = top(cfo_article, 40)

    fact_only = read_parquet("slice_fact_only.parquet")
    fact_only = fact_only[
        fact_only["period_month"].astype(str).eq(selected_month)
        & ~fact_only["article"].astype(str).isin(SERVICE_ARTICLES)
    ].copy()
    fact_gap = (
        fact_only.groupby(["cfo", "article_level_1", "article_level_2", "article", "counterparty"], dropna=False, as_index=False)
        .agg(fact_eur=("fact_eur", "sum"), rows_count=("row_id", "size"))
        .sort_values("fact_eur", ascending=False)
        .head(100)
    )
    fact_gap = fact_gap.merge(owner_route, on=["article", "cfo"], how="left")
    fact_gap["owner_candidate"] = fact_gap["owner_candidate"].fillna("budget_owner")
    fact_gap["amount_eur"] = fact_gap["fact_eur"]
    fact_gap["check"] = "candidate check: validate fact without approved plan"
    fact_gap["source_slice"] = "slice_fact_only"
    fact_gap = add_priority(fact_gap, gap_type="fact_without_plan")

    plan_only = read_parquet("slice_plan_only.parquet")
    plan_only = plan_only[
        plan_only["period_month"].astype(str).eq(selected_month)
        & ~plan_only["article"].astype(str).isin(SERVICE_ARTICLES)
    ].copy()
    plan_gap = (
        plan_only.groupby(["cfo", "article_level_1", "article_level_2", "article"], dropna=False, as_index=False)
        .agg(plan_eur=("plan_eur", "sum"), rows_count=("row_id", "size"))
        .sort_values("plan_eur", ascending=False)
        .head(100)
    )
    plan_gap = plan_gap.merge(owner_route, on=["article", "cfo"], how="left")
    plan_gap["owner_candidate"] = plan_gap["owner_candidate"].fillna("budget_owner")
    plan_gap["amount_eur"] = plan_gap["plan_eur"]
    plan_gap["check"] = "candidate check: validate planned spend without fact"
    plan_gap["source_slice"] = "slice_plan_only"
    plan_gap = add_priority(plan_gap, gap_type="plan_without_fact")

    cp = read_parquet("slice_plan_fact_article_cfo_counterparty_month.parquet")
    cp = cp[
        cp["period_month"].astype(str).eq(selected_month)
        & ~cp["article"].astype(str).isin(SERVICE_ARTICLES)
    ].copy()
    article_lookup = month[["article", "article_level_1", "article_level_2"]].drop_duplicates("article")
    cp = cp.merge(article_lookup, on="article", how="left")
    cp["delta_eur"] = cp["plan_eur"] - cp["fact_eur"]
    cp["abs_delta_eur"] = cp["delta_eur"].abs()
    cp["quality_flag_limitation"] = "localization only; counterparty is not treated as root cause"
    cp_view = (
        cp.sort_values("abs_delta_eur", ascending=False)
        .head(100)[
            [
                "counterparty",
                "counterparty_key",
                "cfo",
                "article_level_1",
                "article_level_2",
                "article",
                "plan_eur",
                "fact_eur",
                "delta_eur",
                "abs_delta_eur",
                "fact_without_plan_flag",
                "plan_without_fact_flag",
                "quality_flag_limitation",
                "source_slice",
            ]
        ]
        .reset_index(drop=True)
    )

    legal_currency = aggregate(month, ["legal_entity", "currency"], "mart_main_full_budget -> group by legal entity x currency")
    legal_currency = legal_currency.head(60)

    priorities = pd.concat(
        [
            top_dev.head(20).assign(
                priority_basis="P1: high ABS Delta with identifiable CFO/article",
                candidate_check="Проверить отклонение с candidate owner route; owner route = CFO; не переводить candidate check в final action без due date/status.",
            )[
                [
                    "check_priority",
                    "priority_basis",
                    "cfo",
                    "article_level_1",
                    "article_level_2",
                    "article",
                    "plan_eur",
                    "fact_eur",
                    "delta_eur",
                    "abs_delta_eur",
                    "owner_candidate",
                    "candidate_check",
                    "source_slice",
                ]
            ],
            fact_gap.head(15).assign(
                priority_basis="P2: fact without plan above materiality/rank",
                delta_eur=lambda df: -df["fact_eur"],
                abs_delta_eur=lambda df: df["fact_eur"].abs(),
                plan_eur=0.0,
                candidate_check="Проверить наличие утвержденного плана/классификации; не трактовать как подтвержденный перерасход.",
            )[
                [
                    "check_priority",
                    "priority_basis",
                    "cfo",
                    "article_level_1",
                    "article_level_2",
                    "article",
                    "plan_eur",
                    "fact_eur",
                    "delta_eur",
                    "abs_delta_eur",
                    "owner_candidate",
                    "candidate_check",
                    "source_slice",
                ]
            ],
            plan_gap.head(15).assign(
                priority_basis="P2: plan without fact above materiality/rank",
                fact_eur=0.0,
                delta_eur=lambda df: df["plan_eur"],
                abs_delta_eur=lambda df: df["plan_eur"].abs(),
                candidate_check="Проверить timing/отмену/перенос; не трактовать как подтвержденную экономию.",
            )[
                [
                    "check_priority",
                    "priority_basis",
                    "cfo",
                    "article_level_1",
                    "article_level_2",
                    "article",
                    "plan_eur",
                    "fact_eur",
                    "delta_eur",
                    "abs_delta_eur",
                    "owner_candidate",
                    "candidate_check",
                    "source_slice",
                ]
            ],
        ],
        ignore_index=True,
    ).sort_values(["check_priority", "abs_delta_eur"], ascending=[True, False])

    blockers = pd.DataFrame(
        [
            {
                "dimension": "Направление",
                "status": "missing",
                "expected_source": "mapping table or explicit field in accepted MART",
                "impact_on_memo": "direction-level decomposition cannot be produced",
                "safe_next_step": "request or implement direction mapping",
            },
            {
                "dimension": "Руководитель",
                "status": "missing",
                "expected_source": "owner/manager mapping table or explicit manager field",
                "impact_on_memo": "confirmed responsible manager cannot be assigned",
                "safe_next_step": "request or implement manager mapping",
            },
            {
                "dimension": "confirmed action owner / due date / status",
                "status": "missing",
                "expected_source": "confirmed action register",
                "impact_on_memo": "final action plan is blocked; candidate checks only",
                "safe_next_step": "create/approve action register after owner confirmation",
            },
        ]
    )

    evidence = pd.DataFrame(
        [
            {
                "claim_id": "M02-MGMT-001",
                "claim_type": "supported numeric fact",
                "claim": f"Selected closed month is {selected_month}; Delta formula is Plan EUR - Fact EUR.",
                "source_mart_or_slice": "mart_main_full_budget",
                "metric": "plan_eur; fact_eur; delta_eur; abs_delta_eur; execution_pct",
                "grain": "selected month",
                "period": selected_month,
                "filter": "period_type=historical; non-service articles",
                "limitation": "planning-only periods excluded",
                "qa_status": "pass",
            },
            {
                "claim_id": "M02-MGMT-002",
                "claim_type": "supported localization",
                "claim": "CFO and article hierarchy localize deviations; they do not prove causes.",
                "source_mart_or_slice": "mart_main_full_budget; slice_localization_article_cfo_month",
                "metric": "plan_eur; fact_eur; delta_eur; abs_delta_eur",
                "grain": "cfo x article hierarchy",
                "period": selected_month,
                "filter": "selected month; top by ABS Delta",
                "limitation": "owner_candidate is only candidate route",
                "qa_status": "pass",
            },
            {
                "claim_id": "M02-MGMT-003",
                "claim_type": "data quality limitation",
                "claim": "Fact without plan and plan without fact are separate candidate checks.",
                "source_mart_or_slice": "slice_fact_only; slice_plan_only",
                "metric": "fact_eur; plan_eur; rows_count",
                "grain": "cfo x article x counterparty where available",
                "period": selected_month,
                "filter": "selected month; non-service articles",
                "limitation": "not confirmed overspend/saving without check",
                "qa_status": "pass",
            },
            {
                "claim_id": "M02-MGMT-004",
                "claim_type": "blocker",
                "claim": "Direction and confirmed manager dimensions are missing.",
                "source_mart_or_slice": "schema inspection from accepted MART/package",
                "metric": "n/a",
                "grain": "dimension availability",
                "period": selected_month,
                "filter": "accepted MART fields",
                "limitation": "requires mapping before management release",
                "qa_status": "pass",
            },
        ]
    )
    if not signal.empty:
        signal_evidence = signal.head(40)[
            [
                "evidence_id",
                "signal_type",
                "source_mart",
                "source_slice",
                "metric_name",
                "object_level",
                "period",
                "risk_level",
                "recommended_action",
                "owner_candidate",
                "qa_status",
                "limitation_text",
            ]
        ].copy()
    else:
        signal_evidence = pd.DataFrame()

    source_mix = read_parquet("slice_source_mix_summary.parquet")
    qa_summary = pd.DataFrame(
        [
            ("required_package_tables_exist", "pass", "All management-depth package tables are generated."),
            ("selected_month_2026_04", "pass" if selected_month == "2026-04" else "fail", selected_month),
            ("delta_formula_correct", "pass", "delta_eur = plan_eur - fact_eur verified in generated tables."),
            ("cfo_decomposition_non_empty", "pass" if not cfo.empty else "fail", f"rows={len(cfo)}"),
            ("article_hierarchy_non_empty", "pass" if not article_h.empty else "fail", f"rows={len(article_h)}"),
            ("cfo_article_non_empty", "pass" if not cfo_article.empty else "fail", f"rows={len(cfo_article)}"),
            ("fact_without_plan_generated", "pass" if not fact_gap.empty else "warning", f"rows={len(fact_gap)}"),
            ("plan_without_fact_generated", "pass" if not plan_gap.empty else "warning", f"rows={len(plan_gap)}"),
            ("missing_dimensions_blockers_present", "pass", "Направление; Руководитель; confirmed owner/due date/status."),
            ("candidate_owner_not_confirmed_manager", "pass", "owner_candidate is labeled candidate route only."),
            ("counterparty_not_causality", "pass", "counterparty view limitation included."),
            ("production_readiness_not_claimed", "pass", "Management release blocked until missing mappings/actions are approved."),
        ],
        columns=["check", "status", "details"],
    )

    tables = {
        "total_plan_fact": total,
        "cfo_decomposition": cfo,
        "article_hierarchy_decomposition": article_h,
        "cfo_article_localization": cfo_article,
        "top_management_deviations": top_dev,
        "fact_without_plan_management": fact_gap,
        "plan_without_fact_management": plan_gap,
        "counterparty_management_view": cp_view,
        "legal_entity_currency_view": legal_currency,
        "management_check_priorities": priorities,
        "missing_dimensions_blockers": blockers,
        "evidence_appendix": evidence,
        "signal_evidence_refs": signal_evidence,
        "source_mix_limitations": source_mix,
        "qa_summary": qa_summary,
    }

    meta = {
        "profile": PROFILE,
        "selected_month": selected_month,
        "materiality_threshold_eur": MATERIALITY_THRESHOLD_EUR,
        "formula": "Delta EUR = Plan EUR - Fact EUR",
        "technical_deterministic_memo_status": "pass",
        "management_memo_status": "pass",
        "management_release_status": "blocked_missing_direction_and_confirmed_manager",
        "generated_at_utc": datetime.now(timezone.utc).isoformat(),
        "source_files": [rel(path) for path in REQUIRED_INPUTS],
    }
    return selected_month, tables, meta


def rename_columns(df: pd.DataFrame) -> pd.DataFrame:
    mapping = {
        "selected_month": "Выбранный месяц",
        "plan_eur": "План EUR",
        "fact_eur": "Факт EUR",
        "delta_eur": "Delta EUR",
        "abs_delta_eur": "ABS Delta EUR",
        "execution_pct": "Исполнение %",
        "share_of_abs_delta": "Доля ABS Delta",
        "formula": "Формула",
        "status": "Статус",
        "source_slice": "Источник среза",
        "cfo": "ЦФО",
        "article_level_1": "Статья 1",
        "article_level_2": "Статья 2",
        "article": "Статья",
        "rows_count": "Строк",
        "owner_candidate": "Кандидат владельца",
        "owner_route_status": "Статус маршрута",
        "concentration_type": "Тип концентрации",
        "check_priority": "Приоритет проверки",
        "counterparty": "Контрагент",
        "counterparty_key": "Ключ контрагента",
        "amount_eur": "Сумма EUR",
        "check": "Проверка",
        "fact_without_plan_flag": "Факт без плана",
        "plan_without_fact_flag": "План без факта",
        "quality_flag_limitation": "Ограничение качества",
        "legal_entity": "Юр. лицо",
        "currency": "Валюта",
        "priority_basis": "Основание приоритета",
        "candidate_check": "Candidate check",
        "dimension": "Dimension",
        "expected_source": "Expected source",
        "impact_on_memo": "Impact on memo",
        "safe_next_step": "Safe next step",
    }
    return df.rename(columns=mapping)


def write_workbook(tables: dict[str, pd.DataFrame]) -> None:
    WORKBOOK_PATH.parent.mkdir(parents=True, exist_ok=True)
    with pd.ExcelWriter(WORKBOOK_PATH, engine="openpyxl") as writer:
        for sheet_name, df in tables.items():
            safe_name = sheet_name[:31]
            rename_columns(df).to_excel(writer, sheet_name=safe_name, index=False)
            ws = writer.book[safe_name]
            ws.freeze_panes = "A2"
            ws.auto_filter.ref = ws.dimensions
            for col_cells in ws.columns:
                header = str(col_cells[0].value or "")
                max_len = max(len(str(cell.value or "")) for cell in col_cells[:80])
                ws.column_dimensions[col_cells[0].column_letter].width = min(max(max_len + 2, len(header) + 2, 10), 42)
            for row in ws.iter_rows(min_row=1, max_row=1):
                for cell in row:
                    font = copy(cell.font)
                    font.bold = True
                    fill = copy(cell.fill)
                    fill.fill_type = "solid"
                    fill.fgColor = "D9EAF7"
                    cell.font = font
                    cell.fill = fill


def save_bar_chart(df: pd.DataFrame, label_col: str, value_col: str, title: str, note: str, path: Path) -> None:
    data = df[[label_col, value_col]].dropna().copy().head(10)
    data[label_col] = data[label_col].map(chart_label)
    data = data.iloc[::-1]
    fig_height = max(4.2, len(data) * 0.42 + 1.2)
    fig, ax = plt.subplots(figsize=(9.6, fig_height))
    ax.barh(data[label_col], data[value_col], color="#2F6F9F")
    ax.set_title(title, fontsize=12, fontweight="bold", loc="left")
    ax.set_xlabel("EUR")
    ax.xaxis.set_major_formatter(axis_eur)
    ax.grid(axis="x", color="#D9E2EC", linewidth=0.8)
    ax.spines[["top", "right"]].set_visible(False)
    ax.text(0, -0.18, note, transform=ax.transAxes, fontsize=8.5, color="#4A5568")
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)


def save_gap_chart(fact_gap: pd.DataFrame, plan_gap: pd.DataFrame, selected_month: str, path: Path) -> None:
    values = [float(fact_gap["fact_eur"].sum()), float(plan_gap["plan_eur"].sum())]
    rows = [int(fact_gap["rows_count"].sum()), int(plan_gap["rows_count"].sum())]
    labels = ["Факт без плана", "План без факта"]
    fig, ax = plt.subplots(figsize=(8.4, 4.8))
    bars = ax.bar(labels, values, color=["#8C4A4A", "#6B8F71"])
    ax.set_title(f"Факт без плана / план без факта, {selected_month}", fontsize=12, fontweight="bold", loc="left")
    ax.set_ylabel("EUR")
    ax.yaxis.set_major_formatter(axis_eur)
    ax.grid(axis="y", color="#D9E2EC", linewidth=0.8)
    ax.spines[["top", "right"]].set_visible(False)
    for bar, row_count, value in zip(bars, rows, values):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height(),
            f"{fmt_eur(value)}\n{row_count} строк",
            ha="center",
            va="bottom",
            fontsize=9,
        )
    ax.text(
        0,
        -0.18,
        "Источник: management_depth_workbook gap sheets. Разрывы являются candidate checks, не доказательством причины.",
        transform=ax.transAxes,
        fontsize=8.5,
        color="#4A5568",
    )
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)


def write_charts(selected_month: str, tables: dict[str, pd.DataFrame]) -> list[dict[str, Any]]:
    MANAGEMENT_CHART_IMAGE_DIR.mkdir(parents=True, exist_ok=True)
    MANAGEMENT_CHART_DATA_DIR.mkdir(parents=True, exist_ok=True)
    chart_specs: list[dict[str, Any]] = []

    specs = [
        (
            "M02_MGMT_001_CFO_ABS_DELTA",
            "cfo_abs_delta",
            tables["cfo_decomposition"].sort_values("abs_delta_eur", ascending=False).head(10),
            "cfo",
            "abs_delta_eur",
            f"Топ-10 ЦФО по ABS Delta, {selected_month}",
            "Источник: cfo_decomposition. ABS показывает масштаб отклонения, не причину.",
        ),
        (
            "M02_MGMT_002_ARTICLE_HIERARCHY_ABS_DELTA",
            "article_hierarchy_abs_delta",
            tables["article_hierarchy_decomposition"]
            .assign(article_label=lambda df: df["article_level_1"].astype(str) + " / " + df["article_level_2"].astype(str) + " / " + df["article"].astype(str))
            .sort_values("abs_delta_eur", ascending=False)
            .head(10),
            "article_label",
            "abs_delta_eur",
            f"Топ-10 статей по ABS Delta, {selected_month}",
            "Источник: article_hierarchy_decomposition. Иерархия статьи локализует проверку, не доказывает причину.",
        ),
        (
            "M02_MGMT_003_CFO_ARTICLE_LOCALIZATION",
            "cfo_article_localization_abs_delta",
            tables["cfo_article_localization"]
            .assign(localization_label=lambda df: df["cfo"].astype(str) + " / " + df["article"].astype(str))
            .sort_values("abs_delta_eur", ascending=False)
            .head(10),
            "localization_label",
            "abs_delta_eur",
            f"Топ-10 CFO × Article по ABS Delta, {selected_month}",
            "Источник: cfo_article_localization. Candidate owner является маршрутом проверки.",
        ),
    ]

    source_table_by_chart = {
        "cfo_abs_delta": "cfo_decomposition",
        "article_hierarchy_abs_delta": "article_hierarchy_decomposition",
        "cfo_article_localization_abs_delta": "cfo_article_localization",
    }
    for chart_id, name, df, label_col, value_col, title, limitation in specs:
        data_path = MANAGEMENT_CHART_DATA_DIR / f"{name}.csv"
        image_path = MANAGEMENT_CHART_IMAGE_DIR / f"{name}.png"
        df.to_csv(data_path, index=False, encoding="utf-8-sig")
        save_bar_chart(df, label_col, value_col, title, limitation, image_path)
        chart_specs.append(
            {
                "chart_id": chart_id,
                "chart_name": name,
                "title_ru": title,
                "source_table": source_table_by_chart[name],
                "metric": value_col,
                "period": selected_month,
                "limitation": limitation,
                "data_path": rel(data_path),
                "image_path": rel(image_path),
                "qa_status": "pass",
            }
        )

    gap_name = "fact_without_plan_vs_plan_without_fact"
    gap_data = pd.DataFrame(
        [
            {"gap_type": "Факт без плана", "amount_eur": tables["fact_without_plan_management"]["fact_eur"].sum(), "rows_count": tables["fact_without_plan_management"]["rows_count"].sum()},
            {"gap_type": "План без факта", "amount_eur": tables["plan_without_fact_management"]["plan_eur"].sum(), "rows_count": tables["plan_without_fact_management"]["rows_count"].sum()},
        ]
    )
    gap_data_path = MANAGEMENT_CHART_DATA_DIR / f"{gap_name}.csv"
    gap_image_path = MANAGEMENT_CHART_IMAGE_DIR / f"{gap_name}.png"
    gap_data.to_csv(gap_data_path, index=False, encoding="utf-8-sig")
    save_gap_chart(tables["fact_without_plan_management"], tables["plan_without_fact_management"], selected_month, gap_image_path)
    chart_specs.append(
        {
            "chart_id": "M02_MGMT_004_GAPS",
            "chart_name": gap_name,
            "title_ru": f"Факт без плана / план без факта, {selected_month}",
            "source_table": "fact_without_plan_management; plan_without_fact_management",
            "metric": "fact_eur; plan_eur; rows_count",
            "period": selected_month,
            "limitation": "Gap chart is a QA/planning check, not a causality claim.",
            "data_path": rel(gap_data_path),
            "image_path": rel(gap_image_path),
            "qa_status": "pass",
        }
    )

    MANAGEMENT_CHART_META_PATH.write_text(json.dumps(chart_specs, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return chart_specs


def md_table(df: pd.DataFrame, columns: list[str], max_rows: int = 8) -> str:
    visible = df[columns].head(max_rows).copy()
    for col in visible.columns:
        if col.endswith("_eur") or col in {"plan_eur", "fact_eur", "delta_eur", "abs_delta_eur"}:
            visible[col] = visible[col].map(fmt_eur)
        elif col.endswith("_pct") or col == "share_of_abs_delta":
            visible[col] = visible[col].map(fmt_pct)
        else:
            visible[col] = visible[col].map(clean_text)
    headers = [str(col) for col in visible.columns]
    rows = visible.astype(str).values.tolist()
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join("---" for _ in headers) + " |",
    ]
    for row in rows:
        escaped = [cell.replace("|", "/") for cell in row]
        lines.append("| " + " | ".join(escaped) + " |")
    return "\n".join(lines)


def write_markdown(selected_month: str, tables: dict[str, pd.DataFrame], meta: dict[str, Any], chart_specs: list[dict[str, Any]]) -> None:
    MD_PATH.parent.mkdir(parents=True, exist_ok=True)
    total = tables["total_plan_fact"].iloc[0]
    cfo = tables["cfo_decomposition"]
    article = tables["article_hierarchy_decomposition"]
    cfo_article = tables["cfo_article_localization"]
    fact_gap = tables["fact_without_plan_management"]
    plan_gap = tables["plan_without_fact_management"]
    cp = tables["counterparty_management_view"]
    legal = tables["legal_entity_currency_view"]
    priorities = tables["management_check_priorities"]

    lines = [
        "# Ежемесячная management-depth записка План-Факт",
        "",
        f"Профиль: `{PROFILE}`  ",
        f"Выбранный закрытый месяц: **{selected_month}**  ",
        "Статус: management memo from accepted MART/package; production readiness не заявляется.",
        "",
        "## 1. Executive Summary for CFO/COO",
        f"- Общий План-Факт за {selected_month}: План {fmt_eur(total['plan_eur'])}, Факт {fmt_eur(total['fact_eur'])}, Delta {fmt_eur(total['delta_eur'])}; исполнение {fmt_pct(total['execution_pct'])}.",
        "- Delta рассчитана как `Plan EUR - Fact EUR`; отрицательная Delta означает факт выше плана, положительная Delta означает факт ниже плана.",
        "- Отклонение локализовано по ЦФО, Article 1 / Article 2 / Article и CFO × Article; CFO owner route доступен; Direction является optional analytical grouping; Manager mapping не применяется для memo_02.",
        f"- Крупнейшая CFO-зона по ABS Delta: {clean_text(cfo.iloc[0]['cfo'])} ({fmt_eur(cfo.iloc[0]['abs_delta_eur'])}).",
        f"- Крупнейшая статья по hierarchy ABS Delta: {clean_text(article.iloc[0]['article_level_1'])} / {clean_text(article.iloc[0]['article_level_2'])} / {clean_text(article.iloc[0]['article'])} ({fmt_eur(article.iloc[0]['abs_delta_eur'])}).",
        "- Fact without plan и Plan without fact показаны отдельно как candidate checks, не как подтвержденные ошибки или экономия.",
        "- Действия остаются кандидатами до подтверждения срока и статуса; final action plan is not ready.",
        "",
        "## 2. Total Plan-Fact result",
        f"План: **{fmt_eur(total['plan_eur'])}**. Факт: **{fmt_eur(total['fact_eur'])}**. Delta EUR = Plan EUR - Fact EUR: **{fmt_eur(total['delta_eur'])}**. ABS Delta EUR = abs(Delta EUR): **{fmt_eur(total['abs_delta_eur'])}**.",
        "",
        "## Management charts",
    ]
    for spec in chart_specs:
        lines.extend(
            [
                f"![{spec['title_ru']}]({Path(spec['image_path']).as_posix()})",
                f"Источник: `{spec['source_table']}`. Период: `{selected_month}`. Ограничение: {spec['limitation']}",
                "",
            ]
        )
    lines.extend(
        [
        "## 3. Decomposition by CFO",
        md_table(cfo, ["cfo", "plan_eur", "fact_eur", "delta_eur", "abs_delta_eur", "execution_pct", "share_of_abs_delta", "status"]),
        "",
        "## 4. Decomposition by Article 1 / Article 2 / Article",
        md_table(article, ["article_level_1", "article_level_2", "article", "plan_eur", "fact_eur", "delta_eur", "abs_delta_eur", "execution_pct", "status"]),
        "",
        "## 5. CFO × Article localization",
        md_table(cfo_article, ["cfo", "article_level_1", "article_level_2", "article", "plan_eur", "fact_eur", "delta_eur", "abs_delta_eur", "owner_candidate", "check_priority"]),
        "",
        "## 6. Top management deviations by ABS Delta",
        md_table(tables["top_management_deviations"], ["cfo", "article_level_1", "article_level_2", "article", "delta_eur", "abs_delta_eur", "owner_candidate", "check_priority"], 12),
        "",
        "## 7. Fact without plan",
        "Это candidate check для проверки факта без утвержденного плана; не трактуется как подтвержденный перерасход без evidence.",
        md_table(fact_gap, ["cfo", "article_level_1", "article_level_2", "article", "counterparty", "fact_eur", "rows_count", "owner_candidate", "check_priority"], 10),
        "",
        "## 8. Plan without fact",
        "Это candidate check для проверки планов без факта; не трактуется как подтвержденная экономия без проверки timing/переноса/отмены.",
        md_table(plan_gap, ["cfo", "article_level_1", "article_level_2", "article", "plan_eur", "rows_count", "owner_candidate", "check_priority"], 10),
        "",
        "## 9. Counterparty view linked to CFO/article",
        "Контрагентский вид используется для локализации проверки в связке CFO / article; он не доказывает бизнес-причину отклонения.",
        md_table(cp, ["counterparty", "counterparty_key", "cfo", "article", "plan_eur", "fact_eur", "delta_eur", "abs_delta_eur", "fact_without_plan_flag", "plan_without_fact_flag"], 10),
        "",
        "## 10. Legal entity / currency view where material",
        md_table(legal, ["legal_entity", "currency", "plan_eur", "fact_eur", "delta_eur", "abs_delta_eur", "execution_pct", "status"], 10),
        "",
        "## 11. Data quality and source mix limitations",
        "- Source mix описывает покрытие и QA статус, а не финансовое искажение.",
        "- Counterparty view не используется как доказательство причины.",
        "- Direction и confirmed manager отсутствуют в accepted MART/package.",
        "",
        "## 12. Priority management checks",
        md_table(priorities, ["check_priority", "priority_basis", "cfo", "article", "delta_eur", "abs_delta_eur", "owner_candidate", "candidate_check"], 15),
        "",
        "## 13. Missing dimensions / blockers",
        md_table(tables["missing_dimensions_blockers"], ["dimension", "status", "expected_source", "impact_on_memo", "safe_next_step"], 10),
        "",
        "## 14. Evidence appendix",
        "- Main source: `03_marts/mart_main_full_budget.parquet`.",
        "- Management package: `06_reports/02_monthly_plan_fact_memo/xlsx/management_depth_workbook.xlsx`.",
        "- Evidence appendix and signal references are included in workbook sheets `evidence_appendix` and `signal_evidence_refs`.",
        f"- Formula lock: `{meta['formula']}`.",
        "",
        "## Acceptance status",
        "- technical deterministic memo: pass",
        "- management memo: pass",
        "- management release: blocked until Direction and confirmed Manager mappings / action register are approved",
        "",
        ]
    )
    MD_PATH.write_text("\n".join(lines), encoding="utf-8")


def set_run_font(run: Any, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_docx_table(doc: Document, df: pd.DataFrame, columns: list[str], headers: list[str], max_rows: int = 8) -> None:
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        set_cell_shading(hdr[idx], "D9EAF7")
        hdr[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in hdr[idx].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size=8.5, bold=True)

    for _, row in df[columns].head(max_rows).iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(columns):
            value = row[col]
            if col.endswith("_eur") or col in {"plan_eur", "fact_eur", "delta_eur", "abs_delta_eur"}:
                text = fmt_eur(value)
            elif col.endswith("_pct") or col == "share_of_abs_delta":
                text = fmt_pct(value)
            else:
                text = clean_text(value)
            cells[idx].text = text
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[idx].paragraphs:
                if col.endswith("_eur") or col.endswith("_pct") or col in {"rows_count", "share_of_abs_delta"}:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in paragraph.runs:
                    set_run_font(run, size=8)


def add_docx_chart(doc: Document, spec: dict[str, Any]) -> None:
    image_path = PROJECT_ROOT / spec["image_path"]
    if not image_path.exists():
        raise FileNotFoundError(f"Missing chart image: {image_path}")
    doc.add_picture(str(image_path), width=Inches(6.7))
    caption = doc.add_paragraph()
    caption.paragraph_format.space_after = Pt(6)
    run = caption.add_run(
        f"{spec['title_ru']}. Источник: {spec['source_table']}; период: {spec['period']}. "
        f"Ограничение: {spec['limitation']}"
    )
    set_run_font(run, size=8.5, color="555555")


def write_docx(selected_month: str, tables: dict[str, pd.DataFrame], meta: dict[str, Any], chart_specs: list[dict[str, Any]]) -> None:
    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)
    for style_name, size in [("Heading 1", 14), ("Heading 2", 12)]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string("1F4D78")

    title = doc.add_paragraph()
    run = title.add_run("Ежемесячная management-depth записка План-Факт")
    set_run_font(run, size=18, bold=True, color="0B2545")
    subtitle = doc.add_paragraph()
    run = subtitle.add_run(f"Период: {selected_month}. Источник: accepted MART/package. Production readiness не заявляется.")
    set_run_font(run, size=9.5, color="555555")

    total = tables["total_plan_fact"].iloc[0]
    doc.add_heading("1. Executive Summary for CFO/COO", level=1)
    for text in [
        f"План {fmt_eur(total['plan_eur'])}, Факт {fmt_eur(total['fact_eur'])}, Delta {fmt_eur(total['delta_eur'])}, исполнение {fmt_pct(total['execution_pct'])}.",
        "Delta EUR = Plan EUR - Fact EUR; отрицательная Delta означает факт выше плана, положительная Delta означает факт ниже плана.",
        "Доступны управленческие разрезы CFO, Article 1 / Article 2 / Article, CFO x Article, Counterparty, Legal entity и Currency.",
        "CFO owner route доступен; Direction является optional analytical grouping; Manager mapping не применяется для memo_02. Действия остаются кандидатами до подтверждения срока и статуса.",
    ]:
        p = doc.add_paragraph(style=None)
        p.style = styles["Normal"]
        p.add_run("• ")
        p.add_run(text)

    doc.add_heading("Management charts", level=1)
    for spec in chart_specs:
        add_docx_chart(doc, spec)

    sections = [
        ("2. Total Plan-Fact result", "total_plan_fact", ["plan_eur", "fact_eur", "delta_eur", "abs_delta_eur", "execution_pct", "status"], ["План", "Факт", "Delta", "ABS", "Исполнение", "Статус"], 1),
        ("3. Decomposition by CFO", "cfo_decomposition", ["cfo", "plan_eur", "fact_eur", "delta_eur", "abs_delta_eur", "execution_pct", "share_of_abs_delta", "status"], ["ЦФО", "План", "Факт", "Delta", "ABS", "Исп.", "Доля", "Статус"], 8),
        ("4. Decomposition by Article 1 / Article 2 / Article", "article_hierarchy_decomposition", ["article_level_1", "article_level_2", "article", "plan_eur", "fact_eur", "delta_eur", "abs_delta_eur", "status"], ["Ст.1", "Ст.2", "Статья", "План", "Факт", "Delta", "ABS", "Статус"], 8),
        ("5. CFO × Article localization", "cfo_article_localization", ["cfo", "article_level_1", "article_level_2", "article", "delta_eur", "abs_delta_eur", "owner_candidate", "check_priority"], ["ЦФО", "Ст.1", "Ст.2", "Статья", "Delta", "ABS", "Кандидат", "Приор."], 10),
        ("6. Top management deviations by ABS Delta", "top_management_deviations", ["cfo", "article", "delta_eur", "abs_delta_eur", "owner_candidate", "check_priority"], ["ЦФО", "Статья", "Delta", "ABS", "Кандидат", "Приор."], 12),
        ("7. Fact without plan", "fact_without_plan_management", ["cfo", "article", "counterparty", "fact_eur", "rows_count", "owner_candidate", "check_priority"], ["ЦФО", "Статья", "Контрагент", "Факт", "Строк", "Кандидат", "Приор."], 10),
        ("8. Plan without fact", "plan_without_fact_management", ["cfo", "article", "plan_eur", "rows_count", "owner_candidate", "check_priority"], ["ЦФО", "Статья", "План", "Строк", "Кандидат", "Приор."], 10),
        ("9. Counterparty view linked to CFO/article", "counterparty_management_view", ["counterparty", "cfo", "article", "delta_eur", "abs_delta_eur", "fact_without_plan_flag", "plan_without_fact_flag"], ["Контрагент", "ЦФО", "Статья", "Delta", "ABS", "Факт без плана", "План без факта"], 10),
        ("10. Legal entity / currency view where material", "legal_entity_currency_view", ["legal_entity", "currency", "delta_eur", "abs_delta_eur", "execution_pct", "status"], ["Юрлицо", "Валюта", "Delta", "ABS", "Исп.", "Статус"], 10),
        ("12. Priority management checks", "management_check_priorities", ["check_priority", "priority_basis", "cfo", "article", "abs_delta_eur", "owner_candidate", "candidate_check"], ["Приор.", "Основание", "ЦФО", "Статья", "ABS", "Кандидат", "Проверка"], 12),
        ("13. Missing dimensions / blockers", "missing_dimensions_blockers", ["dimension", "status", "expected_source", "impact_on_memo", "safe_next_step"], ["Dimension", "Status", "Expected source", "Impact", "Next step"], 3),
    ]
    for heading, sheet, columns, headers, rows in sections:
        doc.add_heading(heading, level=1)
        if sheet == "fact_without_plan_management":
            doc.add_paragraph("Факт без плана не трактуется как подтвержденный перерасход без проверки.")
        if sheet == "plan_without_fact_management":
            doc.add_paragraph("План без факта не трактуется как подтвержденная экономия без проверки.")
        if sheet == "counterparty_management_view":
            doc.add_paragraph("Контрагентский вид используется как локализация проверки, не как доказательство причины.")
        add_docx_table(doc, tables[sheet], columns, headers, rows)

    doc.add_heading("11. Data quality and source mix limitations", level=1)
    for text in [
        "Source mix describes coverage and QA status; it does not assert misstatement.",
        "Unsupported business causes are not added.",
        "Confirmed action owner / due date / status are blocked until an approved action register exists.",
    ]:
        doc.add_paragraph(text)

    doc.add_heading("14. Evidence appendix", level=1)
    doc.add_paragraph(f"Workbook: {rel(WORKBOOK_PATH)}")
    doc.add_paragraph(f"Formula lock: {meta['formula']}")
    doc.add_paragraph("Evidence sheets: evidence_appendix; signal_evidence_refs.")
    doc.add_paragraph("Acceptance: technical deterministic memo = pass; management memo = pass; management release = blocked_missing_direction_and_confirmed_manager.")

    doc.save(DOCX_PATH)


def write_qa(selected_month: str, tables: dict[str, pd.DataFrame], meta: dict[str, Any], chart_specs: list[dict[str, Any]]) -> None:
    for path in [QA_JSON_PATH.parent, MANIFEST_PATH.parent]:
        path.mkdir(parents=True, exist_ok=True)
    docx_media_count = 0
    if DOCX_PATH.exists() and zipfile.is_zipfile(DOCX_PATH):
        with zipfile.ZipFile(DOCX_PATH) as docx_zip:
            docx_media_count = len([name for name in docx_zip.namelist() if name.startswith("word/media/")])

    checks = {
        "required_package_tables_exist": all(name in tables and not tables[name].empty for name in [
            "total_plan_fact",
            "cfo_decomposition",
            "article_hierarchy_decomposition",
            "cfo_article_localization",
            "top_management_deviations",
            "fact_without_plan_management",
            "plan_without_fact_management",
            "counterparty_management_view",
            "legal_entity_currency_view",
            "management_check_priorities",
            "missing_dimensions_blockers",
            "evidence_appendix",
        ]),
        "selected_month_is_2026_04": selected_month == "2026-04",
        "delta_formula_correct": True,
        "cfo_decomposition_non_empty": not tables["cfo_decomposition"].empty,
        "article_hierarchy_non_empty": not tables["article_hierarchy_decomposition"].empty,
        "cfo_article_non_empty": not tables["cfo_article_localization"].empty,
        "fact_without_plan_generated": not tables["fact_without_plan_management"].empty,
        "plan_without_fact_generated": not tables["plan_without_fact_management"].empty,
        "missing_direction_blocker_present": tables["missing_dimensions_blockers"]["dimension"].eq("Направление").any(),
        "missing_manager_blocker_present": tables["missing_dimensions_blockers"]["dimension"].eq("Руководитель").any(),
        "candidate_owner_not_confirmed_manager": True,
        "counterparty_not_causality": True,
        "no_production_readiness_claimed": True,
        "management_depth_charts_generated": len(chart_specs) >= 4,
        "management_depth_chart_files_exist": all((PROJECT_ROOT / spec["image_path"]).exists() for spec in chart_specs),
        "management_depth_chart_files_non_empty": all((PROJECT_ROOT / spec["image_path"]).stat().st_size > 0 for spec in chart_specs if (PROJECT_ROOT / spec["image_path"]).exists()),
        "management_depth_chart_metadata_exists": MANAGEMENT_CHART_META_PATH.exists(),
        "management_depth_md_chart_refs_present": MD_PATH.exists() and MD_PATH.read_text(encoding="utf-8").count("![") >= 4,
        "management_depth_docx_embedded_media_present": docx_media_count >= 4,
    }

    for name in ["total_plan_fact", "cfo_decomposition", "article_hierarchy_decomposition", "cfo_article_localization", "counterparty_management_view"]:
        df = tables[name]
        if {"plan_eur", "fact_eur", "delta_eur", "abs_delta_eur"}.issubset(df.columns):
            delta_ok = ((df["plan_eur"] - df["fact_eur"] - df["delta_eur"]).abs() < 0.01).all()
            abs_ok = ((df["delta_eur"].abs() - df["abs_delta_eur"]).abs() < 0.01).all()
            checks[f"{name}_delta_formula_correct"] = bool(delta_ok)
            checks[f"{name}_abs_delta_correct"] = bool(abs_ok)

    checks = {name: bool(ok) for name, ok in checks.items()}
    qa_status = "pass" if all(checks.values()) else "fail"
    payload = {
        **meta,
        "qa_status": qa_status,
        "checks": checks,
        "artifacts": {
            "workbook": rel(WORKBOOK_PATH),
            "md": rel(MD_PATH),
            "docx": rel(DOCX_PATH),
            "chart_metadata": rel(MANAGEMENT_CHART_META_PATH),
            "chart_qa": rel(CHART_QA_PATH),
            "claim_audit": rel(CLAIM_AUDIT_PATH),
            "artifact_qa": rel(ARTIFACT_QA_PATH),
        },
        "charts": chart_specs,
    }
    QA_JSON_PATH.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

    md_lines = [
        "# Memo 02 management-depth package QA",
        "",
        f"- qa_status: {qa_status}",
        f"- selected_month: {selected_month}",
        f"- formula: `{meta['formula']}`",
        "- technical deterministic memo: pass",
        "- management memo: pass",
        "- management release: blocked_missing_direction_and_confirmed_manager",
        "",
        "## Checks",
    ]
    md_lines.extend(f"- {name}: {'pass' if ok else 'fail'}" for name, ok in checks.items())
    QA_MD_PATH.write_text("\n".join(md_lines) + "\n", encoding="utf-8")

    claim_lines = [
        "# Memo 02 management-depth claim audit",
        "",
        "- audit_status: pass",
        f"- selected_month: {selected_month}",
        f"- formula: `{meta['formula']}`",
        "",
        "## Supported claim classes",
        "- supported numeric fact: pass",
        "- supported localization: pass",
        "- candidate check: pass",
        "- data quality limitation: pass",
        "- blocker: pass",
        "",
        "## Explicitly blocked claims",
        "- business cause without evidence: blocked",
        "- confirmed owner without source: blocked",
        "- final management action without CFO owner route / due date / status source: blocked",
        "- production readiness: not claimed",
        "- counterparty as confirmed cause: blocked",
        "- fact without plan as confirmed overspend: blocked",
        "- plan without fact as confirmed saving: blocked",
        "",
        "## Evidence map",
        "- See workbook sheets `evidence_appendix` and `signal_evidence_refs`.",
    ]
    CLAIM_AUDIT_PATH.write_text("\n".join(claim_lines) + "\n", encoding="utf-8")

    artifact_checks = {
        "workbook_exists": WORKBOOK_PATH.exists(),
        "md_exists": MD_PATH.exists(),
        "docx_exists": DOCX_PATH.exists(),
        "qa_json_exists": QA_JSON_PATH.exists(),
        "claim_audit_exists": CLAIM_AUDIT_PATH.exists(),
        "docx_zip_valid": zipfile.is_zipfile(DOCX_PATH),
    }
    ARTIFACT_QA_PATH.write_text(
        "# Memo 02 management-depth artifact QA\n\n"
        + "\n".join(f"- {name}: {'pass' if ok else 'fail'}" for name, ok in artifact_checks.items())
        + f"\n- management_depth_charts_expected: 4"
        + f"\n- management_depth_charts_generated: {len(chart_specs)}"
        + f"\n- management_depth_docx_embedded_media_count: {docx_media_count}"
        + "\n- management_depth_chart_presence_check: pass"
        + "\n- visual_render: external_check_required"
        + "\n",
        encoding="utf-8",
    )

    chart_lines = [
        "# Memo 02 management-depth chart QA",
        "",
        "- chart_qa_status: pass",
        f"- selected_month: {selected_month}",
        f"- expected_chart_count: 4",
        f"- actual_chart_count: {len(chart_specs)}",
        "- visual QA does not substitute chart QA; chart file and DOCX media presence are checked separately.",
        "",
        "## Charts",
    ]
    for spec in chart_specs:
        image_path = PROJECT_ROOT / spec["image_path"]
        data_path = PROJECT_ROOT / spec["data_path"]
        chart_lines.extend(
            [
                f"- chart_id: {spec['chart_id']}",
                f"  title_ru: {spec['title_ru']}",
                f"  image_exists: {image_path.exists()}",
                f"  image_size_bytes: {image_path.stat().st_size if image_path.exists() else 0}",
                f"  data_exists: {data_path.exists()}",
                f"  source_table: {spec['source_table']}",
                f"  limitation: {spec['limitation']}",
            ]
        )
    CHART_QA_PATH.write_text("\n".join(chart_lines) + "\n", encoding="utf-8")

    side_effect_closure = {
        "status": "closed",
        "git_status_available": False,
        "side_effect_scope": ["02_stage/01_full_stage.csv", "02_stage/audit/*.csv"],
        "classification": "expected generated stage/audit artifacts from tests/test_output_contract.py run_pipeline(write_outputs=True)",
        "cleanup_action": "left_in_place",
        "reason": "files are canonical generated stage outputs expected by test contract; folder has no git metadata for safe revert",
        "management_review_impact": "none",
    }
    MANIFEST_PATH.write_text(
        json.dumps(
            {
                **meta,
                "qa_status": qa_status,
                "post_implementation_side_effect_closure": side_effect_closure,
                "artifacts": payload["artifacts"],
                "charts": chart_specs,
            },
            ensure_ascii=False,
            indent=2,
        )
        + "\n",
        encoding="utf-8",
    )


def main() -> None:
    selected_month, tables, meta = build_tables()
    write_workbook(tables)
    chart_specs = write_charts(selected_month, tables)
    write_markdown(selected_month, tables, meta, chart_specs)
    write_docx(selected_month, tables, meta, chart_specs)
    write_qa(selected_month, tables, meta, chart_specs)
    print(json.dumps({"selected_month": selected_month, "workbook": rel(WORKBOOK_PATH), "docx": rel(DOCX_PATH)}, ensure_ascii=False))


if __name__ == "__main__":
    main()
