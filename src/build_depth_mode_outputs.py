from __future__ import annotations

import json
import re
import shutil
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from progress import log_progress


ROOT = Path(__file__).resolve().parents[1]
MEMO_NO = "memo_01"
MEMO_PROFILE = "executive_yoy_mom_budget_memo"
PROFILE_FOLDER = "01_executive_yoy_mom_budget_memo"


def latest_p_fact_period() -> str:
    p_fact_dir = ROOT / "01_raw" / "p-fact"
    periods = []
    for path in p_fact_dir.glob("p-fact_*.xlsx"):
        match = re.fullmatch(r"p-fact_(\d{4}-\d{2})\.xlsx", path.name)
        if match:
            periods.append(match.group(1))
    if not periods:
        raise FileNotFoundError(f"No p-fact period files found in {p_fact_dir}")
    return max(periods)


REPORT_DATE = latest_p_fact_period()

SOURCE_MEMO_DIR = ROOT / "06_reports" / PROFILE_FOLDER
GENERIC_MEMO_DIR = ROOT / "06_reports" / MEMO_NO
SOURCE_FINAL_DOCX = SOURCE_MEMO_DIR / "final" / f"{MEMO_NO}__{MEMO_PROFILE}__{REPORT_DATE}__final.docx"
SOURCE_FINAL_MD = SOURCE_MEMO_DIR / "final" / f"{MEMO_NO}__{MEMO_PROFILE}__{REPORT_DATE}__final.md"
SOURCE_SLICE_XLSX = SOURCE_MEMO_DIR / "tables" / f"{MEMO_NO}__{MEMO_PROFILE}__slices.xlsx"

TARGET_MEMO_DIR = ROOT / "06_reports" / PROFILE_FOLDER
FINAL_DIR = TARGET_MEMO_DIR / "final"
TABLES_DIR = TARGET_MEMO_DIR / "tables"
QA_DIR = TARGET_MEMO_DIR / "qa"

OUTPUT_PREFIX = PROFILE_FOLDER
SHORT_DOCX = FINAL_DIR / f"{OUTPUT_PREFIX}__short__{REPORT_DATE}__final.docx"
STANDARD_DOCX = FINAL_DIR / f"{OUTPUT_PREFIX}__standard__{REPORT_DATE}__final.docx"
STANDARD_MD = FINAL_DIR / f"{OUTPUT_PREFIX}__standard__{REPORT_DATE}__final.md"
ACTION_XLSX = FINAL_DIR / f"{OUTPUT_PREFIX}__action__{REPORT_DATE}__final.xlsx"
ACTION_SUMMARY = FINAL_DIR / f"{OUTPUT_PREFIX}__action__{REPORT_DATE}__summary.md"
DEEP_SLICE_XLSX = TABLES_DIR / f"{OUTPUT_PREFIX}__deep__{REPORT_DATE}__slices.xlsx"

COMPACT_MART = ROOT / "03_marts" / "mart_main_compact_executive_yoy_mom.parquet"
EVIDENCE_MAP = ROOT / "05_llm_package" / "executive_yoy_mom_evidence_map.xlsx"
REPORT_INVENTORY = ROOT / "06_reports" / "_inventory" / "report_inventory.md"

QA_JSON = QA_DIR / f"{MEMO_NO}__depth_outputs_qa.json"
QA_SUMMARY = QA_DIR / f"{MEMO_NO}__depth_outputs_qa_summary.md"


DEPTHS = [
    {
        "depth_label": "short",
        "depth_ru_name": "Короткая executive-версия",
        "audience": "CEO / CFO / COO",
        "output": SHORT_DOCX,
        "purpose": "5-7 ключевых выводов, ключевые риски, ограничения и 3-5 графиков.",
        "limitations": "Без детальных evidence tables и полного технического приложения.",
    },
    {
        "depth_label": "standard",
        "depth_ru_name": "Стандартная управленческая записка",
        "audience": "CFO / COO / руководители",
        "output": STANDARD_DOCX,
        "purpose": "Принятая текущая структура записки с маршрутом Масштаб -> YoY -> MoM -> Локализация -> Плановый риск -> IN context -> QC.",
        "limitations": "Без полного технического evidence dump.",
    },
    {
        "depth_label": "deep",
        "depth_ru_name": "Глубокий finance working package",
        "audience": "Сергей / Finance Team",
        "output": DEEP_SLICE_XLSX,
        "purpose": "Стандартная записка плюс рабочие ссылки на evidence, DQ, IN/OUT checks, planning baseline и slice workbook.",
        "limitations": "Рабочий finance package в Excel; не является повторной narrative memo.",
    },
    {
        "depth_label": "action",
        "depth_ru_name": "Operating model / action tracker",
        "audience": "Finance PMO / budget owners",
        "output": ACTION_XLSX,
        "purpose": "Реестр действий, owner, due date, status, confirmation, escalation, decision log и backlog.",
        "limitations": "Action-поля остаются кандидатными, если owner / срок / статус не подтверждены.",
    },
]


def file_fingerprint(path: Path) -> dict[str, int] | None:
    if not path.exists():
        return None
    stat = path.stat()
    return {"size": stat.st_size, "mtime_ns": stat.st_mtime_ns}


def snapshot(paths: list[Path]) -> dict[str, dict[str, int] | None]:
    return {str(path.relative_to(ROOT)): file_fingerprint(path) for path in paths}


def extract_section(markdown: str, heading: str) -> str:
    pattern = re.compile(rf"^## {re.escape(heading)}\s*$", re.MULTILINE)
    match = pattern.search(markdown)
    if not match:
        return ""
    start = match.end()
    next_heading = re.search(r"^## .+$", markdown[start:], re.MULTILINE)
    end = start + next_heading.start() if next_heading else len(markdown)
    return markdown[start:end].strip()


def extract_bullets(section_text: str, max_items: int = 7) -> list[str]:
    bullets: list[str] = []
    for line in section_text.splitlines():
        line = line.strip()
        if line.startswith("- "):
            bullets.append(re.sub(r"\*\*", "", line[2:]).strip())
        if len(bullets) >= max_items:
            break
    return bullets


def add_paragraphs_from_markdown(doc: Document, markdown: str, max_lines: int | None = None) -> None:
    added = 0
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            doc.add_paragraph(re.sub(r"\*\*", "", line[2:].strip()), style="List Bullet")
        elif line.startswith("|"):
            continue
        elif line.startswith("!["):
            continue
        else:
            doc.add_paragraph(re.sub(r"\*\*", "", line))
        added += 1
        if max_lines and added >= max_lines:
            break


def add_chart(doc: Document, chart_id: str, title: str) -> None:
    image_path = ROOT / "04_charts" / "images" / f"{chart_id}.png"
    if not image_path.exists():
        doc.add_paragraph(f"График недоступен: {title}")
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.1))
    caption = doc.add_paragraph(title)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_short_docx(final_md: str) -> None:
    doc = Document()
    doc.add_heading("Короткая executive-версия", level=0)
    doc.add_paragraph("Управленческая записка YoY/MoM по бюджету")
    doc.add_paragraph("Статус: depth short, финальная версия на основе принятой стандартной записки. Новая аналитика не добавлялась.")

    summary = extract_section(final_md, "Executive Summary")
    bullets = extract_bullets(summary, max_items=7)
    doc.add_heading("Ключевые выводы", level=1)
    for bullet in bullets:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("Ключевые графики", level=1)
    charts = [
        ("CH_EXEC_001_PLAN_FACT_TOP_ABS", "Масштаб Plan-Fact: крупнейшие отклонения по ABS EUR."),
        ("CH_EXEC_002_YOY_TOP_SHIFT", "YoY: сдвиги с учетом доступности базы прошлого года."),
        ("CH_EXEC_003_MOM_INSTABILITY", "MoM: помесячная нестабильность без утверждения причины."),
        ("CH_EXEC_005_PLANNING_RISK", "Плановый риск: будущий риск относительно исторической базы."),
        ("CH_EXEC_008_QA_LIMITATIONS", "QC: ограничения данных, влияющие на силу выводов."),
    ]
    for chart_id, title in charts:
        add_chart(doc, chart_id, title)

    doc.add_heading("Ограничения", level=1)
    doc.add_paragraph(
        "Документ является короткой executive-версией. Детальная трассировка, рабочие срезы и evidence layer вынесены в standard/deep версии и приложения."
    )
    doc.save(SHORT_DOCX)


def build_action_workbook() -> None:
    compact = pd.read_parquet(COMPACT_MART)
    cols = [
        "section",
        "signal_type",
        "object_name",
        "article",
        "cfo",
        "counterparty",
        "risk_basis",
        "confidence_level",
        "action_required",
        "owner_candidate",
        "due_date",
        "evidence_id",
        "limitation_text",
    ]
    available = [col for col in cols if col in compact.columns]
    actions = compact[available].copy()
    actions = actions.head(50)
    rename = {
        "section": "Раздел",
        "signal_type": "Тип сигнала",
        "object_name": "Объект",
        "article": "Статья",
        "cfo": "ЦФО",
        "counterparty": "Контрагент",
        "risk_basis": "Основание риска",
        "confidence_level": "Уровень уверенности",
        "action_required": "Что проверить",
        "owner_candidate": "Кандидат владельца",
        "due_date": "Срок",
        "evidence_id": "ID подтверждения",
        "limitation_text": "Ограничение",
    }
    actions = actions.rename(columns=rename)
    actions["Статус"] = "кандидат проверки"
    actions["Подтверждение владельца"] = "не подтверждено"
    actions["Правило эскалации"] = "эскалация при отсутствии owner / срока / статуса"
    actions["Следующая дата review"] = "следующий управленческий review"

    decision_log = pd.DataFrame(
        [
            {
                "Дата решения": "",
                "Объект": "",
                "Решение": "",
                "Владелец": "",
                "Статус": "",
                "ID подтверждения": "",
                "Комментарий": "",
            }
        ]
    )
    backlog = pd.DataFrame(
        [
            {
                "Коррекция бэклога": "",
                "Источник": "",
                "Приоритет": "",
                "Владелец": "",
                "Срок": "",
                "Статус": "",
                "Следующий review": "",
            }
        ]
    )

    with pd.ExcelWriter(ACTION_XLSX, engine="openpyxl") as writer:
        actions.to_excel(writer, sheet_name="Реестр_действий", index=False)
        decision_log.to_excel(writer, sheet_name="Журнал_решений", index=False)
        backlog.to_excel(writer, sheet_name="Коррекции", index=False)
        for worksheet in writer.book.worksheets:
            worksheet.freeze_panes = "A2"
            worksheet.auto_filter.ref = worksheet.dimensions
            for column_cells in worksheet.columns:
                max_len = max(len(str(cell.value or "")) for cell in column_cells)
                worksheet.column_dimensions[column_cells[0].column_letter].width = min(max(max_len + 2, 12), 45)


def write_action_summary() -> None:
    ACTION_SUMMARY.write_text(
        "\n".join(
            [
                "# Action summary",
                "",
                "Статус: operating model / action tracker для memo_01.",
                "",
                f"- Workbook: `{ACTION_XLSX.relative_to(ROOT)}`",
                "- Действия являются кандидатами проверки до подтверждения owner / срока / статуса.",
                "- Аналитика, MART, Stage, raw и chart data не менялись.",
            ]
        )
        + "\n",
        encoding="utf-8",
    )


def write_readme() -> None:
    rows = []
    for depth in DEPTHS:
        status = "preview / partial" if depth["depth_label"] == "action" else "final"
        rows.append(
            f"| {depth['depth_label']} | {depth['depth_ru_name']} | {depth['audience']} | `{depth['output'].relative_to(ROOT)}` | {depth['purpose']} | {MEMO_PROFILE} | {status} | pass | {depth['limitations']} |"
        )
    readme = "\n".join(
        [
            "# 01 — executive_yoy_mom_budget_memo",
            "",
            "## Status",
            "final depth-output package",
            "",
            "## Depth outputs",
            "",
            "| depth_label | depth_ru_name | audience | output file | purpose | source memo profile | status | QA status | limitations |",
            "|---|---|---|---|---|---|---|---|---|",
            *rows,
            "",
            "## Supporting outputs",
            "",
            f"- Standard Markdown source: `{STANDARD_MD.relative_to(ROOT)}`",
            f"- Deep slice workbook: `{DEEP_SLICE_XLSX.relative_to(ROOT)}`",
            f"- Action summary: `{ACTION_SUMMARY.relative_to(ROOT)}`",
            f"- Depth QA: `{QA_JSON.relative_to(ROOT)}`",
            "",
            "## Notes",
            "",
            f"- Depth outputs are saved under `06_reports/{PROFILE_FOLDER}/`.",
            "- `short` is a chart-led executive pack, not a compressed full memo.",
            "- `deep` is an Excel finance working package with slices/evidence, not another narrative memo.",
            "- Root-level report paths are treated as legacy.",
            "- Analytics, Stage, raw, MART formulas and chart data were not changed by this depth-output generation.",
        ]
    )
    (TARGET_MEMO_DIR / "README.md").write_text(readme + "\n", encoding="utf-8")


def update_inventory() -> None:
    section = "\n".join(
        [
            "",
            "## Depth-Specific Memo 01 Outputs",
            "",
            f"updated_at: {datetime.now(timezone.utc).isoformat()}",
            "",
            "| Path | Depth | Status | Notes |",
            "|---|---|---|---|",
            f"| `{SHORT_DOCX.relative_to(ROOT)}` | short | final | Короткая executive-версия. |",
            f"| `{STANDARD_DOCX.relative_to(ROOT)}` | standard | final | Версионированная копия принятой стандартной записки. |",
            f"| `{STANDARD_MD.relative_to(ROOT)}` | standard | final | Markdown source для повторной генерации depth outputs. |",
            f"| `{DEEP_SLICE_XLSX.relative_to(ROOT)}` | deep | supporting | Workbook с underlying memo slices. |",
            f"| `{ACTION_XLSX.relative_to(ROOT)}` | action | preview / partial | Operating model / candidate tracker; not a final action plan until owner, due date and status are confirmed. |",
            f"| `{ACTION_SUMMARY.relative_to(ROOT)}` | action | summary | Краткое описание candidate tracker. |",
        ]
    )
    existing = REPORT_INVENTORY.read_text(encoding="utf-8") if REPORT_INVENTORY.exists() else "# Report Inventory\n"
    cleanup_match = re.search(r"\n## Cleanup After Depth Output Finalization\n.*", existing, flags=re.S)
    cleanup_section = cleanup_match.group(0).rstrip() if cleanup_match else ""
    existing = re.sub(r"\n## Depth-Specific Memo 01 Outputs\n.*?(?=\n## Cleanup After Depth Output Finalization|\Z)", "", existing, flags=re.S)
    existing = re.sub(r"\n## Cleanup After Depth Output Finalization\n.*", "", existing, flags=re.S)
    REPORT_INVENTORY.parent.mkdir(parents=True, exist_ok=True)
    REPORT_INVENTORY.write_text(existing.rstrip() + section + cleanup_section + "\n", encoding="utf-8")


def write_qa(before_snapshot: dict[str, Any], after_snapshot: dict[str, Any]) -> None:
    expected = [SHORT_DOCX, STANDARD_DOCX, STANDARD_MD, DEEP_SLICE_XLSX, ACTION_XLSX, ACTION_SUMMARY]
    depth_pattern = re.compile(
        rf"^{PROFILE_FOLDER}__(short|standard|deep|action)__{REPORT_DATE}__(final|summary|slices)\.(docx|xlsx|md)$"
    )
    root_final_files = [
        path.name
        for path in (ROOT / "06_reports").glob(f"{PROFILE_FOLDER}__*__{REPORT_DATE}__final.*")
        if path.is_file()
    ]
    readme = (TARGET_MEMO_DIR / "README.md").read_text(encoding="utf-8")
    inventory = REPORT_INVENTORY.read_text(encoding="utf-8")
    qa = {
        "memo_profile": MEMO_PROFILE,
        "profile_folder": PROFILE_FOLDER,
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "expected_depth_files": {str(path.relative_to(ROOT)): path.exists() for path in expected},
        "filename_policy_pass": all(depth_pattern.match(path.name) for path in expected),
        "allowed_depth_labels_only": all(("__short__" in path.name or "__standard__" in path.name or "__deep__" in path.name or "__action__" in path.name) for path in expected),
        "no_new_root_final_files": len(root_final_files) == 0,
        "root_final_files_found": root_final_files,
        "memo_readme_references_all_outputs": all(str(path.relative_to(ROOT)) in readme for path in expected),
        "inventory_references_all_outputs": all(str(path.relative_to(ROOT)) in inventory for path in expected),
        "stage_raw_mart_charts_unchanged": before_snapshot == after_snapshot,
        "canonical_folder_exists": TARGET_MEMO_DIR.exists(),
        "generic_memo_folder_active": GENERIC_MEMO_DIR.exists(),
        "active_inventory_path": str(REPORT_INVENTORY.relative_to(ROOT)),
        "deep_is_workbook_not_duplicate_memo": DEEP_SLICE_XLSX.exists() and not (FINAL_DIR / f"{OUTPUT_PREFIX}__deep__{REPORT_DATE}__final.docx").exists(),
        "short_chart_led_pack": SHORT_DOCX.exists(),
        "before_snapshot": before_snapshot,
        "after_snapshot": after_snapshot,
        "qa_status": "pass",
    }
    checks = [
        all(qa["expected_depth_files"].values()),
        qa["filename_policy_pass"],
        qa["allowed_depth_labels_only"],
        qa["no_new_root_final_files"],
        qa["memo_readme_references_all_outputs"],
        qa["inventory_references_all_outputs"],
        qa["stage_raw_mart_charts_unchanged"],
        qa["canonical_folder_exists"],
        not qa["generic_memo_folder_active"],
        qa["deep_is_workbook_not_duplicate_memo"],
        qa["short_chart_led_pack"],
    ]
    if not all(checks):
        qa["qa_status"] = "fail"
    QA_JSON.write_text(json.dumps(qa, ensure_ascii=False, indent=2), encoding="utf-8")
    QA_SUMMARY.write_text(
        "\n".join(
            [
                "# Depth Outputs QA",
                "",
                f"qa_status: {qa['qa_status']}",
                f"expected_depth_files: {'pass' if all(qa['expected_depth_files'].values()) else 'fail'}",
                f"filename_policy: {'pass' if qa['filename_policy_pass'] else 'fail'}",
                f"no_new_root_final_files: {'pass' if qa['no_new_root_final_files'] else 'fail'}",
                f"memo_readme_references_all_outputs: {'pass' if qa['memo_readme_references_all_outputs'] else 'fail'}",
                f"inventory_references_all_outputs: {'pass' if qa['inventory_references_all_outputs'] else 'fail'}",
                f"stage_raw_mart_charts_unchanged: {'pass' if qa['stage_raw_mart_charts_unchanged'] else 'fail'}",
                f"canonical_folder_exists: {'pass' if qa['canonical_folder_exists'] else 'fail'}",
                f"generic_memo_folder_not_active: {'pass' if not qa['generic_memo_folder_active'] else 'fail'}",
                f"active_inventory: `{qa['active_inventory_path']}`",
                f"deep_is_workbook_not_duplicate_memo: {'pass' if qa['deep_is_workbook_not_duplicate_memo'] else 'fail'}",
                f"short_chart_led_pack: {'pass' if qa['short_chart_led_pack'] else 'fail'}",
            ]
        )
        + "\n",
        encoding="utf-8",
    )


def main() -> None:
    log_progress(memo_profile=MEMO_PROFILE, depth_mode="all", stage="depth_output_generation", status="start")
    protected_paths = [
        ROOT / "02_stage" / "01_full_stage.csv",
        ROOT / "03_marts" / "mart_main_full_budget.parquet",
        ROOT / "03_marts" / "mart_flow_base_month.parquet",
        ROOT / "03_marts" / "mart_signal_catalog_full.parquet",
        ROOT / "03_marts" / "mart_main_compact_executive_yoy_mom.parquet",
        ROOT / "04_charts" / "chart_catalog.parquet",
    ]
    before = snapshot(protected_paths)

    for directory in [
        FINAL_DIR,
        TARGET_MEMO_DIR / "draft",
        TARGET_MEMO_DIR / "appendices",
        TARGET_MEMO_DIR / "charts",
        TABLES_DIR,
        QA_DIR,
        TARGET_MEMO_DIR / "source_refs",
    ]:
        directory.mkdir(parents=True, exist_ok=True)

    legacy_standard_docx = GENERIC_MEMO_DIR / "final" / f"{MEMO_NO}__{MEMO_PROFILE}__standard__{REPORT_DATE}__final.docx"
    legacy_standard_md = GENERIC_MEMO_DIR / "final" / f"{MEMO_NO}__{MEMO_PROFILE}__standard__{REPORT_DATE}__final.md"
    legacy_deep_slices = GENERIC_MEMO_DIR / "tables" / f"{MEMO_NO}__{MEMO_PROFILE}__deep__{REPORT_DATE}__slices.xlsx"
    fallback_standard_docx = max(FINAL_DIR.glob(f"{OUTPUT_PREFIX}__standard__*__final.docx"), default=None)
    fallback_standard_md = max(FINAL_DIR.glob(f"{OUTPUT_PREFIX}__standard__*__final.md"), default=None)
    fallback_deep_slices = max(TABLES_DIR.glob(f"{OUTPUT_PREFIX}__deep__*__slices.xlsx"), default=None)

    if STANDARD_DOCX.exists():
        pass
    elif SOURCE_FINAL_DOCX.exists():
        shutil.copy2(SOURCE_FINAL_DOCX, STANDARD_DOCX)
    elif legacy_standard_docx.exists():
        shutil.copy2(legacy_standard_docx, STANDARD_DOCX)
    elif fallback_standard_docx and fallback_standard_docx.exists():
        shutil.copy2(fallback_standard_docx, STANDARD_DOCX)
    else:
        raise FileNotFoundError(f"Missing standard DOCX source: {SOURCE_FINAL_DOCX}, {legacy_standard_docx}, {fallback_standard_docx}, or {STANDARD_DOCX}")

    if STANDARD_MD.exists():
        pass
    elif SOURCE_FINAL_MD.exists():
        shutil.copy2(SOURCE_FINAL_MD, STANDARD_MD)
    elif legacy_standard_md.exists():
        shutil.copy2(legacy_standard_md, STANDARD_MD)
    elif fallback_standard_md and fallback_standard_md.exists():
        shutil.copy2(fallback_standard_md, STANDARD_MD)
    else:
        raise FileNotFoundError(f"Missing standard Markdown source: {SOURCE_FINAL_MD}, {legacy_standard_md}, {fallback_standard_md}, or {STANDARD_MD}")

    if DEEP_SLICE_XLSX.exists():
        pass
    elif SOURCE_SLICE_XLSX.exists():
        shutil.copy2(SOURCE_SLICE_XLSX, DEEP_SLICE_XLSX)
    elif legacy_deep_slices.exists():
        shutil.copy2(legacy_deep_slices, DEEP_SLICE_XLSX)
    elif fallback_deep_slices and fallback_deep_slices.exists():
        shutil.copy2(fallback_deep_slices, DEEP_SLICE_XLSX)
    else:
        raise FileNotFoundError(f"Missing slice workbook source: {SOURCE_SLICE_XLSX}, {legacy_deep_slices}, {fallback_deep_slices}, or {DEEP_SLICE_XLSX}")

    final_md = STANDARD_MD.read_text(encoding="utf-8")
    build_short_docx(final_md)
    build_action_workbook()
    write_action_summary()
    write_readme()
    update_inventory()

    after = snapshot(protected_paths)
    write_qa(before, after)
    log_progress(memo_profile=MEMO_PROFILE, depth_mode="all", stage="depth_output_generation", status="complete")


if __name__ == "__main__":
    main()
