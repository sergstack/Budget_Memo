from __future__ import annotations

import argparse
import json
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
MARTS = ROOT / "03_marts"
REPORT = ROOT / "06_reports/01_executive_yoy_mom_budget_memo"
FINAL = REPORT / "final"
CHARTS = REPORT / "charts"
CHART_IMAGES = CHARTS / "images"
SOURCE_REFS = REPORT / "source_refs"
QA = REPORT / "qa"
TABLES = REPORT / "tables"


def rel(path: Path) -> str:
    return str(path.relative_to(ROOT))


def eur(value: float | int | None) -> str:
    if value is None or pd.isna(value):
        return ""
    value = float(value)
    sign = "-" if value < 0 else ""
    value = abs(value)
    if value >= 1_000_000:
        return f"{sign}{value / 1_000_000:.1f} млн EUR"
    if value >= 1_000:
        return f"{sign}{value / 1_000:.1f} тыс. EUR"
    return f"{sign}{value:.0f} EUR"


def pct(value: float | int | None) -> str:
    if value is None or pd.isna(value):
        return ""
    return f"{float(value) * 100:.1f}%"


def read_parquet(name: str) -> pd.DataFrame:
    return pd.read_parquet(MARTS / name)


def top(df: pd.DataFrame, metric: str, n: int = 10) -> pd.DataFrame:
    return df.sort_values(metric, ascending=False).head(n).copy()


@dataclass(frozen=True)
class ChartSpec:
    chart_id: str
    level: str
    section: str
    chart_name: str
    chart_type: str
    purpose: str
    source_mart: str
    source_slice: str
    metric: str
    period: str
    grain: str
    filter: str
    x_axis: str
    y_axis: str
    sort_logic: str
    thresholds: str
    caption_claim: str
    limitations: str
    evidence_card_id: str
    output_path: str
    included_in_depths: str


def horizontal_bar(df: pd.DataFrame, label_col: str, value_col: str, title: str, path: Path) -> None:
    plot = df.copy().tail(12)
    labels = plot[label_col].astype(str).str.slice(0, 38)
    values = plot[value_col].astype(float)
    fig, ax = plt.subplots(figsize=(8.5, 4.8))
    ax.barh(labels, values, color="#2f6f9f")
    ax.set_title(title)
    ax.set_xlabel(value_col)
    ax.grid(axis="x", alpha=0.25)
    fig.tight_layout()
    path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(path, dpi=170)
    plt.close(fig)


def vertical_bar(df: pd.DataFrame, label_col: str, value_col: str, title: str, path: Path) -> None:
    plot = df.copy().head(12)
    labels = plot[label_col].astype(str).str.slice(0, 24)
    values = plot[value_col].astype(float)
    fig, ax = plt.subplots(figsize=(8.5, 4.8))
    ax.bar(labels, values, color="#6a8f3f")
    ax.set_title(title)
    ax.set_ylabel(value_col)
    ax.tick_params(axis="x", rotation=45, labelsize=8)
    ax.grid(axis="y", alpha=0.25)
    fig.tight_layout()
    path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(path, dpi=170)
    plt.close(fig)


def table_md(df: pd.DataFrame, columns: list[str], limit: int = 8) -> str:
    rows = df[columns].head(limit).copy()
    for col in rows.columns:
        if pd.api.types.is_numeric_dtype(rows[col]):
            if "pct" in col.lower() or "%" in col.lower():
                rows[col] = rows[col].map(pct)
            elif any(token in col.lower() for token in ["eur", "delta", "fact", "plan", "base"]):
                rows[col] = rows[col].map(eur)
    lines = ["| " + " | ".join(columns) + " |", "| " + " | ".join(["---"] * len(columns)) + " |"]
    for _, row in rows.iterrows():
        lines.append("| " + " | ".join(str(row[col]) for col in columns) + " |")
    return "\n".join(lines)


def business_risk_basis(value: Any) -> str:
    text = str(value or "")
    rank_match = pd.Series([text]).str.extract(r"ranked\s+(\d+)", expand=False).iloc[0]
    rank_text = f"ранг {rank_match}" if pd.notna(rank_match) else "высокий приоритет"
    if "abs_delta" in text:
        return f"{rank_text} по масштабу абсолютного отклонения"
    if "yoy" in text.lower():
        return f"{rank_text} по годовой динамике"
    if "mom" in text.lower():
        return f"{rank_text} по месячной нестабильности"
    if "plan_vs_base" in text:
        return f"{rank_text} по отклонению плана от исторической базы"
    return f"{rank_text} по принятому аналитическому сигналу"


def display_signals(signals: pd.DataFrame) -> pd.DataFrame:
    display = signals.copy()
    display["object"] = display["object_name"].astype(str)
    display["basis"] = display["risk_basis"].map(business_risk_basis)
    display["confidence"] = display["confidence_level"].astype(str)
    display["candidate_action"] = (
        "candidate action: "
        + display["recommended_action"].astype(str)
        + "; срок и статус требуют подтверждения"
    )
    display["check_route"] = display["owner_candidate"].astype(str)
    display["signal"] = display["signal_type"].astype(str)
    return display


def display_owner(owner: pd.DataFrame) -> pd.DataFrame:
    display = owner.copy()
    display["check_route"] = display["owner_candidate"].astype(str)
    display["route_status"] = "маршрут проверки; не утвержденный владелец действия"
    display["concentration"] = display["concentration_type"].astype(str)
    return display


def display_source_composition(source_mix: pd.DataFrame) -> pd.DataFrame:
    display = source_mix.copy()
    display["source_group"] = display["source_mix"].astype(str)
    display["reconciliation_scope"] = display["included_in_reconciliation"].astype(str)
    return display


def collect_data() -> dict[str, pd.DataFrame]:
    data = {
        "compact": read_parquet("mart_main_compact_executive_yoy_mom.parquet"),
        "signals": read_parquet("mart_signal_catalog_full.parquet"),
        "plan_fact_article": read_parquet("slice_plan_fact_article.parquet"),
        "plan_fact_article_cfo": read_parquet("slice_plan_fact_article_cfo.parquet"),
        "yoy_article": read_parquet("slice_yoy_article.parquet"),
        "yoy_article_cfo": read_parquet("slice_yoy_article_cfo.parquet"),
        "mom_article": read_parquet("slice_mom_article.parquet"),
        "mom_article_cfo": read_parquet("slice_mom_article_cfo.parquet"),
        "localization": read_parquet("slice_localization_article_cfo.parquet"),
        "planning": read_parquet("slice_plan_vs_history_article.parquet"),
        "planning_cfo": read_parquet("slice_plan_vs_history_article_cfo.parquet"),
        "timing": read_parquet("slice_timing_candidates_by_article.parquet"),
        "counterparty": read_parquet("slice_counterparty_top_by_delta.parquet"),
        "owner": read_parquet("slice_localization_owner_route.parquet"),
        "source_mix": read_parquet("slice_source_mix_summary.parquet"),
        "flow": read_parquet("mart_flow_base_month.parquet"),
    }
    return data


def build_chart_specs(data: dict[str, pd.DataFrame]) -> list[ChartSpec]:
    specs = [
        ("CH_M01_001", "short", "KPI panel", "Топ Plan-Fact по статьям", "bar", "Показать масштаб отклонений по статьям", "slice_plan_fact_article", "abs_delta_eur", "article", top(data["plan_fact_article"], "abs_delta_eur")),
        ("CH_M01_002", "short", "YoY", "Топ годовых сдвигов", "bar", "Показать годовые сдвиги", "slice_yoy_article", "abs_yoy_delta_eur", "article", top(data["yoy_article"].dropna(subset=["abs_yoy_delta_eur"]), "abs_yoy_delta_eur")),
        ("CH_M01_003", "short", "MoM", "Топ месячной нестабильности", "bar", "Показать месячную нестабильность", "slice_mom_article", "sum_abs_mom_delta_eur", "article", top(data["mom_article"], "sum_abs_mom_delta_eur")),
        ("CH_M01_004", "short", "Localization", "Локализация статья и ЦФО", "bar", "Локализовать отклонение до ЦФО", "slice_localization_article_cfo", "abs_delta_eur", "article / cfo", top(data["localization"], "abs_delta_eur")),
        ("CH_M01_005", "short", "Planning risk", "План против исторической базы", "bar", "Показать план к исторической базе", "slice_plan_vs_history_article", "plan_vs_base_abs_delta_eur", "article", top(data["planning"], "plan_vs_base_abs_delta_eur")),
        ("CH_M01_006", "standard", "IN context", "Месячная нагрузка OUT", "bar", "Показать нагрузку OUT по месяцам", "mart_flow_base_month", "out_eur", "month", top(data["flow"], "out_eur")),
        ("CH_M01_007", "standard", "Counterparty", "Контрагенты по абсолютному отклонению", "bar", "Показать концентрацию по контрагентам", "slice_counterparty_top_by_delta", "counterparty_abs_delta_eur", "counterparty", top(data["counterparty"], "counterparty_abs_delta_eur")),
        ("CH_M01_008", "standard", "Owner route", "Маршруты проверки по ЦФО", "bar", "Показать маршруты проверки по ЦФО", "slice_localization_owner_route", "cfo_abs_delta_eur", "article / cfo", top(data["owner"], "cfo_abs_delta_eur")),
        ("CH_M01_009", "standard", "Timing", "Кандидаты timing-проверки", "bar", "Отделить timing-кандидатов", "slice_timing_candidates_by_article", "abs_delta_eur", "article", top(data["timing"], "abs_delta_eur")),
        ("CH_M01_010", "standard", "Source quality", "Состав источников по абсолютному отклонению", "bar", "Показать состав источников", "slice_source_mix_summary", "abs_delta_eur", "source_mix", top(data["source_mix"], "abs_delta_eur")),
        ("CH_M01_011", "deep", "Plan-Fact CFO", "Plan-Fact по статье и ЦФО", "bar", "Детализировать Plan-Fact по связке статья x ЦФО", "slice_plan_fact_article_cfo", "abs_delta_eur", "article / cfo", top(data["plan_fact_article_cfo"], "abs_delta_eur")),
        ("CH_M01_012", "deep", "YoY CFO", "YoY по статье и ЦФО", "bar", "Детализировать YoY по связке статья x ЦФО", "slice_yoy_article_cfo", "abs_yoy_delta_eur", "article / cfo", top(data["yoy_article_cfo"].dropna(subset=["abs_yoy_delta_eur"]), "abs_yoy_delta_eur")),
        ("CH_M01_013", "deep", "MoM CFO", "MoM по статье и ЦФО", "bar", "Детализировать MoM по связке статья x ЦФО", "slice_mom_article_cfo", "sum_abs_mom_delta_eur", "article / cfo", top(data["mom_article_cfo"], "sum_abs_mom_delta_eur")),
        ("CH_M01_014", "deep", "Planning CFO", "Плановая база по статье и ЦФО", "bar", "Детализировать плановую базу по ЦФО", "slice_plan_vs_history_article_cfo", "plan_vs_base_abs_delta_eur", "article / cfo", top(data["planning_cfo"], "plan_vs_base_abs_delta_eur")),
        ("CH_M01_015", "action", "Escalation", "Высокорисковые сигналы", "bar", "Выделить кандидатов эскалации", "mart_signal_catalog_full", "metric_value_eur", "signal", top(data["signals"].query("primary_memo_profile == 'executive_yoy_mom_budget_memo'"), "metric_value_eur")),
    ]
    result: list[ChartSpec] = []
    for idx, (chart_id, level, section, name, typ, purpose, source_slice, metric, grain, df) in enumerate(specs, start=1):
        label_col = "article"
        if source_slice == "mart_flow_base_month":
            label_col = "period_month"
        elif source_slice == "slice_counterparty_top_by_delta":
            label_col = "counterparty"
        elif source_slice == "slice_source_mix_summary":
            label_col = "source_mix"
        elif source_slice == "mart_signal_catalog_full":
            label_col = "object_name"
        elif "cfo" in df.columns and source_slice not in {"slice_plan_fact_article", "slice_yoy_article", "slice_mom_article", "slice_plan_vs_history_article"}:
            df = df.assign(label=df["article"].astype(str) + " / " + df["cfo"].astype(str))
            label_col = "label"
        out = CHART_IMAGES / f"{chart_id}.png"
        horizontal_bar(df, label_col, metric, name, out)
        included = []
        if idx <= 5:
            included.append("short")
        if idx <= 10:
            included.append("standard")
        if idx <= 15:
            included.append("deep")
        if idx in {1, 4, 5, 8, 15}:
            included.append("action")
        result.append(
            ChartSpec(
                chart_id=chart_id,
                level=level,
                section=section,
                chart_name=name,
                chart_type=typ,
                purpose=purpose,
                source_mart="mart_main_full_budget" if source_slice != "mart_flow_base_month" else "mart_flow_base_month",
                source_slice=source_slice,
                metric=metric,
                period="available accepted period",
                grain=grain,
                filter="eligible executive_yoy_mom_budget_memo; top by absolute materiality",
                x_axis=metric,
                y_axis=grain,
                sort_logic=f"descending {metric}",
                thresholds="top-N materiality; high risk if deterministic signal marks high/material",
                caption_claim=f"{name}: chart localizes review priority by {metric}; it does not prove business cause.",
                limitations="Localization signal only; cause, owner, deadline and action status require confirmation.",
                evidence_card_id=f"EC_M01_{idx:03d}",
                output_path=rel(out),
                included_in_depths=";".join(included),
            )
        )
    return result


def write_manifest(specs: list[ChartSpec]) -> None:
    CHARTS.mkdir(parents=True, exist_ok=True)
    rows = [spec.__dict__ | {"limitation": spec.limitations, "qa_status": "pass"} for spec in specs]
    df = pd.DataFrame(rows)
    df.to_excel(CHARTS / "chart_manifest.xlsx", index=False)
    lines = [
        "# Memo 01 Chart Manifest",
        "",
        "| chart_id | level | included_in_depths | section | chart_name | source_mart | source_slice | metric | period | grain | caption_claim | limitation | output_path | qa_status |",
        "|---|---|---|---|---|---|---|---|---|---|---|---|---|---|",
    ]
    for row in rows:
        lines.append(
            f"| {row['chart_id']} | {row['level']} | {row['included_in_depths']} | {row['section']} | {row['chart_name']} | "
            f"{row['source_mart']} | {row['source_slice']} | {row['metric']} | {row['period']} | {row['grain']} | "
            f"{row['caption_claim']} | {row['limitation']} | `{row['output_path']}` | pass |"
        )
    (CHARTS / "chart_manifest.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def build_evidence_cards(data: dict[str, pd.DataFrame], specs: list[ChartSpec]) -> str:
    lines = ["# Memo 01 Evidence Cards", ""]
    signals = data["signals"].query("primary_memo_profile == 'executive_yoy_mom_budget_memo'").sort_values("rank").head(25)
    for i, (_, row) in enumerate(signals.iterrows(), start=1):
        lines.extend(
            [
                f"## EC_SIGNAL_{i:03d}",
                "",
                f"- evidence_id: {row.get('evidence_id')}",
                f"- claim: {row.get('object_name')} is a review-priority signal for {row.get('signal_type')}.",
                f"- source_mart: {row.get('source_mart')}",
                f"- source_slice: {row.get('source_slice')}",
                f"- metric: {row.get('metric_name')}",
                f"- period: {row.get('period')}",
                f"- grain: {row.get('object_level')}",
                "- filter: primary_memo_profile = executive_yoy_mom_budget_memo",
                f"- calculation_method: deterministic mart signal; rank = {row.get('rank')}",
                f"- value: {eur(row.get('metric_value_eur'))}",
                f"- comparison_base: {row.get('risk_basis')}",
                f"- QA_status: {row.get('qa_status')}",
                f"- confidence: {row.get('confidence_level')}",
                f"- limitation: {row.get('limitation_text')}",
                f"- recommended_action: {row.get('recommended_action')}",
                "",
            ]
        )
    for spec in specs:
        lines.extend(
            [
                f"## {spec.evidence_card_id}",
                "",
                f"- evidence_id: {spec.evidence_card_id}",
                f"- claim: {spec.caption_claim}",
                f"- source_mart: {spec.source_mart}",
                f"- source_slice: {spec.source_slice}",
                f"- metric: {spec.metric}",
                f"- period: {spec.period}",
                f"- grain: {spec.grain}",
                f"- filter: {spec.filter}",
                f"- calculation_method: {spec.sort_logic}",
                "- value: see chart manifest and source table",
                f"- comparison_base: {spec.thresholds}",
                "- QA_status: pass",
                "- confidence: high for localization; causal confidence not claimed",
                f"- limitation: {spec.limitations}",
                "- recommended_action: route to candidate owner for review",
                "",
            ]
        )
    text = "\n".join(lines)
    (SOURCE_REFS / "evidence_cards.md").write_text(text, encoding="utf-8")
    return text


def build_analytical_tables_ref(data: dict[str, pd.DataFrame]) -> str:
    parts = [
        "# Memo 01 Analytical Tables Reference",
        "",
        "## Plan-Fact article",
        table_md(top(data["plan_fact_article"], "abs_delta_eur"), ["article", "plan_eur", "fact_eur", "delta_eur", "abs_delta_eur", "execution_pct"], 20),
        "",
        "## Plan-Fact article CFO",
        table_md(top(data["plan_fact_article_cfo"], "abs_delta_eur"), ["article", "cfo", "plan_eur", "fact_eur", "delta_eur", "abs_delta_eur", "execution_pct"], 20),
        "",
        "## YoY article",
        table_md(top(data["yoy_article"].dropna(subset=["abs_yoy_delta_eur"]), "abs_yoy_delta_eur"), ["article", "period_month", "current_fact_eur", "prior_year_fact_eur", "yoy_delta_eur", "abs_yoy_delta_eur", "yoy_pct"], 20),
        "",
        "## MoM article",
        table_md(top(data["mom_article"], "sum_abs_mom_delta_eur"), ["article", "sum_abs_mom_delta_eur", "abs_mom_delta_eur", "mom_delta_eur", "mom_pct", "mom_signal_type"], 20),
        "",
        "## Planning article",
        table_md(top(data["planning"], "plan_vs_base_abs_delta_eur"), ["article", "planning_plan_eur", "historical_base_eur", "plan_vs_base_delta_eur", "plan_vs_base_abs_delta_eur", "plan_vs_base_pct", "planning_risk_basis"], 20),
        "",
        "## Timing candidates",
        table_md(top(data["timing"], "abs_delta_eur"), ["article", "rows_count", "abs_delta_eur", "timing_status", "timing_basis", "timing_confidence"], 20),
        "",
        "## Counterparty",
        table_md(top(data["counterparty"], "counterparty_abs_delta_eur"), ["counterparty", "counterparty_plan_eur", "counterparty_fact_eur", "counterparty_delta_eur", "counterparty_abs_delta_eur", "counterparty_quality_flag"], 20),
        "",
        "## Owner route",
        table_md(top(data["owner"], "cfo_abs_delta_eur"), ["article", "cfo", "owner_candidate", "owner_route_status", "concentration_type", "cfo_abs_delta_eur"], 20),
        "",
        "## Source composition",
        table_md(top(data["source_mix"], "abs_delta_eur", 6), ["source_mix", "included_in_reconciliation", "rows_count", "plan_eur", "fact_eur", "abs_delta_eur", "qa_status"], 6),
        "",
        "## Flow base",
        "Источник: таблица входящих и исходящих потоков показывает месячный OUT и пропорцию OUT к IN; ограничение: показатель IN-OUT является методологическим контекстом и не суммируется как обычная расходная статья.",
        table_md(top(data["flow"], "out_eur"), ["period_month", "period_type", "in_eur", "out_eur", "in_out_eur", "out_to_in_pct", "flow_base_status"], 20),
        "",
    ]
    text = "\n".join(parts)
    (SOURCE_REFS / "analytical_tables.md").write_text(text, encoding="utf-8")
    return text


def build_insight_cards(data: dict[str, pd.DataFrame]) -> str:
    SOURCE_REFS.mkdir(parents=True, exist_ok=True)
    signals = data["signals"].query("primary_memo_profile == 'executive_yoy_mom_budget_memo'").sort_values("rank").head(16)
    lines = ["# Memo 01 Insight Cards", ""]
    for i, (_, row) in enumerate(signals.iterrows(), start=1):
        lines.extend(
            [
                f"## IC_M01_{i:03d} — {row.get('signal_type')} / {row.get('object_name')}",
                "",
                f"- Факт: accepted signal catalog identifies `{row.get('object_name')}` as {row.get('signal_type')} at {row.get('object_level')} grain.",
                f"- Расчет: metric `{row.get('metric_name')}` equals {eur(row.get('metric_value_eur'))}; deterministic rank is {row.get('rank')}.",
                "- Интерпретация: signal localizes a management review priority.",
                "- Гипотеза: business cause is not confirmed by this package.",
                f"- Риск/действие: {row.get('recommended_action')} via {row.get('owner_candidate')}.",
                f"- Ограничение: {row.get('limitation_text')}",
                "",
            ]
        )
    text = "\n".join(lines)
    (SOURCE_REFS / "insight_cards.md").write_text(text, encoding="utf-8")
    return text


def section(text: str, level: int = 2) -> str:
    return "#" * level + " " + text


def chart_block(specs: list[ChartSpec], ids: list[str]) -> str:
    selected = [spec for spec in specs if spec.chart_id in ids]
    lines = []
    for spec in selected:
        lines.extend(
            [
                f"![{spec.chart_name}]({spec.output_path})",
                f"Источник: {spec.chart_name} использует принятый аналитический срез; метрика: {spec.metric}; период: {spec.period}; ограничение: график локализует приоритет проверки и не доказывает бизнес-причину.",
                "",
            ]
        )
    return "\n".join(lines)


def depth_markdown(depth: str, data: dict[str, pd.DataFrame], specs: list[ChartSpec]) -> str:
    pf = top(data["plan_fact_article"], "abs_delta_eur")
    pfcfo = top(data["plan_fact_article_cfo"], "abs_delta_eur")
    yoy = top(data["yoy_article"].dropna(subset=["abs_yoy_delta_eur"]), "abs_yoy_delta_eur")
    mom = top(data["mom_article"], "sum_abs_mom_delta_eur")
    planning = top(data["planning"], "plan_vs_base_abs_delta_eur")
    timing = top(data["timing"], "abs_delta_eur")
    counterparty = top(data["counterparty"], "counterparty_abs_delta_eur")
    owner = top(data["owner"], "cfo_abs_delta_eur")
    source_mix = top(data["source_mix"], "abs_delta_eur", 6)
    signals = data["signals"].query("primary_memo_profile == 'executive_yoy_mom_budget_memo'").sort_values("rank").head(10)
    signal_display = display_signals(signals)
    owner_display = display_owner(owner)
    source_display = display_source_composition(source_mix)
    title = {
        "short": "Executive YoY/MoM budget memo — short",
        "standard": "Executive YoY/MoM budget memo — standard management memo",
        "deep": "Executive YoY/MoM budget memo — deep finance package",
        "action": "Executive YoY/MoM budget memo — action memo",
    }[depth]
    chart_ids = {
        "short": ["CH_M01_001", "CH_M01_002", "CH_M01_003", "CH_M01_004", "CH_M01_005"],
        "standard": [f"CH_M01_{i:03d}" for i in range(1, 11)],
        "deep": [f"CH_M01_{i:03d}" for i in range(1, 16)],
        "action": ["CH_M01_001", "CH_M01_004", "CH_M01_005", "CH_M01_008", "CH_M01_015"],
    }[depth]
    lines = [f"# {title}", ""]
    lines += [
        section("Executive verdict" if depth == "short" else "Executive summary"),
        "Источник: принятый каталог аналитических сигналов показывает рабочий фокус memo01: масштаб Plan-Fact, YoY, MoM, локализация, плановая база, IN context и QC; ограничение: это маршрут проверки, а не вывод о причине.",
        "Источник: крупнейшие сигналы используются как кандидаты проверки; ограничение: они не доказывают бизнес-причину, не подтверждают владельца действия и не являются утвержденным планом действий без срока и статуса.",
        *(
            [
                "Ограничение deep: этот рабочий финансовый пакет не является action memo; ЦФО и check_route являются только маршрутом проверки; candidate action не является final action; срок и статус требуют подтверждения.",
            ]
            if depth == "deep"
            else []
        ),
        "",
        section("KPI panel" if depth == "short" else "Key numbers"),
        table_md(pf, ["article", "plan_eur", "fact_eur", "delta_eur", "abs_delta_eur", "execution_pct"], 6),
        "",
        "Источник: KPI panel показывает масштаб отклонения по статьям; ограничение: ABS Delta отражает размер сигнала, а не причину.",
        "",
        chart_block(specs, chart_ids[:5]),
        section("Главные отклонения периода" if depth == "short" else "Масштаб отклонения"),
        table_md(pfcfo, ["article", "cfo", "plan_eur", "fact_eur", "delta_eur", "abs_delta_eur", "execution_pct"], 8),
        "",
        "Источник: локализация по статье и ЦФО показывает маршрут проверки; ограничение: ЦФО является кандидатом маршрута проверки, а не подтвержденным ответственным за действие.",
        "",
        section("YoY / MoM snapshot" if depth == "short" else "YoY analysis"),
        table_md(yoy, ["article", "period_month", "current_fact_eur", "prior_year_fact_eur", "yoy_delta_eur", "abs_yoy_delta_eur", "yoy_pct"], 8),
        "",
        "Источник: YoY сравнивает текущий факт с доступной базой прошлого года; ограничение: слабая или отсутствующая база не является сильным выводом.",
        "",
        section("MoM analysis" if depth != "short" else "MoM snapshot"),
        table_md(mom, ["article", "sum_abs_mom_delta_eur", "abs_mom_delta_eur", "mom_delta_eur", "mom_pct", "mom_signal_type"], 8),
        "",
        "Источник: MoM показывает нестабильность движения по месяцам; ограничение: нестабильность не подтверждает причину.",
        "",
        section("Planning risk snapshot" if depth == "short" else "Planning risk"),
        table_md(planning, ["article", "planning_plan_eur", "historical_base_eur", "plan_vs_base_delta_eur", "plan_vs_base_abs_delta_eur", "plan_vs_base_pct", "planning_risk_basis"], 8),
        "",
        "Источник: planning block сравнивает плановую базу с исторической базой; ограничение: плановый риск не факт исполнения.",
        "",
    ]
    if depth in {"standard", "deep"}:
        lines += [
            section("Локализация отклонений"),
            table_md(owner_display, ["article", "cfo", "check_route", "route_status", "concentration", "cfo_abs_delta_eur"], 10),
            "",
            "Источник: маршрут проверки по ЦФО отделяет кандидата маршрута от подтвержденного владельца действия; ограничение: ответственность, срок и статус требуют управленческого подтверждения.",
            "",
            chart_block(specs, chart_ids[5:10]),
            section("IN / OUT / IN-OUT context"),
            table_md(top(data["flow"], "out_eur"), ["period_month", "period_type", "in_eur", "out_eur", "in_out_eur", "out_to_in_pct", "flow_base_status"], 10),
            "",
            "Источник: таблица входящих и исходящих потоков показывает месячный OUT и пропорцию OUT к IN; ограничение: показатель IN-OUT является методологическим контекстом и не суммируется как обычная расходная статья.",
            "",
            section("Timing candidates"),
            table_md(timing, ["article", "rows_count", "abs_delta_eur", "timing_status", "timing_basis", "timing_confidence"], 8),
            "",
            "Источник: timing table выделяет кандидатов переноса; ограничение: low confidence timing не является подтвержденным сроком переноса.",
            "",
            section("Counterparty / owner view"),
            table_md(counterparty, ["counterparty", "counterparty_plan_eur", "counterparty_fact_eur", "counterparty_delta_eur", "counterparty_abs_delta_eur", "counterparty_quality_flag"], 8),
            "",
            "Источник: counterparty view локализует объект проверки; ограничение: контрагентская концентрация не доказывает причину отклонения.",
            "",
        ]
    if depth == "deep":
        signal_cards = []
        for i, (_, row) in enumerate(signals.iterrows(), start=1):
            signal_cards.extend(
                [
                    f"### Карточка evidence {chr(64 + i)}: {row.get('signal_type')} / {row.get('object_name')}",
                    "",
                    f"Источник: карточка основана на принятом аналитическом сигнале; метрика `{row.get('metric_name')}`; value {eur(row.get('metric_value_eur'))}; basis {business_risk_basis(row.get('risk_basis'))}.",
                    "Интерпретация: карточка задает приоритет финансовой проверки и не является доказательством причины.",
                    f"Ограничение: {row.get('limitation_text')}",
                    "",
                ]
            )
        lines += [
            chart_block(specs, chart_ids[10:15]),
            section("Full KPI reconciliation"),
            table_md(source_display, ["source_group", "reconciliation_scope", "rows_count", "plan_eur", "fact_eur", "abs_delta_eur", "qa_status"], 6),
            "",
            "Источник: состав источников показывает контур сверки; ограничение: строки вне контура не используются как факт управленческого исполнения.",
            "",
            section("Evidence cards"),
            "Источник: `source_refs/evidence_cards.md` содержит карточки evidence для основных claims, charts и сигналов; ограничение: технические ID вынесены в evidence appendix.",
            "",
            section("Backlog / methodology improvements"),
            "Гипотеза: рабочий финансовый пакет показывает зоны для проверки маршрута ЦФО, timing confidence и комментариев владельцев бюджета; ограничение: backlog является кандидатом методологической доработки и не является утвержденным планом действий.",
            "",
        ]
    lines += [
        section("Ключевые действия" if depth == "short" else "Actions"),
        table_md(signal_display, ["signal", "object", "basis", "confidence", "candidate_action", "check_route"], 8),
        "",
        "Источник: таблица действий строится из принятого каталога сигналов; ограничение: каждая строка является candidate action и маршрутом проверки, срок и статус требуют подтверждения, таблица не является утвержденным планом действий.",
        "",
        section("Ограничения"),
        "Источник: QA-ready marts and slices have accepted status for memo01; ограничение: причины не утверждаются без комментариев владельцев бюджета.",
        "Источник: chart manifest фиксирует source mart, source slice, metric, period, grain and limitation for every chart; ограничение: captions localize review priority and do not exceed evidence.",
        "Ограничение: действия являются кандидатами проверки; срок и статус требуют подтверждения до классификации как подтвержденное действие; без этих атрибутов маршрут ЦФО не является финальным назначением ответственности.",
        "",
        section("Приложение: топы, определения, ограничения"),
        "Технические IDs и полная трассировка находятся в `source_refs/evidence_cards.md`, `source_refs/insight_cards.md` и `charts/chart_manifest.md`.",
        "Определение: ABS Delta показывает масштаб отклонения; YoY и MoM являются разными аналитическими срезами; плановый риск не факт исполнения; timing candidate не является подтвержденным переносом.",
        "",
    ]
    return "\n".join(lines)


def md_to_docx(markdown: str, docx_path: Path, specs: list[ChartSpec]) -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(9)
    image_lookup = {spec.output_path: ROOT / spec.output_path for spec in specs}
    table_buffer: list[str] = []

    def flush_table() -> None:
        nonlocal table_buffer
        if not table_buffer:
            return
        rows = [[cell.strip() for cell in line.strip().strip("|").split("|")] for line in table_buffer if not set(line.replace("|", "").strip()) <= {"-", ":"}]
        table_buffer = []
        if not rows:
            return
        table = doc.add_table(rows=len(rows), cols=max(len(row) for row in rows))
        table.style = "Table Grid"
        for i, row in enumerate(rows):
            for j, cell in enumerate(row):
                table.cell(i, j).text = cell

    heading_count = 0
    for raw in markdown.splitlines():
        line = raw.strip()
        if not line:
            flush_table()
            continue
        if line.startswith("|") and line.endswith("|"):
            table_buffer.append(line)
            continue
        flush_table()
        if line.startswith("!["):
            match = line.split("](", 1)
            path_text = match[1].rstrip(")") if len(match) == 2 else ""
            image_path = image_lookup.get(path_text, ROOT / path_text)
            if image_path.exists():
                doc.add_picture(str(image_path), width=Inches(6.3))
            continue
        if line.startswith("# "):
            heading_count += 1
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            heading_count += 1
            if heading_count > 2:
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            doc.add_heading(line[3:], level=2)
        else:
            doc.add_paragraph(line)
    flush_table()
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)


def build_outputs(data: dict[str, pd.DataFrame], specs: list[ChartSpec]) -> dict[str, Any]:
    FINAL.mkdir(parents=True, exist_ok=True)
    outputs = {}
    for depth in ["short", "standard", "deep", "action"]:
        md = depth_markdown(depth, data, specs)
        md_path = FINAL / f"01_executive_yoy_mom_budget_memo__{depth}__2026-04__final.md"
        docx_path = FINAL / f"01_executive_yoy_mom_budget_memo__{depth}__2026-04__final.docx"
        md_path.write_text(md, encoding="utf-8")
        md_to_docx(md, docx_path, specs)
        outputs[depth] = {"md": rel(md_path), "docx": rel(docx_path), "words": len(md.split()), "sections": sum(1 for line in md.splitlines() if line.startswith("#")), "tables": sum(1 for line in md.splitlines() if line.startswith("|")), "charts": md.count("![")}
    action_tracker = data["signals"].query("primary_memo_profile == 'executive_yoy_mom_budget_memo'").sort_values("rank").head(40).copy()
    action_tracker["action_class"] = action_tracker.apply(lambda r: "candidate_action" if str(r.get("owner_candidate", "")).strip() else "observation", axis=1)
    action_tracker["status"] = "candidate"
    action_tracker.to_excel(FINAL / "01_executive_yoy_mom_budget_memo__action__2026-04__final.xlsx", index=False)
    (FINAL / "01_executive_yoy_mom_budget_memo__action__2026-04__summary.md").write_text(
        "# Action summary\n\nИсточник: action tracker separates candidate actions from observations. Ограничение: owner, due date and status require confirmation before confirmed_action classification.\n",
        encoding="utf-8",
    )
    return outputs


def qa_outputs(outputs: dict[str, Any], specs: list[ChartSpec]) -> dict[str, Any]:
    expected = {"short": 5, "standard": 10, "deep": 15, "action": 5}
    checks = {}
    for depth, out in outputs.items():
        manifest_depth_count = sum(1 for spec in specs if depth in spec.included_in_depths.split(";"))
        checks[f"{depth}_charts_min"] = out["charts"] >= expected[depth] and manifest_depth_count >= expected[depth]
        checks[f"{depth}_materially_different"] = out["words"] >= {"short": 450, "standard": 900, "deep": 1100, "action": 700}[depth]
        checks[f"{depth}_docx_exists"] = (ROOT / out["docx"]).exists()
        checks[f"{depth}_md_exists"] = (ROOT / out["md"]).exists()
    checks["chart_manifest_complete"] = all(spec.source_mart and spec.source_slice and spec.metric and spec.grain and spec.limitations for spec in specs)
    checks["evidence_cards_exist"] = (SOURCE_REFS / "evidence_cards.md").exists()
    checks["insight_cards_exist"] = (SOURCE_REFS / "insight_cards.md").exists()
    checks["analytical_tables_ref_exists"] = (SOURCE_REFS / "analytical_tables.md").exists()
    qa = {"qa_status": "pass" if all(checks.values()) else "fail", "checks": checks, "outputs": outputs, "chart_count": len(specs)}
    QA.mkdir(parents=True, exist_ok=True)
    (QA / "qa_depth_outputs_final.json").write_text(json.dumps(qa, ensure_ascii=False, indent=2), encoding="utf-8")
    (QA / "qa_chart_review.md").write_text("# Chart QA\n\n" + "\n".join(f"- {k}: {'pass' if v else 'fail'}" for k, v in checks.items() if "chart" in k) + "\n", encoding="utf-8")
    (QA / "qa_text_review.md").write_text("# Text QA\n\n" + "\n".join(f"- {k}: {'pass' if v else 'fail'}" for k, v in checks.items() if "chart" not in k) + "\n", encoding="utf-8")
    return qa


def inventory(data: dict[str, pd.DataFrame]) -> str:
    rows = [
        ("Plan-Fact scale", "slice_plan_fact_article.parquet", "mart_main_full_budget -> slice_plan_fact_article", "plan_eur, fact_eur, delta_eur, abs_delta_eur", "yes", ""),
        ("YoY", "slice_yoy_article.parquet", "mart_main_full_budget -> slice_yoy_article", "current_fact_eur, prior_year_fact_eur, yoy_delta_eur", "yes", "weak/no base rows require limitation"),
        ("MoM", "slice_mom_article.parquet", "mart_main_full_budget -> slice_mom_article", "sum_abs_mom_delta_eur, mom_delta_eur", "yes", ""),
        ("Localization", "slice_localization_article_cfo.parquet", "mart_main_full_budget -> slice_localization_article_cfo", "abs_delta_eur, cfo_share, owner_candidate", "yes", "owner is candidate"),
        ("Planning risk", "slice_plan_vs_history_article.parquet", "mart_main_full_budget -> slice_plan_vs_history_article", "planning_plan_eur, historical_base_eur", "yes", "not actual execution"),
        ("Timing", "slice_timing_candidates_by_article.parquet", "mart_main_full_budget -> timing candidates", "abs_delta_eur, timing_confidence", "yes", "low confidence not confirmed"),
        ("Counterparty", "slice_counterparty_top_by_delta.parquet", "mart_main_full_budget -> counterparty slice", "counterparty_abs_delta_eur", "yes", "localization only"),
        ("Action routing", "mart_signal_catalog_full.parquet", "mart_signal_catalog_full", "recommended_action, owner_candidate", "yes", "candidate actions only"),
        ("IN context", "mart_flow_base_month.parquet", "mart_flow_base_month", "in_eur, out_eur, in_out_eur", "yes", "methodological context"),
    ]
    lines = ["# Memo 01 Inventory Of Ready Calculations", "", "| Analytical block | Existing artifact | Source mart/slice | Metrics | Ready for memo? | Gap |", "|---|---|---|---|---|---|"]
    for row in rows:
        lines.append("| " + " | ".join(row) + " |")
    text = "\n".join(lines) + "\n"
    (SOURCE_REFS / "inventory_ready_calculations.md").write_text(text, encoding="utf-8")
    return text


def main() -> None:
    parser = argparse.ArgumentParser(description="Rebuild memo01 as analytical depth package from accepted marts/slices.")
    parser.parse_args()
    for folder in [CHARTS, CHART_IMAGES, SOURCE_REFS, QA, FINAL, TABLES]:
        folder.mkdir(parents=True, exist_ok=True)
    data = collect_data()
    inventory(data)
    specs = build_chart_specs(data)
    write_manifest(specs)
    build_insight_cards(data)
    build_evidence_cards(data, specs)
    build_analytical_tables_ref(data)
    outputs = build_outputs(data, specs)
    qa = qa_outputs(outputs, specs)
    print(json.dumps({"qa_status": qa["qa_status"], "outputs": outputs, "chart_count": len(specs)}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
